from __future__ import annotations

import math
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional, Tuple

try:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt  # type: ignore

    MATPLOTLIB_AVAILABLE = True
except Exception:  # pragma: no cover - env specific
    MATPLOTLIB_AVAILABLE = False
    plt = None  # type: ignore

try:
    from docx import Document  # type: ignore
    from docx.shared import Inches  # type: ignore

    DOCX_AVAILABLE = True
except Exception:  # pragma: no cover - env specific
    DOCX_AVAILABLE = False
    Document = None  # type: ignore
    Inches = None  # type: ignore

TEMP_DIR = Path("temp")
OUTPUT_DIR = Path("outputs")


def ensure_output_dirs(
    temp_dir: Path = TEMP_DIR,
    output_dir: Path = OUTPUT_DIR,
) -> Tuple[Path, Path]:
    """Create temp/output directories if missing."""
    temp_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)
    return temp_dir, output_dir


def _draw_block(ax, center: Tuple[float, float], text: str, width: float = 1.4, height: float = 0.6):
    x, y = center
    rect = plt.Rectangle((x - width / 2, y - height / 2), width, height, fill=False, linewidth=2, color="#23496b")
    ax.add_patch(rect)
    ax.text(x, y, text, ha="center", va="center", fontsize=9, color="#23496b")


def _safe_stage_value(stage13: Dict[str, Any], key: str, default: float = 0.0) -> float:
    raw = stage13.get(key, default)
    try:
        return float(raw)
    except Exception:
        return default


def render_block_sld(
    stage13: Dict[str, Any],
    *,
    temp_dir: Path = TEMP_DIR,
    output_dir: Path = OUTPUT_DIR,
    filename: str = "block_sld.png",
) -> Tuple[Path, Path]:
    """Render a lightweight block-level SLD and save to temp/output folders."""
    if not MATPLOTLIB_AVAILABLE:
        raise RuntimeError("matplotlib is required for SLD rendering but is not installed.")

    temp_dir, output_dir = ensure_output_dirs(temp_dir, output_dir)

    poi_voltage = _safe_stage_value(stage13, "poi_nominal_voltage_kv", 0.0)
    dc_blocks = int(stage13.get("dc_block_total_qty", 0) or stage13.get("container_count", 0) or 1)
    ac_blocks = max(1, int((dc_blocks + 3) // 4))
    eff_dc_to_poi = _safe_stage_value(stage13, "eff_dc_to_poi_frac", 0.0) * 100

    fig, ax = plt.subplots(figsize=(9, 4.2))
    ax.axis("off")

    _draw_block(ax, (0, 0), f"POI / MV Bus\n{poi_voltage:.1f} kV")
    _draw_block(ax, (2.2, 0), "MV/LV\nTransformer")
    _draw_block(ax, (4.4, 0), f"PCS\n{ac_blocks}× blocks")
    _draw_block(ax, (6.6, 0.9), f"DC Busbar A\nDC Blocks per AC ≈ {math.ceil(dc_blocks / ac_blocks / 2)}")
    _draw_block(ax, (6.6, -0.9), f"DC Busbar B\nDC Blocks per AC ≈ {math.floor(dc_blocks / ac_blocks / 2)}")
    _draw_block(ax, (8.8, 0.9), f"DC Blocks\n{dc_blocks} total")
    _draw_block(ax, (8.8, -0.9), f"RTE Chain\n{eff_dc_to_poi:.1f}% est.")

    xs = [0.7, 2.9, 5.1, 7.3, 8.3]
    for i in range(len(xs) - 1):
        ax.arrow(xs[i], 0, xs[i + 1] - xs[i] - 0.6, 0, width=0.04, head_width=0.18, length_includes_head=True, color="#5cc3e4")

    temp_path = temp_dir / filename
    output_path = output_dir / filename
    fig.savefig(temp_path, dpi=160, bbox_inches="tight")
    fig.savefig(output_path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return temp_path, output_path


def render_layout_plot(
    stage13: Dict[str, Any],
    *,
    temp_dir: Path = TEMP_DIR,
    output_dir: Path = OUTPUT_DIR,
    filename: str = "layout.png",
) -> Tuple[Path, Path]:
    """Plot a simple block layout and save to temp/output folders."""
    if not MATPLOTLIB_AVAILABLE:
        raise RuntimeError("matplotlib is required for layout plotting but is not installed.")

    temp_dir, output_dir = ensure_output_dirs(temp_dir, output_dir)
    containers = int(stage13.get("container_count", 0) or 0)
    cabinets = int(stage13.get("cabinet_count", 0) or 0)
    total = max(containers + cabinets, 1)

    cols = max(1, int(math.ceil(math.sqrt(total))))
    rows = int(math.ceil(total / cols))

    xs, ys, colors, labels = [], [], [], []
    for idx in range(total):
        r = idx // cols
        c = idx % cols
        xs.append(c)
        ys.append(rows - r)
        if idx < containers:
            colors.append("#23496b")
            labels.append("Container")
        else:
            colors.append("#5cc3e4")
            labels.append("Cabinet")

    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.scatter(xs, ys, c=colors, s=240, edgecolor="#1e1e1e", linewidth=1.2)
    for i, (x, y) in enumerate(zip(xs, ys)):
        ax.text(x, y, labels[i][0], ha="center", va="center", color="#ffffff", fontsize=10, weight="bold")

    ax.set_xticks(range(cols))
    ax.set_yticks(range(1, rows + 1))
    ax.set_xlim(-0.6, cols - 0.4)
    ax.set_ylim(0.4, rows + 0.6)
    ax.set_title(f"Site Layout Preview – {containers} containers / {cabinets} cabinets")
    ax.grid(True, linestyle="--", alpha=0.35)

    temp_path = temp_dir / filename
    output_path = output_dir / filename
    fig.savefig(temp_path, dpi=160, bbox_inches="tight")
    fig.savefig(output_path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return temp_path, output_path


def _resolve_report_name(stage13: Dict[str, Any], filename: Optional[str]) -> str:
    if filename:
        return filename
    project = stage13.get("project_name", "CALB_ESS_Project").replace(" ", "_")
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{project}_Stage4_Report_{stamp}.docx"


def generate_stage4_report(
    stage13: Dict[str, Any],
    *,
    sld_path: Optional[Path] = None,
    layout_path: Optional[Path] = None,
    output_dir: Path = OUTPUT_DIR,
    filename: Optional[str] = None,
) -> Path:
    """Generate a DOCX summary and save it to the outputs directory."""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is required for report generation but is not installed.")

    _, output_dir = ensure_output_dirs(output_dir=output_dir)

    doc = Document()
    project_name = stage13.get("project_name", "CALB ESS Project")
    doc.add_heading("Stage 4 AC Block Summary", level=1)
    doc.add_paragraph(f"Project: {project_name}")
    doc.add_paragraph(f"Selected scenario: {stage13.get('selected_scenario', 'unknown')}")
    doc.add_paragraph(f"POI Requirement: {stage13.get('poi_power_req_mw', 0.0):.2f} MW / {stage13.get('poi_energy_req_mwh', 0.0):.2f} MWh")
    doc.add_paragraph(
        f"Containers: {int(stage13.get('container_count', 0))} | "
        f"Cabinets: {int(stage13.get('cabinet_count', 0))} | "
        f"Busbars: {int(stage13.get('busbars_needed', 0))}"
    )
    doc.add_paragraph(
        f"DC Blocks Total: {int(stage13.get('dc_block_total_qty', 0))} | "
        f"DC Nameplate @BOL: {stage13.get('dc_nameplate_bol_mwh', 0.0):.2f} MWh | "
        f"Oversize: {stage13.get('oversize_mwh', 0.0):.2f} MWh"
    )

    if sld_path and sld_path.exists():
        doc.add_heading("Block-Level SLD", level=2)
        doc.add_picture(str(sld_path), width=Inches(6.4))

    if layout_path and layout_path.exists():
        doc.add_heading("Layout Preview", level=2)
        doc.add_picture(str(layout_path), width=Inches(6.4))

    report_name = _resolve_report_name(stage13, filename)
    docx_path = output_dir / report_name
    doc.save(docx_path)
    return docx_path
