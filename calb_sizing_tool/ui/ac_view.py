from __future__ import annotations
import importlib.util
import io
import math
from typing import Any, Dict, List, Optional

import graphviz
import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle
import streamlit as st
import pandas as pd
import numpy as np

# --- 核心修改: 引入配置和包内模块 ---
from calb_sizing_tool.config import AC_DATA_PATH

# 尝试导入 simulation，如果未迁移成功则容错处理
try:
    from calb_sizing_tool.sizing import simulation
except ImportError:
    simulation = None

# -----------------------------------------------------
# Global Constants & Checks
# -----------------------------------------------------
AC_BLOCK_CANDIDATES: List[Dict[str, float]] = [
    {"pcs_units": 2, "pcs_unit_kw": 1250, "ac_block_mw": 2.5},
    {"pcs_units": 2, "pcs_unit_kw": 1725, "ac_block_mw": 3.45},
    {"pcs_units": 4, "pcs_unit_kw": 1250, "ac_block_mw": 5.0},
    {"pcs_units": 4, "pcs_unit_kw": 1725, "ac_block_mw": 6.9},
]

DOCX_AVAILABLE = False
if importlib.util.find_spec("docx"):
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True

# -----------------------------------------------------
# Helper Functions (Logic)
# -----------------------------------------------------

def find_ac_block_container_only(
    poi_mw: float,
    container_cnt: int,
    *,
    candidates: List[Dict[str, float]] = AC_BLOCK_CANDIDATES,
    search_extra: int = 40
) -> Optional[Dict[str, Any]]:
    """
    Try container-only DC Blocks to match AC Blocks.
    """
    if container_cnt <= 0:
        return None

    best = None
    best_score = None

    for cand in candidates:
        p_ac = cand["ac_block_mw"]
        n_min = max(1, int(math.ceil(poi_mw / p_ac)))

        for ac_qty in range(n_min, n_min + search_extra + 1):
            if container_cnt % ac_qty != 0:
                continue
            dc_per_block = container_cnt // ac_qty
            # Exclude exactly 3 DC Blocks per AC Block (design rule)
            if dc_per_block == 3:
                continue
            total_ac = ac_qty * p_ac
            if total_ac < poi_mw:
                continue
            oversize = total_ac - poi_mw
            score = (oversize, ac_qty)
            if best_score is None or score < best_score:
                best_score = score
                best = {
                    "strategy": "container_only",
                    "ac_block_qty": ac_qty,
                    "ac_block_rated_mw": p_ac,
                    "pcs_units": cand["pcs_units"],
                    "pcs_unit_kw": cand["pcs_unit_kw"],
                    "dc_blocks_per_block": dc_per_block,
                    "total_ac_mw": total_ac,
                    "oversize_mw": oversize,
                }
    return best


def find_ac_block_mixed(
    poi_mw: float,
    container_cnt: int,
    cabinet_cnt: int,
    *,
    candidates: List[Dict[str, float]] = AC_BLOCK_CANDIDATES,
    search_extra: int = 40
) -> Optional[Dict[str, Any]]:
    """
    Try mixed (container + cabinet) DC Blocks to match AC Blocks.
    """
    dc_total = container_cnt + cabinet_cnt
    if dc_total <= 0:
        return None

    best = None
    best_score = None

    for cand in candidates:
        p_ac = cand["ac_block_mw"]
        n_min = max(1, int(math.ceil(poi_mw / p_ac)))

        for ac_qty in range(n_min, n_min + search_extra + 1):
            cont_per_block = container_cnt // ac_qty
            cab_per_block = cabinet_cnt // ac_qty
            cont_rem = container_cnt % ac_qty
            cab_rem = cabinet_cnt % ac_qty
            base_dc_each = cont_per_block + cab_per_block
            if base_dc_each == 0 and (cont_rem + cab_rem) == 0:
                continue
            max_dc_each = base_dc_each + (1 if cont_rem > 0 or cab_rem > 0 else 0)
            # Exclude any configuration with exactly 3 DC per block
            if base_dc_each == 3 or max_dc_each == 3:
                continue
            total_dc_calc = (cont_per_block * ac_qty + cont_rem) + (cab_per_block * ac_qty + cab_rem)
            if total_dc_calc != dc_total:
                continue
            total_ac = ac_qty * p_ac
            if total_ac < poi_mw:
                continue
            oversize = total_ac - poi_mw
            score = (oversize, max_dc_each - base_dc_each, ac_qty)
            if best_score is None or score < best_score:
                best_score = score
                best = {
                    "strategy": "mixed",
                    "ac_block_qty": ac_qty,
                    "ac_block_rated_mw": p_ac,
                    "pcs_units": cand["pcs_units"],
                    "pcs_unit_kw": cand["pcs_unit_kw"],
                    "container_per_block": cont_per_block,
                    "cabinet_per_block": cab_per_block,
                    "container_rem": cont_rem,
                    "cabinet_rem": cab_rem,
                    "dc_blocks_per_block_base": base_dc_each,
                    "dc_blocks_per_block_max": max_dc_each,
                    "total_ac_mw": total_ac,
                    "oversize_mw": oversize,
                }
    return best


def get_dc_block_total(stage13: Dict[str, Any]) -> int:
    return int(
        stage13.get("dc_block_total_qty")
        or stage13.get("dc_total_blocks")
        or (stage13.get("container_count", 0) + stage13.get("cabinet_count", 0))
    )


def build_block_sld(dc_block_count: int, res: Dict[str, Any]) -> graphviz.Digraph:
    ac_label = (
        "AC Block\n"
        f"{res.get('ac_block_rated_mw', 0):.2f} MW\n"
        f"{res.get('pcs_units', 0)} × {res.get('pcs_unit_kw', 0)} kW"
    )

    dot = graphviz.Digraph("block_sld")
    dot.attr(rankdir="LR", splines="ortho", nodesep="0.65", ranksep="0.9")
    dot.attr("node", shape="box", style="rounded,filled", fillcolor="#E8F1FF")

    dot.node("ac_block", ac_label)
    dot.node("xfmr", "Transformer\n(MV/LV)")
    dot.node("pcs", "PCS")
    dot.node("ac_bus", "AC Busbar")

    dot.edge("ac_block", "xfmr")
    dot.edge("xfmr", "pcs")
    dot.edge("pcs", "ac_bus")

    for idx in range(dc_block_count or 1):
        name = f"DC Block {idx + 1}" if dc_block_count else "DC Block"
        node_id = f"dc_{idx + 1}" if dc_block_count else "dc"
        dot.node(node_id, name)
        dot.edge("ac_bus", node_id)

    return dot


def compute_layout_positions(dc_block_count: int, *, columns: int = 4, spacing: float = 2.4) -> Dict[str, Any]:
    positions: Dict[str, Any] = {}

    span = (max(columns, 1) - 1) * spacing
    center_x = span / 2

    positions["AC Block"] = (center_x, spacing * 2.4)
    positions["Transformer"] = (center_x - spacing, spacing * 1.6)
    positions["PCS"] = (center_x + spacing, spacing * 1.6)
    positions["AC Busbar"] = (center_x, spacing * 0.9)

    for idx in range(dc_block_count):
        row = idx // max(columns, 1)
        col = idx % max(columns, 1)
        positions[f"DC Block {idx + 1}"] = (col * spacing, -row * spacing)

    return positions


def render_layout_plot(positions: Dict[str, Any], *, ac_color: str = "#CFE2FF", dc_color: str = "#E2F6CF"):
    fig, ax = plt.subplots(figsize=(10, 6))

    def draw_box(name: str, xy: tuple[float, float], color: str) -> None:
        width, height = 1.5, 0.9
        rect = Rectangle((xy[0] - width / 2, xy[1] - height / 2), width, height, facecolor=color, edgecolor="#4B5563")
        ax.add_patch(rect)
        ax.text(xy[0], xy[1], name, ha="center", va="center", fontsize=9)

    for name, xy in positions.items():
        if name.startswith("DC Block"):
            draw_box(name, xy, dc_color)
        else:
            draw_box(name, xy, ac_color)

    bus_xy = positions.get("AC Busbar")
    if bus_xy:
        for name, xy in positions.items():
            if name.startswith("DC Block"):
                ax.plot([bus_xy[0], xy[0]], [bus_xy[1], xy[1]], color="#9CA3AF", linestyle="--", linewidth=1)

    chain = ["AC Block", "Transformer", "PCS", "AC Busbar"]
    for left, right in zip(chain, chain[1:]):
        if left in positions and right in positions:
            lx, ly = positions[left]
            rx, ry = positions[right]
            ax.plot([lx, rx], [ly, ry], color="#2563EB", linewidth=2)

    xs = [xy[0] for xy in positions.values()]
    ys = [xy[1] for xy in positions.values()]
    margin = 1.5
    ax.set_xlim(min(xs) - margin, max(xs) + margin)
    ax.set_ylim(min(ys) - margin, max(ys) + margin)
    ax.set_aspect("equal", adjustable="datalim")
    ax.axis("off")
    ax.set_title("Physical Layout: 1 AC Block with Connected DC Blocks")
    return fig


def require_stage13_output() -> Dict[str, Any]:
    stage13 = st.session_state.get("stage13_output")
    if not stage13:
        # Check if we have DC sizing results in session state from dc_view.py
        # If user ran DC sizing but stage13_output wasn't formed yet, we can try to form it or warn.
        # For now, consistent with the request, just stop.
        st.warning("Stage 1–3 output not found. Please complete 'DC Block Sizing' first.")
        # Return empty dict to prevent crash, but UI will likely stop logic
        return {}
    return stage13


def render_step1_summary(res: Dict[str, Any]) -> None:
    st.markdown("## Step 1 · AC Block Sizing Summary (V0.4)")
    st.write(f"**Strategy**: {res.get('strategy', '').replace('_', ' ').title()}")

    cA, cB, cC, cD, cE = st.columns(5)
    with cA:
        st.metric("AC Blocks Quantity", f"{res['ac_block_qty']}")
    with cB:
        st.metric("AC Block Rating (MW)", f"{res['ac_block_rated_mw']:.2f}")
    with cC:
        st.metric("PCS per Block", f"{res['pcs_units']} × {res['pcs_unit_kw']} kW")
    with cD:
        st.metric("Total AC Capacity (MW)", f"{res['total_ac_mw']:.2f}")
    with cE:
        st.metric("Oversize vs POI (MW)", f"{res['oversize_mw']:.2f}")

    if res.get("strategy") == "container_only":
        st.write(f"**DC Blocks per AC Block**: {res.get('dc_blocks_per_block', 0)}")
        return

    st.write("**Mixed DC Distribution per AC Block**")
    st.write(f"• Containers / Block: {res.get('container_per_block', 0)}")
    st.write(f"• Cabinets / Block: {res.get('cabinet_per_block', 0)}")
    if res.get("container_rem", 0) or res.get("cabinet_rem", 0):
        st.write(
            f"  ⚠ Remainder not evenly divisible: container_rem={res.get('container_rem', 0)}, "
            f"cabinet_rem={res.get('cabinet_rem', 0)}"
        )
    st.write(
        f"  • DC Blocks per Block (base/max): "
        f"{res.get('dc_blocks_per_block_base', 0)} / {res.get('dc_blocks_per_block_max', 0)}"
    )


def make_ac_report_filename(project_name: str) -> str:
    safe = "".join(c if c.isalnum() or c in ("-", "_") else "_" for c in project_name) or "CALB_ESS_Project"
    return f"{safe}_AC_Block_Report.docx"


def build_ac_report(stage13: Dict[str, Any], res: Dict[str, Any]) -> bytes | None:
    if not DOCX_AVAILABLE:
        return None

    project_name = stage13.get("project_name", "CALB ESS Project")
    doc = Document()
    doc.add_heading("AC Block Sizing Report", level=0)
    doc.add_paragraph(f"Project: {project_name}")

    strategy = res.get("strategy", "").replace("_", " ").title()
    doc.add_paragraph(f"Strategy: {strategy}")

    p_meta = doc.add_paragraph()
    p_meta.add_run("Input Snapshot").bold = True
    doc.add_paragraph(
        f"POI Power Requirement: {stage13.get('poi_power_req_mw', 0.0):.2f} MW\n"
        f"POI Nominal Voltage: {stage13.get('poi_nominal_voltage_kv', '')} kV\n"
        f"DC Containers: {stage13.get('container_count', 0)}\n"
        f"DC Cabinets: {stage13.get('cabinet_count', 0)}"
    )

    doc.add_heading("AC Block Summary", level=1)
    summary_items = {
        "AC Blocks Quantity": res.get("ac_block_qty", 0),
        "AC Block Rating (MW)": f"{res.get('ac_block_rated_mw', 0.0):.2f}",
        "PCS per Block": f"{res.get('pcs_units', 0)} × {res.get('pcs_unit_kw', 0)} kW",
        "Total AC Capacity (MW)": f"{res.get('total_ac_mw', 0.0):.2f}",
        "Oversize vs POI (MW)": f"{res.get('oversize_mw', 0.0):.2f}",
    }

    if res.get("strategy") == "mixed":
        summary_items.update(
            {
                "Containers / Block": res.get("container_per_block", 0),
                "Cabinets / Block": res.get("cabinet_per_block", 0),
                "DC Blocks per Block (base/max)": f"{res.get('dc_blocks_per_block_base', 0)} / {res.get('dc_blocks_per_block_max', 0)}",
            }
        )
    else:
        summary_items["DC Blocks per AC Block"] = res.get("dc_blocks_per_block", 0)

    table = doc.add_table(rows=1, cols=2)
    header_cells = table.rows[0].cells
    header_cells[0].text = "Metric"
    header_cells[1].text = "Value"

    for key, value in summary_items.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def run_step1_sizing(
    poi_mw: float,
    container_cnt: int,
    cabinet_cnt: int,
    search_extra: int,
) -> Optional[Dict[str, Any]]:
    best_container = find_ac_block_container_only(
        poi_mw=float(poi_mw),
        container_cnt=int(container_cnt),
        search_extra=int(search_extra),
    )
    if best_container:
        st.success("Container-only DC Blocks matched AC Block configuration.")
        return best_container

    best_mixed = find_ac_block_mixed(
        poi_mw=float(poi_mw),
        container_cnt=int(container_cnt),
        cabinet_cnt=int(cabinet_cnt),
        search_extra=int(search_extra),
    )
    if best_mixed:
        st.warning("Container-only did not satisfy rules, using mixed DC Blocks for AC Block matching.")
        return best_mixed

    st.error(
        "Neither container-only nor mixed DC Blocks can satisfy AC Block rules. "
        "Please revise DC configuration in Stage 1–3 (e.g., adjust container/cabinet counts)."
    )
    return None

# -----------------------------------------------------
# Main Interface Function (Called by app.py)
# -----------------------------------------------------
def show() -> None:
    st.title("Stage 4 – AC Block (V0.4)")
    st.caption("Chain: RMU/SW → Transformer (MV/LV) → PCS → DC Busbar → DC Block")

    # 1. 尝试读取 AC 数据字典 (用于后续逻辑扩展)
    try:
        df = pd.read_excel(AC_DATA_PATH)
        # st.success("Loaded AC Block Data Dictionary") # Optional: Debug msg
    except Exception as e:
        st.error(f"无法读取数据文件 (AC_DATA_PATH): {e}")
        st.stop()

    # 2. 检查前置步骤数据
    stage13 = require_stage13_output()
    if not stage13:
        st.stop()

    # 3. 页面 Tabs
    tab1, tab2, tab3 = st.tabs(
        [
            "Step 1 · AC Block Sizing",
            "Step 2 · Block SLD + Layout",
            "Step 3 · Site + Simulation",
        ]
    )

    with tab1:
        poi_mw = stage13.get("poi_power_req_mw", 0.0)
        container_cnt = stage13.get("container_count", 0)
        cabinet_cnt = stage13.get("cabinet_count", 0)
        dc_total = stage13.get("dc_total_blocks", container_cnt + cabinet_cnt)
        poi_voltage = stage13.get("poi_nominal_voltage_kv", "")

        c1, c2, c3 = st.columns(3)
        c1.metric("POI Power Requirement (MW)", f"{poi_mw:.2f}")
        c2.metric("DC Blocks Total", f"{dc_total}")
        c3.metric("POI Nominal Voltage (kV)", f"{poi_voltage}")

        st.write(f"Current DC configuration: container = {container_cnt}, cabinet = {cabinet_cnt}")
        st.divider()

        search_extra = st.number_input(
            "Search Extra AC Block Qty (V0.4)",
            min_value=0,
            value=40,
            step=5,
        )

        if st.button("Run Step 1 · AC Block Sizing"):
            result = run_step1_sizing(
                poi_mw=poi_mw,
                container_cnt=container_cnt,
                cabinet_cnt=cabinet_cnt,
                search_extra=search_extra,
            )
            if result:
                st.session_state["stage4_step1_result"] = result

        res = st.session_state.get("stage4_step1_result")
        if res:
            render_step1_summary(res)
            if DOCX_AVAILABLE:
                report_bytes = build_ac_report(stage13, res)
                if report_bytes:
                    st.download_button(
                        "📄 Export AC Block Report (DOCX)",
                        data=report_bytes,
                        file_name=make_ac_report_filename(stage13.get("project_name", "CALB_ESS_Project")),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
            else:
                st.info("python-docx not installed; cannot export AC Block report in this environment.")

    with tab2:
        res = st.session_state.get("stage4_step1_result")
        if not res:
            st.info("Run Step 1 first to populate the Block-level SLD and layout.")
        else:
            dc_block_count = get_dc_block_total(stage13)
            st.subheader("Block-level Single Line Diagram (AC + multiple DC Blocks)")
            dot = build_block_sld(dc_block_count, res)
            st.graphviz_chart(dot, use_container_width=True)

            st.divider()
            st.subheader("Physical Layout (not to scale)")
            columns = st.slider("DC Blocks per row", min_value=1, max_value=6, value=4, step=1)
            spacing = st.slider("Spacing between Blocks (m)", min_value=1.0, max_value=4.0, value=2.4, step=0.2)
            positions = compute_layout_positions(dc_block_count, columns=columns, spacing=spacing)
            fig = render_layout_plot(positions)
            st.pyplot(fig, use_container_width=True)

    with tab3:
        res = st.session_state.get("stage4_step1_result")
        if not res:
            st.info("Run Step 1 first to enable power flow and fault simulations.")
        else:
            st.subheader("Site Power Flow Snapshot")
            poi_mw = stage13.get("poi_power_req_mw", 0.0)
            eff_chain = stage13.get("eff_dc_to_poi_frac", 0.95) or 0.95
            dc_flow = poi_mw * eff_chain
            st.metric("POI Power Requirement (MW)", f"{poi_mw:.2f}")
            st.metric("Estimated DC Block Power (MW)", f"{dc_flow:.2f}")
            st.write(
                "Power flow assumes steady-state operation across Transformer → PCS → AC Busbar → DC Blocks. "
                "Adjust Stage 1–3 efficiency to refine the snapshot."
            )

            st.divider()
            st.subheader("Fault Scenario Simulation (conceptual)")
            scenario = st.selectbox(
                "Select a fault scenario",
                ["normal_operation", "short_circuit", "transformer_failure"],
                format_func=lambda v: v.replace("_", " ").title(),
            )

            def simulate_fault(s: str) -> str:
                if s == "short_circuit":
                    return "Short circuit detected at AC Busbar. Isolate bus section and re-route via remaining feeders."
                if s == "transformer_failure":
                    return "Transformer failure detected. Shed load on affected AC Block and engage redundant transformer if available."
                return "System operating normally. Power routed through AC Block to connected DC Blocks."

            st.write(simulate_fault(scenario))
            st.caption("Detailed time-series simulations can be integrated with dispatch logic in future iterations.")