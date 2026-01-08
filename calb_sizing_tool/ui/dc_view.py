import streamlit as st
import pandas as pd
import numpy as np
import math
import altair as alt
import os
import io
from pathlib import Path

# --- 适配新架构的引用 ---
from calb_sizing_tool.config import DC_DATA_PATH, PROJECT_ROOT
from calb_sizing_tool.ui.stage4_interface import pack_stage13_output
# 引入数据模型用于传递给 AC/SLD
from calb_sizing_tool.models import DCBlockResult
from calb_sizing_tool.state.project_state import bump_run_id_dc, init_project_state
from calb_sizing_tool.state.session_state import init_shared_state, set_run_time

# ==========================================
# 0. SETUP & LIBRARY CHECK
# ==========================================
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Optional: Matplotlib for DOCX chart export
try:
    import matplotlib.pyplot as plt
    MATPLOTLIB_AVAILABLE = True
except Exception:
    MATPLOTLIB_AVAILABLE = False

# Design rule: max 418kWh cabinets per DC busbar
K_MAX_FIXED = 10

# ==========================================
# CALB VI COLORS
# ==========================================
CALB_SKY_BLUE   = "#5cc3e4"
CALB_DEEP_BLUE  = "#23496b"
CALB_BLACK      = "#1e1e1e"
CALB_GREY       = "#58595b"
CALB_LIGHT_GREY = "#bebec3"
CALB_WISE_GREY  = "#cedbea"
CALB_WHITE      = "#ffffff"

# ==========================================
# HELPERS
# ==========================================
def inject_css():
    base_theme = st.get_option("theme.base") or "light"
    is_dark = (base_theme == "dark")

    if is_dark:
        BG_MAIN      = "#0f1116"
        TITLE_COLOR  = CALB_SKY_BLUE
        CARD_BG      = "#1b1e24"
        CARD_BORDER  = "#2c2f36"
        METRIC_LABEL = "#9ca3af"
        METRIC_VALUE = CALB_SKY_BLUE
        BUTTON_BG    = CALB_SKY_BLUE
        BUTTON_FG    = CALB_BLACK
        BUTTON_HOVER_BG = CALB_WHITE
        BUTTON_HOVER_FG = CALB_DEEP_BLUE
    else:
        BG_MAIN      = CALB_WISE_GREY
        TITLE_COLOR  = CALB_DEEP_BLUE
        CARD_BG      = CALB_WHITE
        CARD_BORDER  = CALB_LIGHT_GREY
        METRIC_LABEL = CALB_GREY
        METRIC_VALUE = CALB_DEEP_BLUE
        BUTTON_BG    = CALB_SKY_BLUE
        BUTTON_FG    = CALB_WHITE
        BUTTON_HOVER_BG = CALB_DEEP_BLUE
        BUTTON_HOVER_FG = CALB_WHITE

    global CHART_TEXT_COLOR
    CHART_TEXT_COLOR = CALB_DEEP_BLUE if not is_dark else CALB_SKY_BLUE

    st.markdown(
        f"""
    <style>
    .calb-page-title {{
        font-family: "Segoe UI", "Helvetica Neue", Arial, sans-serif;
        color: {TITLE_COLOR};
        margin-bottom: 0.0rem;
    }}
    .calb-card {{
        background-color: {CARD_BG};
        padding: 1.2rem 1.6rem;
        border-radius: 18px;
        border: 1px solid {CARD_BORDER};
        box-shadow: 0 6px 18px rgba(0, 0, 0, 0.15);
        margin-bottom: 1.4rem;
    }}
    .metric-label {{
        color: {METRIC_LABEL} !important;
        font-size: 0.85rem;
        text-transform: uppercase;
        letter-spacing: 0.04em;
        font-family: "Segoe UI", "Helvetica Neue", Arial, sans-serif;
    }}
    .metric-value {{
        font-size: 1.3rem;
        font-weight: 600;
        color: {METRIC_VALUE} !important;
        font-family: "Segoe UI", "Helvetica Neue", Arial, sans-serif;
    }}
    .stButton>button {{
        background-color: {BUTTON_BG};
        color: {BUTTON_FG};
        border-radius: 999px;
        border: none;
        padding: 0.7rem 2.3rem;
        font-weight: 650;
        font-size: 1.0rem;
        font-family: "Segoe UI", "Helvetica Neue", Arial, sans-serif;
        box-shadow: 0 4px 10px rgba(0,0,0,0.25);
    }}
    .stButton>button:hover {{
        background-color: {BUTTON_HOVER_BG};
        color: {BUTTON_HOVER_FG};
    }}
    </style>
    """,
        unsafe_allow_html=True,
    )

def to_float(x, default=0.0):
    try:
        if isinstance(x, str):
            x = x.replace("%", "").replace(",", "").strip()
        return float(x)
    except Exception:
        return default

def to_frac(x, default=1.0):
    v = to_float(x, default)
    if v > 1.5:
        return v / 100.0
    return v

def safe_div(a: float, b: float, default: float = 0.0) -> float:
    try:
        if b is None or abs(float(b)) < 1e-12:
            return default
        return float(a) / float(b)
    except Exception:
        return default

def calc_sc_loss_pct(sc_months: float) -> float:
    m = int(round(to_float(sc_months, 0.0)))
    if m <= 0:
        return 0.0

    mapping = {
        1: 2.0, 2: 2.0, 3: 2.0,
        4: 2.5, 5: 2.8, 6: 3.0,
        7: 3.2, 8: 3.5, 9: 3.8,
        10: 4.1, 11: 4.3, 12: 4.5,
    }

    if m in mapping:
        return mapping[m]

    if m > 12:
        base_loss = 4.5
        extra_months = sc_months - 12.0
        return base_loss + (extra_months * 0.05)

    lower = int(math.floor(sc_months))
    upper = int(math.ceil(sc_months))
    if lower <= 0:
        return mapping.get(upper, 2.0)
    if upper == lower:
        return mapping.get(lower, 2.0)
    v_low = mapping.get(lower, 2.0)
    v_up = mapping.get(upper, v_low)
    ratio = (sc_months - lower) / (upper - lower)
    return v_low + (v_up - v_low) * ratio

def first_success_key(results: dict, preferred_order: list):
    for k in preferred_order:
        v = results.get(k)
        if isinstance(v, tuple) and len(v) > 0 and v[0] != "ERROR":
            return k
    return None

# ==========================================
# 3. LOAD DATA FROM EXCEL
# ==========================================
@st.cache_data
def load_data(path: Path):
    if not path.is_file():
         raise FileNotFoundError(f"Data file not found at {path}")
         
    xls = pd.ExcelFile(path)

    df_case = pd.read_excel(xls, "ess_sizing_case")
    defaults = {}
    for _, row in df_case.iterrows():
        field_name = row.get("Field Name")
        if isinstance(field_name, str) and field_name.strip():
            defaults[field_name.strip()] = row.get("Default Value")

    df_blocks = pd.read_excel(xls, "dc_block_template_314_data")
    df_soh_profile = pd.read_excel(xls, "soh_profile_314_data")
    df_soh_curve = pd.read_excel(xls, "soh_curve_314_template")
    df_rte_profile = pd.read_excel(xls, "rte_profile_314_data")
    df_rte_curve = pd.read_excel(xls, "rte_curve_314_template")

    return defaults, df_blocks, df_soh_profile, df_soh_curve, df_rte_profile, df_rte_curve

# ==========================================
# 4. CORE CALC LOGIC
# ==========================================
def run_stage1(inputs: dict, defaults: dict) -> dict:
    def get(name, fallback=None):
        if name in inputs and inputs[name] is not None:
            return inputs[name]
        if name in defaults:
            return defaults[name]
        return fallback

    project_name = str(get("project_name", "CALB ESS Project"))

    poi_mw = to_float(get("poi_power_req_mw", 100.0))
    poi_mwh = to_float(get("poi_energy_req_mwh", 400.0))
    project_life_years = int(to_float(get("project_life_years", 20)))
    cycles_per_year = int(to_float(get("cycles_per_year", 365)))
    poi_guarantee_year = int(to_float(get("poi_guarantee_year", 0)))

    eff_dc_cables = to_frac(get("eff_dc_cables", 0.995))
    eff_pcs       = to_frac(get("eff_pcs", 0.985))
    eff_mvt       = to_frac(get("eff_mvt", 0.995))
    eff_ac_sw     = to_frac(get("eff_ac_cables_sw_rmu", 0.992))
    eff_hvt       = to_frac(get("eff_hvt_others", 1.0))
    eff_chain = eff_dc_cables * eff_pcs * eff_mvt * eff_ac_sw * eff_hvt

    sc_val = to_float(get("sc_time_months", 3.0))
    if sc_val < 3.0:
        sc_val = 3.0
    sc_time_months = sc_val
    sc_loss_pct = calc_sc_loss_pct(sc_time_months)
    sc_loss_frac = sc_loss_pct / 100.0

    dod_frac    = to_frac(get("dod_pct", 97.0))
    dc_rte_frac = to_frac(get("dc_round_trip_efficiency_pct", 94.0))
    dc_one_way_eff = math.sqrt(dc_rte_frac) if dc_rte_frac >= 0 else 0.0
    dc_usable_bol_frac = dod_frac * dc_one_way_eff

    denom = (1.0 - sc_loss_frac) * dc_usable_bol_frac * eff_chain
    dc_energy_required = safe_div(poi_mwh, denom, default=0.0)

    dc_power_required_mw = safe_div(poi_mw, eff_chain, default=0.0) if eff_chain > 0 else 0.0

    return {
        "project_name": project_name,
        "poi_power_req_mw": poi_mw,
        "poi_energy_req_mwh": poi_mwh,
        "project_life_years": project_life_years,
        "cycles_per_year": cycles_per_year,
        "poi_guarantee_year": poi_guarantee_year,
        "eff_dc_cables_frac": eff_dc_cables,
        "eff_pcs_frac": eff_pcs,
        "eff_mvt_frac": eff_mvt,
        "eff_ac_cables_sw_rmu_frac": eff_ac_sw,
        "eff_hvt_others_frac": eff_hvt,
        "eff_dc_to_poi_frac": eff_chain,
        "sc_time_months": sc_time_months,
        "sc_loss_pct": sc_loss_pct,
        "sc_loss_frac": sc_loss_frac,
        "dod_frac": dod_frac,
        "dc_round_trip_efficiency_frac": dc_rte_frac,
        "dc_one_way_efficiency_frac": dc_one_way_eff,
        "dc_usable_bol_frac": dc_usable_bol_frac,
        "dc_energy_capacity_required_mwh": dc_energy_required,
        "dc_power_required_mw": dc_power_required_mw,
    }

def _pick_dc_block(df_blocks: pd.DataFrame, form: str):
    df = df_blocks.copy()
    df["Block_Form_L"] = df["Block_Form"].astype(str).str.lower()
    cand = df[(df["Block_Form_L"] == form.lower()) & (df["Is_Active"] == 1)]
    if cand.empty:
        return None

    pref = cand[cand["Is_Default_Option"] == 1]
    if pref.empty:
        pref = cand

    pref = pref.sort_values("Block_Nameplate_Capacity_Mwh", ascending=False)
    row = pref.iloc[0]
    return str(row["Dc_Block_Code"]), str(row["Dc_Block_Name"]), float(row["Block_Nameplate_Capacity_Mwh"])

def _make_config_table(rows: list):
    df = pd.DataFrame(rows)
    if df.empty:
        return df

    df["Subtotal (MWh)"] = df["Unit Capacity (MWh)"] * df["Count"]
    total = float(df["Subtotal (MWh)"].sum())
    df["Total DC Nameplate @BOL (MWh)"] = total
    return df

def build_config_container_only(required_dc_mwh: float, container_unit: float, container_code: str, container_name: str):
    cnt = int(math.ceil(required_dc_mwh / container_unit)) if required_dc_mwh > 0 else 0
    rows = [{
        "Block Code": container_code,
        "Block Name": container_name,
        "Form": "container",
        "Unit Capacity (MWh)": float(container_unit),
        "Count": int(cnt),
    }]
    df = _make_config_table(rows)
    total = float(df["Subtotal (MWh)"].sum()) if not df.empty else 0.0
    oversize = total - required_dc_mwh
    return {
        "mode": "container_only",
        "dc_nameplate_bol_mwh": total,
        "oversize_mwh": oversize,
        "config_adjustment_frac": (total / required_dc_mwh - 1.0) if required_dc_mwh > 0 else 0.0,
        "block_config_table": df,
        "container_count": cnt,
        "cabinet_count": 0,
        "busbars_needed": 0,
    }

def build_config_cabinet_only(required_dc_mwh: float, cab_unit: float, cab_code: str, cab_name: str, k_max: int = K_MAX_FIXED):
    cab = int(math.ceil(required_dc_mwh / cab_unit)) if required_dc_mwh > 0 else 0
    busbars = int(math.ceil(cab / k_max)) if cab > 0 else 0
    rows = [{
        "Block Code": cab_code,
        "Block Name": cab_name,
        "Form": "cabinet",
        "Unit Capacity (MWh)": float(cab_unit),
        "Count": int(cab),
    }]
    df = _make_config_table(rows)
    total = float(df["Subtotal (MWh)"].sum()) if not df.empty else 0.0
    oversize = total - required_dc_mwh
    return {
        "mode": "cabinet_only",
        "dc_nameplate_bol_mwh": total,
        "oversize_mwh": oversize,
        "config_adjustment_frac": (total / required_dc_mwh - 1.0) if required_dc_mwh > 0 else 0.0,
        "block_config_table": df,
        "container_count": 0,
        "cabinet_count": cab,
        "busbars_needed": busbars,
    }

def build_config_hybrid(required_dc_mwh: float,
                        container_unit: float, container_code: str, container_name: str,
                        cab_unit: float, cab_code: str, cab_name: str,
                        k_max: int = K_MAX_FIXED):
    if required_dc_mwh <= 0:
        cont = 0
        cab = 0
    else:
        cont = int(math.floor(required_dc_mwh / container_unit))
        remainder = required_dc_mwh - cont * container_unit
        if remainder <= 1e-9:
            cab = 0
        else:
            cab = int(math.ceil(remainder / cab_unit))
            if cab > k_max:
                cont += 1
                cab = 0

    rows = []
    if cont > 0:
        rows.append({
            "Block Code": container_code,
            "Block Name": container_name,
            "Form": "container",
            "Unit Capacity (MWh)": float(container_unit),
            "Count": int(cont),
        })
    if cab > 0:
        rows.append({
            "Block Code": cab_code,
            "Block Name": cab_name,
            "Form": "cabinet",
            "Unit Capacity (MWh)": float(cab_unit),
            "Count": int(cab),
        })

    df = _make_config_table(rows)
    total = float(df["Subtotal (MWh)"].sum()) if not df.empty else 0.0
    oversize = total - required_dc_mwh
    busbars = 1 if cab > 0 else 0

    return {
        "mode": "hybrid",
        "dc_nameplate_bol_mwh": total,
        "oversize_mwh": oversize,
        "config_adjustment_frac": (total / required_dc_mwh - 1.0) if required_dc_mwh > 0 else 0.0,
        "block_config_table": df,
        "container_count": cont,
        "cabinet_count": cab,
        "busbars_needed": busbars,
    }

def select_soh_profile(effective_c_rate: float, cycles_per_year: int, df_soh_profile: pd.DataFrame):
    df = df_soh_profile.copy()
    df["c_rate_diff"] = (df["C_Rate"] - effective_c_rate).abs()
    df["cycles_diff"] = (df["Cycles_Per_Year"] - cycles_per_year).abs()
    df["score"] = df["c_rate_diff"] * 10.0 + df["cycles_diff"] / 365.0
    best = df.sort_values("score").iloc[0]
    return int(best["Profile_Id"]), float(best["C_Rate"]), int(best["Cycles_Per_Year"])

def select_rte_profile(effective_c_rate: float, df_rte_profile: pd.DataFrame):
    df = df_rte_profile.copy()
    df["c_rate_diff"] = (df["C_Rate"] - effective_c_rate).abs()
    best = df.sort_values("c_rate_diff").iloc[0]
    return int(best["Profile_Id"]), float(best["C_Rate"])

def run_stage3(stage1: dict,
               stage2: dict,
               df_soh_profile: pd.DataFrame,
               df_soh_curve: pd.DataFrame,
               df_rte_profile: pd.DataFrame,
               df_rte_curve: pd.DataFrame):

    dc_nameplate_bol_mwh = stage2["dc_nameplate_bol_mwh"]
    if dc_nameplate_bol_mwh <= 0:
        raise ValueError("DC nameplate @BOL must be > 0 for Stage 3.")

    poi_mw = stage1["poi_power_req_mw"]
    poi_energy_mwh = stage1["poi_energy_req_mwh"]
    project_life_years = int(stage1["project_life_years"])
    cycles_per_year = int(stage1["cycles_per_year"])
    sc_loss_frac = stage1["sc_loss_frac"]
    dod_frac = stage1["dod_frac"]
    eff_chain = stage1["eff_dc_to_poi_frac"]
    guarantee_year = int(stage1.get("poi_guarantee_year", 0))

    dc_power_mw = safe_div(poi_mw, eff_chain, default=0.0) if eff_chain > 0 else 0.0
    effective_c_rate = safe_div(dc_power_mw, dc_nameplate_bol_mwh, default=0.0)

    soh_profile_id, chosen_soh_c_rate, chosen_cycles_per_year = select_soh_profile(
        effective_c_rate, cycles_per_year, df_soh_profile
    )
    rte_profile_id, chosen_rte_c_rate = select_rte_profile(effective_c_rate, df_rte_profile)

    soh_curve_sel = df_soh_curve[df_soh_curve["Profile_Id"] == soh_profile_id].copy()
    if "Soh_Dc_Pct" in soh_curve_sel.columns:
        soh_curve_sel["Soh_Dc_Pct"] = soh_curve_sel["Soh_Dc_Pct"].apply(lambda x: to_frac(x))
    soh_curve_sel = soh_curve_sel.sort_values("Life_Year_Index")

    rte_curve_sel = df_rte_curve[df_rte_curve["Profile_Id"] == rte_profile_id].copy()
    if "Soh_Band_Min_Pct" in rte_curve_sel.columns:
        rte_curve_sel["Soh_Band_Min_Pct"] = rte_curve_sel["Soh_Band_Min_Pct"].apply(lambda x: to_frac(x))
    if "Rte_Dc_Pct" in rte_curve_sel.columns:
        rte_curve_sel["Rte_Dc_Pct"] = rte_curve_sel["Rte_Dc_Pct"].apply(lambda x: to_frac(x))
    rte_curve_sel = rte_curve_sel.sort_values("Soh_Band_Min_Pct", ascending=False)

    records = []

    for y in range(0, project_life_years + 1):
        row = soh_curve_sel[soh_curve_sel["Life_Year_Index"] == y]
        if not row.empty:
            soh_rel = float(row["Soh_Dc_Pct"].iloc[0])
        else:
            soh_rel = float(soh_curve_sel["Soh_Dc_Pct"].iloc[-1])

        soh_abs = soh_rel * (1.0 - sc_loss_frac)

        rte_row = rte_curve_sel[rte_curve_sel["Soh_Band_Min_Pct"] <= soh_abs].head(1)
        if rte_row.empty:
            rte_row = rte_curve_sel.tail(1)

        raw_rte = float(rte_row["Rte_Dc_Pct"].iloc[0])
        dc_rte_frac_year = min(1.0, max(0.0, raw_rte))
        dc_one_way_eff_year = math.sqrt(dc_rte_frac_year)

        dc_gross_capacity_mwh_year = dc_nameplate_bol_mwh * soh_abs
        dc_usable_mwh_year = dc_gross_capacity_mwh_year * dod_frac * dc_one_way_eff_year
        poi_usable_mwh_year = max(dc_usable_mwh_year * eff_chain, 0.0)

        system_rte_frac_year = min(1.0, max(0.0, dc_rte_frac_year * (eff_chain ** 2)))
        meets_poi_energy = poi_usable_mwh_year >= poi_energy_mwh

        records.append(
            {
                "Year_Index": int(y),
                "SOH_Relative": soh_rel,
                "SOH_Absolute": soh_abs,
                "DC_Nameplate_BOL_MWh": dc_nameplate_bol_mwh,
                "DC_Gross_Capacity_MWh": dc_gross_capacity_mwh_year,
                "DC_Usable_MWh": dc_usable_mwh_year,
                "DC_RTE_Frac": dc_rte_frac_year,
                "System_RTE_Frac": system_rte_frac_year,
                "POI_Usable_Energy_MWh": poi_usable_mwh_year,
                "Meets_POI_Req": meets_poi_energy,
                "Is_Guarantee_Year": (y == guarantee_year),
            }
        )

    df_years = pd.DataFrame(records)
    df_years["SOH_Display_Pct"] = df_years["SOH_Relative"] * 100.0
    df_years["SOH_Absolute_Pct"] = df_years["SOH_Absolute"] * 100.0
    df_years["DC_RTE_Pct"] = df_years["DC_RTE_Frac"] * 100.0
    df_years["System_RTE_Pct"] = df_years["System_RTE_Frac"] * 100.0

    meta = {
        "poi_power_mw": poi_mw,
        "dc_power_mw": dc_power_mw,
        "effective_c_rate": effective_c_rate,
        "soh_profile_id": soh_profile_id,
        "rte_profile_id": rte_profile_id,
        "chosen_soh_c_rate": chosen_soh_c_rate,
        "chosen_soh_cycles_per_year": chosen_cycles_per_year,
        "chosen_rte_c_rate": chosen_rte_c_rate,
    }
    return df_years, meta

def size_with_guarantee(stage1: dict,
                        mode: str,
                        df_blocks: pd.DataFrame,
                        df_soh_profile: pd.DataFrame,
                        df_soh_curve: pd.DataFrame,
                        df_rte_profile: pd.DataFrame,
                        df_rte_curve: pd.DataFrame,
                        k_max: int = K_MAX_FIXED,
                        max_iter: int = 60):

    required_dc = float(stage1["dc_energy_capacity_required_mwh"])
    poi_energy_req = float(stage1["poi_energy_req_mwh"])
    guarantee_year = int(stage1.get("poi_guarantee_year", 0))
    project_life_years = int(stage1["project_life_years"])
    guarantee_year = max(0, min(guarantee_year, project_life_years))

    picked_container = _pick_dc_block(df_blocks, "container")
    picked_cabinet   = _pick_dc_block(df_blocks, "cabinet")

    if picked_container is None:
        raise ValueError("No active 'container' DC block template found in Excel.")
    if picked_cabinet is None:
        if mode != "container_only":
            raise ValueError("No active 'cabinet' DC block template found in Excel (needed for hybrid/cabinet-only).")

    cont_code, cont_name, cont_unit = picked_container
    if picked_cabinet:
        cab_code, cab_name, cab_unit = picked_cabinet
    else:
        cab_code, cab_name, cab_unit = "", "", 0.0

    if mode == "container_only":
        s2 = build_config_container_only(required_dc, cont_unit, cont_code, cont_name)
    elif mode == "hybrid":
        s2 = build_config_hybrid(required_dc, cont_unit, cont_code, cont_name, cab_unit, cab_code, cab_name, k_max=k_max)
    elif mode == "cabinet_only":
        s2 = build_config_cabinet_only(required_dc, cab_unit, cab_code, cab_name, k_max=k_max)
    else:
        raise ValueError(f"Unknown mode: {mode}")

    s3_df, s3_meta = run_stage3(stage1, s2, df_soh_profile, df_soh_curve, df_rte_profile, df_rte_curve)

    def poi_at_year(df, year):
        row = df[df["Year_Index"] == int(year)]
        if row.empty:
            return None
        return float(row["POI_Usable_Energy_MWh"].iloc[0])

    poi_g = poi_at_year(s3_df, guarantee_year)
    if poi_g is None:
        return s2, s3_df, s3_meta, 1, None, False

    iter_count = 1
    converged = (poi_g + 1e-6 >= poi_energy_req)

    while not converged and iter_count < max_iter:
        if mode == "container_only":
            s2["container_count"] += 1
        elif mode == "cabinet_only":
            s2["cabinet_count"] += 1
        elif mode == "hybrid":
            if s2["cabinet_count"] < k_max and s2["cabinet_count"] > 0:
                s2["cabinet_count"] += 1
            elif s2["cabinet_count"] == 0:
                s2["cabinet_count"] = 1
            else:
                s2["container_count"] += 1

        rows = []
        if s2["container_count"] > 0:
            rows.append({
                "Block Code": cont_code,
                "Block Name": cont_name,
                "Form": "container",
                "Unit Capacity (MWh)": float(cont_unit),
                "Count": int(s2["container_count"]),
            })
        if mode in ("hybrid", "cabinet_only") and s2["cabinet_count"] > 0:
            rows.append({
                "Block Code": cab_code,
                "Block Name": cab_name,
                "Form": "cabinet",
                "Unit Capacity (MWh)": float(cab_unit),
                "Count": int(s2["cabinet_count"]),
            })

        df_cfg = _make_config_table(rows)
        s2["block_config_table"] = df_cfg
        total = float(df_cfg["Subtotal (MWh)"].sum()) if not df_cfg.empty else 0.0
        s2["dc_nameplate_bol_mwh"] = total
        s2["oversize_mwh"] = total - required_dc
        s2["config_adjustment_frac"] = (total / required_dc - 1.0) if required_dc > 0 else 0.0

        if mode == "cabinet_only":
            s2["busbars_needed"] = int(math.ceil(s2["cabinet_count"] / k_max)) if s2["cabinet_count"] > 0 else 0
        elif mode == "hybrid":
            s2["busbars_needed"] = 1 if s2["cabinet_count"] > 0 else 0
        else:
            s2["busbars_needed"] = 0

        s3_df, s3_meta = run_stage3(stage1, s2, df_soh_profile, df_soh_curve, df_rte_profile, df_rte_curve)
        poi_g = poi_at_year(s3_df, guarantee_year)

        iter_count += 1
        if poi_g is not None and (poi_g + 1e-6 >= poi_energy_req):
            converged = True

    return s2, s3_df, s3_meta, iter_count, poi_g, converged

# ==========================================
# 5. REPORT EXPORT HELPERS
# ==========================================
def find_logo_for_report():
    try:
        candidates = [
            PROJECT_ROOT / "calb_assets" / "logo" / "calb_logo.png",
            PROJECT_ROOT / "calb_logo.png",
        ]
        for candidate in candidates:
            if candidate.exists():
                return str(candidate)
        data_dir = os.path.dirname(os.path.abspath(DC_DATA_PATH))
        for fname in os.listdir(data_dir):
            lower = fname.lower()
            if lower.endswith((".png", ".jpg", ".jpeg")) and ("logo" in lower or "calb" in lower):
                return os.path.join(data_dir, fname)
    except Exception:
        return None
    return None

def make_report_filename(project_name: str) -> str:
    base = (project_name or "CALB_ESS_Project").strip()
    if not base:
        base = "CALB_ESS_Project"
    safe = "".join(c if c.isalnum() or c in (" ", "_", "-") else "_" for c in base)
    safe = "_".join(safe.split())
    return f"{safe}_ESS_Sizing_Report.docx"

def _docx_add_config_table(doc: Document, df_conf: pd.DataFrame):
    if df_conf is None or df_conf.empty:
        doc.add_paragraph("No DC block configuration selected.", style="Intense Quote")
        return

    drop_cols = [
        "Block Code",
        "Form",
        "Config Adjustment (%)",
        "Busbars Needed (K=10)",
        "Oversize (MWh)",
        "Oversize_mwh",
        "Busbars Needed",
    ]
    df_show = df_conf.copy()
    cols_to_drop = [c for c in drop_cols if c in df_show.columns]
    if cols_to_drop:
        df_show = df_show.drop(columns=cols_to_drop)

    TOTAL_COL = "Total DC Nameplate @BOL (MWh)"
    has_total_col = TOTAL_COL in df_show.columns

    tbl = doc.add_table(rows=1, cols=len(df_show.columns))
    tbl.style = "Table Grid"

    hdr = tbl.rows[0].cells
    for j, col in enumerate(df_show.columns):
        hdr[j].text = str(col)
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for i, (_, row) in enumerate(df_show.iterrows()):
        rc = tbl.add_row().cells
        for j, col in enumerate(df_show.columns):
            if has_total_col and col == TOTAL_COL and i > 0:
                rc[j].text = ""
                continue

            val = row[col]
            if isinstance(val, float):
                rc[j].text = f"{val:.3f}"
            else:
                rc[j].text = f"{val}"

def _plot_poi_usable_png(s3_df: pd.DataFrame, poi_target: float, title: str) -> io.BytesIO:
    df = s3_df.sort_values("Year_Index").copy()
    x = df["Year_Index"].astype(int).tolist()
    y = df["POI_Usable_Energy_MWh"].astype(float).tolist()

    fig = plt.figure(figsize=(7.0, 3.2))
    ax = fig.add_subplot(111)
    ax.bar(x, y, color=CALB_SKY_BLUE)
    ax.axhline(poi_target, linewidth=2, color="#ff0000")
    ax.set_title(title)
    ax.set_xlabel("Year (from COD)")
    ax.set_ylabel("POI Usable Energy (MWh)")
    ax.set_xticks(x)
    ax.grid(True, axis="y", linestyle="--", alpha=0.35)

    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf


def _plot_dc_capacity_bar_png(
    s2: dict, s3_df: pd.DataFrame, guarantee_year: int, title: str
) -> io.BytesIO | None:
    if not MATPLOTLIB_AVAILABLE:
        return None
    try:
        bol = None
        if isinstance(s2, dict):
            bol = s2.get("dc_nameplate_bol_mwh")
        cod = None
        yx = None
        if s3_df is not None and not s3_df.empty:
            year0 = s3_df[s3_df["Year_Index"] == 0]
            if not year0.empty:
                cod = float(year0["POI_Usable_Energy_MWh"].iloc[0])
            g_row = s3_df[s3_df["Year_Index"] == int(guarantee_year)]
            if not g_row.empty:
                yx = float(g_row["POI_Usable_Energy_MWh"].iloc[0])

        labels = ["BOL", "COD", f"Y{int(guarantee_year)}"]
        values = [
            float(bol) if bol is not None else 0.0,
            float(cod) if cod is not None else 0.0,
            float(yx) if yx is not None else 0.0,
        ]

        fig = plt.figure(figsize=(6.6, 3.0))
        ax = fig.add_subplot(111)
        ax.bar(labels, values, color=CALB_SKY_BLUE)
        ax.set_title(title)
        ax.set_xlabel("Stage")
        ax.set_ylabel("Energy (MWh)")
        ax.grid(True, axis="y", linestyle="--", alpha=0.35)

        buf = io.BytesIO()
        fig.tight_layout()
        fig.savefig(buf, format="png", dpi=150)
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception:
        return None

def _docx_add_lifetime_table(doc: Document, s3_df: pd.DataFrame):
    cols = [
        "Year_Index",
        "SOH_Display_Pct",
        "SOH_Absolute_Pct",
        "DC_Usable_MWh",
        "POI_Usable_Energy_MWh",
        "DC_RTE_Pct",
        "System_RTE_Pct",
    ]
    for c in cols:
        if c not in s3_df.columns:
            s3_df[c] = np.nan

    headers_map = {
        "Year_Index": "Year (From COD)",
        "SOH_Display_Pct": "SOH @ COD Baseline (%)",
        "SOH_Absolute_Pct": "SOH Vs FAT (%)",
        "DC_Usable_MWh": "DC Usable (MWh)",
        "POI_Usable_Energy_MWh": "POI Usable (MWh)",
        "DC_RTE_Pct": "DC RTE (%)",
        "System_RTE_Pct": "System RTE (%)",
    }

    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for j, c in enumerate(cols):
        hdr[j].text = headers_map[c]
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    df_sorted = s3_df.sort_values("Year_Index")
    for _, r in df_sorted.iterrows():
        rc = tbl.add_row().cells
        rc[0].text = str(int(r["Year_Index"]))
        rc[1].text = f"{r['SOH_Display_Pct']:.1f}" if not pd.isna(r["SOH_Display_Pct"]) else ""
        rc[2].text = f"{r['SOH_Absolute_Pct']:.1f}" if not pd.isna(r["SOH_Absolute_Pct"]) else ""
        rc[3].text = f"{r['DC_Usable_MWh']:.2f}" if not pd.isna(r["DC_Usable_MWh"]) else ""
        rc[4].text = f"{r['POI_Usable_Energy_MWh']:.2f}" if not pd.isna(r["POI_Usable_Energy_MWh"]) else ""
        rc[5].text = f"{r['DC_RTE_Pct']:.1f}%" if not pd.isna(r["DC_RTE_Pct"]) else ""
        rc[6].text = f"{r['System_RTE_Pct']:.1f}%" if not pd.isna(r["System_RTE_Pct"]) else ""

def build_report_bytes(stage1: dict, results_dict: dict, report_order: list):
    if not DOCX_AVAILABLE:
        return None

    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    logo_path = find_logo_for_report()
    header = section.header
    header.is_linked_to_previous = False
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.9))
    hdr_cells = header_table.rows[0].cells

    if logo_path:
        p_logo = hdr_cells[0].paragraphs[0]
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(1.2))

    p_info = hdr_cells[1].paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_info = p_info.add_run(
        "CALB Group Co., Ltd.\n"
        "Utility-Scale Energy Storage Systems\n"
        "Confidential Sizing Report"
    )
    run_info.font.size = Pt(9)
    run_info.font.bold = True

    project_name = stage1.get("project_name", "CALB ESS Project")

    doc.add_heading("CALB Utility-Scale ESS Sizing Report", level=1)
    doc.add_paragraph(f"Project: {project_name}")
    doc.add_paragraph(
        f"POI Requirement: {stage1['poi_power_req_mw']:.2f} MW / "
        f"{stage1['poi_energy_req_mwh']:.2f} MWh"
    )
    doc.add_paragraph("")

    doc.add_heading("1. Project Summary", level=2)
    p = doc.add_paragraph()
    p.add_run(f"Project life: {int(stage1['project_life_years'])} years\n")
    p.add_run(f"POI guarantee year: {int(stage1.get('poi_guarantee_year', 0))}\n")
    p.add_run(f"Cycles per year (assumed): {int(stage1['cycles_per_year'])}\n")
    p.add_run(f"S&C time from FAT to COD: {int(round(stage1.get('sc_time_months', 0)))} months\n")
    p.add_run(f"DC→POI efficiency chain (one-way): {stage1.get('eff_dc_to_poi_frac', 0.0)*100:.2f}%\n")
    p.add_run(f"POI→DC equivalent power: {stage1.get('dc_power_required_mw', 0.0):.2f} MW")

    doc.add_paragraph(
        "This sizing report is based on the 314 Ah cell database and the internal "
        "CALB SOH/RTE profiles for the selected operating conditions."
    )

    doc.add_heading("2. Equipment Summary (DC Blocks)", level=2)

    for key, title in report_order:
        if key not in results_dict:
            continue
        s2, _, _, iter_count, poi_g, converged = results_dict[key]
        doc.add_paragraph(title, style=None)
        _docx_add_config_table(doc, s2.get("block_config_table"))
        doc.add_paragraph(f"Iterations: {iter_count} | Guarantee met: {bool(converged)}")
        if poi_g is not None:
            doc.add_paragraph(f"POI usable energy @ guarantee year: {poi_g:.2f} MWh")
        doc.add_paragraph("")

    doc.add_heading("3. Lifetime POI Usable Energy & SOH (Per Configuration)", level=2)

    poi_target = float(stage1["poi_energy_req_mwh"])
    guarantee_year = int(stage1.get("poi_guarantee_year", 0))

    for key, title in report_order:
        if key not in results_dict:
            continue
        _, s3_df, s3_meta, _, _, _ = results_dict[key]

        doc.add_paragraph(title, style=None)
        doc.add_paragraph(
            f"POI Power = {s3_meta.get('poi_power_mw', 0.0):.2f} MW | "
            f"DC-equivalent Power = {s3_meta.get('dc_power_mw', 0.0):.2f} MW | "
            f"Effective C-rate (DC-side) = {s3_meta.get('effective_c_rate', 0.0):.3f} C"
        )
        doc.add_paragraph(
            f"SOH profile ID = {s3_meta.get('soh_profile_id')} "
            f"(C-rate ≈ {s3_meta.get('chosen_soh_c_rate')}, cycles/year = {s3_meta.get('chosen_soh_cycles_per_year')}); "
            f"RTE profile ID = {s3_meta.get('rte_profile_id')} (C-rate ≈ {s3_meta.get('chosen_rte_c_rate')})."
        )
        doc.add_paragraph(f"Guarantee Year (from COD) = {guarantee_year} | POI Energy Target = {poi_target:.2f} MWh")

        if MATPLOTLIB_AVAILABLE:
            try:
                png = _plot_poi_usable_png(
                    s3_df=s3_df,
                    poi_target=poi_target,
                    title=f"POI Usable Energy vs Year – {key}"
                )
                doc.add_picture(png, width=Inches(6.7))
            except Exception:
                doc.add_paragraph("Chart export skipped due to plotting error.")
        else:
            doc.add_paragraph("Chart export skipped (matplotlib not available).")

        _docx_add_lifetime_table(doc, s3_df)
        doc.add_paragraph("")

    p_final = doc.add_paragraph()
    p_fmt = p_final.paragraph_format
    p_fmt.space_before = Pt(0)
    p_fmt.space_after = Pt(0)
    p_fmt.line_spacing = Pt(0)
    run_final = p_final.add_run()
    run_final.font.size = Pt(1)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ==========================================
# 6. MAIN VIEW FUNCTION
# ==========================================
def show():
    state = init_shared_state()
    init_project_state()
    dc_inputs = state.dc_inputs
    dc_results = state.dc_results

    # Inject CSS
    inject_css()
    
    # Title
    st.markdown(
        """
        <div style="padding-top:0.6rem; padding-bottom:0.6rem;">
            <h1 class="calb-page-title">Utility-Scale ESS Sizing Tool V1.0 (DC)</h1>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.markdown("<br/>", unsafe_allow_html=True)
    
    # Load Data
    try:
        defaults, df_blocks, df_soh_profile, df_soh_curve, df_rte_profile, df_rte_curve = load_data(DC_DATA_PATH)
    except Exception as exc:
        st.error(f"❌ Failed to load data file: {exc}")
        return

    def get_default_numeric(key, fallback):
        return to_float(defaults.get(key, fallback), fallback)

    def get_default_str(key, fallback):
        return str(defaults.get(key, fallback))
        
    def get_default_percent_val(key, fallback):
         raw = to_float(defaults.get(key, fallback), fallback)
         if raw <= 1.5 and raw > 0: return raw * 100.0
         return raw

    def _init_input(field: str, default_value):
        key = f"dc_inputs.{field}"
        if key not in st.session_state:
            st.session_state[key] = default_value
        if field not in dc_inputs:
            dc_inputs[field] = st.session_state[key]
        return key

    # --- UI Form ---
    with st.container():
        st.markdown("<div class='calb-card'>", unsafe_allow_html=True)
        st.subheader("1 · Project Inputs")

        with st.form("main_form"):
            project_name_default = (
                dc_inputs.get("project_name")
                or st.session_state.get("project_name")
                or get_default_str("project_name", "CALB ESS Project")
            )
            project_name = st.text_input(
                "Project Name",
                key=_init_input("project_name", project_name_default),
            )
            dc_inputs["project_name"] = project_name
            st.session_state["project_name"] = project_name

            c1, c2, c3 = st.columns(3)
            poi_power = c1.number_input(
                "POI Required Power (MW)",
                key=_init_input("poi_power_req_mw", get_default_numeric("poi_power_req_mw", 100.0)),
            )
            dc_inputs["poi_power_req_mw"] = poi_power
            poi_energy = c2.number_input(
                "POI Required Capacity (MWh)",
                key=_init_input("poi_energy_req_mwh", get_default_numeric("poi_energy_req_mwh", 400.0)),
            )
            dc_inputs["poi_energy_req_mwh"] = poi_energy
            project_life = int(
                c3.number_input(
                    "Project Life (Years)",
                    key=_init_input(
                        "project_life_years",
                        int(get_default_numeric("project_life_years", 20)),
                    ),
                )
            )
            dc_inputs["project_life_years"] = project_life

            mv_default = dc_inputs.get("poi_nominal_voltage_kv")
            if mv_default is None:
                mv_default = st.session_state.get("poi_nominal_voltage_kv", 33.0)
            poi_nominal_voltage_kv = st.number_input(
                "POI / MV Voltage (kV)",
                key=_init_input("poi_nominal_voltage_kv", float(mv_default)),
                step=0.1,
            )
            dc_inputs["poi_nominal_voltage_kv"] = poi_nominal_voltage_kv

            freq_options = ["TBD", 50.0, 60.0]
            freq_default = dc_inputs.get("poi_frequency_option")
            if freq_default is None:
                freq_default = "TBD"
            if freq_default in (50, 50.0):
                freq_index = 1
            elif freq_default in (60, 60.0):
                freq_index = 2
            else:
                freq_index = 0
            poi_frequency = st.selectbox(
                "POI Frequency (Hz)",
                freq_options,
                index=freq_index,
                key=_init_input("poi_frequency_option", freq_default),
                help="Optional; used for reporting only.",
            )
            dc_inputs["poi_frequency_option"] = poi_frequency

            c4, c5, c6 = st.columns(3)
            cycles_year = int(
                c4.number_input(
                    "Cycles Per Year",
                    key=_init_input(
                        "cycles_per_year",
                        int(get_default_numeric("cycles_per_year", 365)),
                    ),
                )
            )
            dc_inputs["cycles_per_year"] = cycles_year
            guarantee_year = int(
                c5.number_input(
                    "POI Guarantee Year",
                    key=_init_input("poi_guarantee_year", 0),
                )
            )
            dc_inputs["poi_guarantee_year"] = guarantee_year
            sc_time_months = int(
                c6.number_input(
                    "S&C Time (Months)",
                    key=_init_input("sc_time_months", 3),
                )
            )
            dc_inputs["sc_time_months"] = sc_time_months

            st.markdown("---")
            st.subheader("2 · DC Parameters")
            
            c7, c8 = st.columns(2)
            dod_pct = c7.number_input(
                "DOD (%)",
                key=_init_input("dod_pct", 97.0),
            )
            dc_inputs["dod_pct"] = dod_pct
            dc_rte_pct = c8.number_input(
                "DC RTE (%)",
                key=_init_input("dc_round_trip_efficiency_pct", 94.0),
            )
            dc_inputs["dc_round_trip_efficiency_pct"] = dc_rte_pct
            
            st.info(f"Design Rule: Max 418kWh Cabinets per DC Busbar (K) = {K_MAX_FIXED} (fixed)")

            st.markdown("**3 · Configuration Options**")
            copt1, copt2, copt3 = st.columns([2, 2, 3])
            enable_hybrid = copt1.checkbox(
                "Enable Hybrid Mode",
                key=_init_input("enable_hybrid", True),
            )
            dc_inputs["enable_hybrid"] = enable_hybrid
            enable_cabinet_only = copt2.checkbox(
                "Enable Cabinet-Only Mode",
                key=_init_input("enable_cabinet_only", True),
            )
            dc_inputs["enable_cabinet_only"] = enable_cabinet_only
            hybrid_disable_threshold = copt3.number_input(
                "Disable Hybrid Threshold (MWh)",
                key=_init_input("hybrid_disable_threshold_mwh", 9999.0),
            )
            dc_inputs["hybrid_disable_threshold_mwh"] = hybrid_disable_threshold

            with st.expander("Advanced: Efficiency Chain"):
                poi_is_dc_side = st.checkbox(
                    "POI Is Located At DC Side (Force 100%)",
                    key=_init_input("poi_is_dc_side", False),
                )
                dc_inputs["poi_is_dc_side"] = poi_is_dc_side
                eff_dc_cables = st.number_input(
                    "DC Cables (%)",
                    key=_init_input("eff_dc_cables", 99.5),
                )
                dc_inputs["eff_dc_cables"] = eff_dc_cables
                eff_pcs = st.number_input(
                    "PCS (%)",
                    key=_init_input("eff_pcs", 98.5),
                )
                dc_inputs["eff_pcs"] = eff_pcs
                eff_mvt = st.number_input(
                    "MV Transformer (%)",
                    key=_init_input("eff_mvt", 99.5),
                )
                dc_inputs["eff_mvt"] = eff_mvt
                eff_ac_sw = st.number_input(
                    "AC Cables + SW (%)",
                    key=_init_input("eff_ac_cables_sw_rmu", 99.2),
                )
                dc_inputs["eff_ac_cables_sw_rmu"] = eff_ac_sw
                eff_hvt = st.number_input(
                    "HVT (%)",
                    key=_init_input("eff_hvt_others", 100.0),
                )
                dc_inputs["eff_hvt_others"] = eff_hvt

            st.markdown("---")
            run_btn = st.form_submit_button("🔄 Run Sizing")
        
        st.markdown("</div>", unsafe_allow_html=True)

    # --- Logic Execution ---
    if run_btn:
        bump_run_id_dc()
        st.session_state["poi_nominal_voltage_kv"] = float(poi_nominal_voltage_kv)
        poi_frequency_hz = None if poi_frequency == "TBD" else float(poi_frequency)
        st.session_state["poi_frequency_hz"] = poi_frequency_hz
        dc_inputs["poi_frequency_hz"] = poi_frequency_hz
        dc_inputs["poi_nominal_voltage_kv"] = float(poi_nominal_voltage_kv)
        dc_inputs["project_name"] = project_name
        dc_inputs["poi_power_req_mw"] = poi_power
        dc_inputs["poi_energy_req_mwh"] = poi_energy
        
        if poi_is_dc_side:
            eff_dc_cables = eff_pcs = eff_mvt = eff_ac_sw = eff_hvt = 100.0

        sc_loss_pct = calc_sc_loss_pct(sc_time_months)
        inputs = {
            "project_name": project_name,
            "poi_power_req_mw": poi_power,
            "poi_energy_req_mwh": poi_energy,
            "eff_dc_cables": eff_dc_cables,
            "eff_pcs": eff_pcs,
            "eff_mvt": eff_mvt,
            "eff_ac_cables_sw_rmu": eff_ac_sw,
            "eff_hvt_others": eff_hvt,
            "sc_time_months": sc_time_months,
            "sc_loss_frac": sc_loss_pct / 100.0,
            "dod_pct": dod_pct,
            "dc_round_trip_efficiency_pct": dc_rte_pct,
            "project_life_years": project_life,
            "cycles_per_year": cycles_year,
            "poi_guarantee_year": guarantee_year
        }

        s1 = run_stage1(inputs, defaults)

        # Stage 1 Display
        st.markdown("<div class='calb-card'>", unsafe_allow_html=True)
        st.subheader("3 · Stage 1 – DC Requirement")
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Theoretical DC Required", f"{s1['dc_energy_capacity_required_mwh']:.2f} MWh")
        m2.metric("Efficiency Chain", f"{s1['eff_dc_to_poi_frac']*100:.2f} %")
        m3.metric("DC Usable @BOL", f"{s1['dc_usable_bol_frac']*100:.2f} %")
        m4.metric("DC Power Req", f"{s1['dc_power_required_mw']:.2f} MW")
        st.markdown("</div>", unsafe_allow_html=True)

        # Run Options
        modes_to_run = ["container_only"]
        if enable_cabinet_only: modes_to_run.insert(0, "cabinet_only")
        if enable_hybrid:
            if hybrid_disable_threshold > 0 and poi_energy >= hybrid_disable_threshold: pass
            else: modes_to_run.insert(0, "hybrid")

        results = {}
        for mode in modes_to_run:
            try:
                results[mode] = size_with_guarantee(
                    s1, mode,
                    df_blocks,
                    df_soh_profile, df_soh_curve,
                    df_rte_profile, df_rte_curve,
                    k_max=K_MAX_FIXED
                )
            except Exception as e:
                results[mode] = ("ERROR", str(e))

        ok_results = {
            k: v for k, v in results.items() if isinstance(v, tuple) and v[0] != "ERROR"
        }
        report_order = [
            (k, k.replace("_", " ").title()) for k in modes_to_run if k in ok_results
        ]
        dc_results["results_dict"] = ok_results
        dc_results["report_order"] = report_order

        # Tabs Display
        st.markdown("<div class='calb-card'>", unsafe_allow_html=True)
        st.subheader("4 · Configuration Comparison")
        
        tabs = st.tabs([m.replace("_", " ").title() for m in modes_to_run])
        for i, mode in enumerate(modes_to_run):
            with tabs[i]:
                res = results.get(mode)
                if isinstance(res, tuple) and res[0] != "ERROR":
                    s2, s3_df, s3_meta, iter_c, poi_g, conv = res
                    
                    c1, c2, c3 = st.columns(3)
                    c1.metric("Installed DC Nameplate", f"{s2['dc_nameplate_bol_mwh']:.3f} MWh")
                    c2.metric("Container Count", int(s2.get('container_count', 0)))
                    c3.metric("Cabinet Count", int(s2.get('cabinet_count', 0)))

                    st.dataframe(s2.get("block_config_table"), use_container_width=True)
                    
                    # Chart
                    s3_df_sorted = s3_df.sort_values("Year_Index")
                    bars = alt.Chart(s3_df_sorted).mark_bar(color=CALB_SKY_BLUE).encode(
                        x=alt.X("Year_Index:O", title="Year"),
                        y=alt.Y("POI_Usable_Energy_MWh:Q", title="POI Usable (MWh)"),
                        tooltip=["Year_Index", "POI_Usable_Energy_MWh", "SOH_Display_Pct"]
                    )
                    rule = alt.Chart(pd.DataFrame({"y": [float(poi_energy)]})).mark_rule(color="red").encode(y="y")
                    st.altair_chart((bars + rule).properties(height=350), use_container_width=True)

                    # --- RESTORED YEARLY DATA TABLE ---
                    with st.expander("Show yearly data table", expanded=False):
                        display_cols = [
                            "Year_Index",
                            "DC_Usable_MWh",
                            "POI_Usable_Energy_MWh",
                            "DC_RTE_Pct",
                            "System_RTE_Pct",
                            "Meets_POI_Req",
                            "Is_Guarantee_Year",
                        ]
                        # Ensure cols exist
                        for c in display_cols:
                            if c not in s3_df_sorted.columns:
                                s3_df_sorted[c] = np.nan

                        disp = s3_df_sorted[display_cols].copy()
                        # Formatting for display
                        disp = disp.rename(columns={
                            "Year_Index": "Year",
                            "DC_Usable_MWh": "DC Usable (MWh)",
                            "POI_Usable_Energy_MWh": "POI Usable (MWh)",
                            "DC_RTE_Pct": "DC RTE (%)",
                            "System_RTE_Pct": "System RTE (%)",
                            "Meets_POI_Req": "Meets Req?",
                            "Is_Guarantee_Year": "Guarantee Year"
                        })
                        try:
                            st.dataframe(disp, use_container_width=True)
                        except TypeError:
                            st.dataframe(disp)
                    # ----------------------------------

                    # Pack data for Session State (For AC/SLD)
                    # We only pack the FIRST valid result as the 'active' one for now, or the user preferred one
                    if mode == "container_only": 
                        # Construct DCBlockResult Pydantic Object
                        container_unit = 5.015 # Default if unavailable
                        if not s2['block_config_table'].empty:
                            # Try to extract actual unit capacity from the first row
                            try:
                                container_unit = float(s2['block_config_table'].iloc[0]['Unit Capacity (MWh)'])
                            except:
                                pass

                        dc_res = DCBlockResult(
                            block_id="DC-Block",
                            container_model="CALB-314Ah", # Ideal: read from config
                            capacity_mwh=container_unit,
                            voltage_v=1200.0,
                            count=int(s2.get('container_count', 0))
                        )

                        st.session_state['dc_result_summary'] = {
                            "mwh": s2['dc_nameplate_bol_mwh'],
                            "target_mw": poi_power,
                            "voltage": 1200, 
                            "container_count": int(s2.get('container_count', 0)),
                            "dc_block": dc_res  # Pass the object for AC view
                        }
                        
                        # Also pack for legacy interface if needed
                        st.session_state["stage13_output"] = pack_stage13_output(
                            s1, s2, s3_meta, 
                            dc_block_total_qty=int(s2.get('container_count', 0)),
                            selected_scenario=mode,
                            poi_nominal_voltage_kv=poi_nominal_voltage_kv,
                            poi_frequency_hz=poi_frequency_hz,
                            stage3_df=s3_df,
                        )
                        dc_results.update(
                            {
                                "stage13_output": st.session_state.get("stage13_output"),
                                "dc_result_summary": st.session_state.get("dc_result_summary"),
                            }
                        )
                        set_run_time("dc_results")

                else:
                    st.error(f"Error: {res[1]}")
        st.markdown("</div>", unsafe_allow_html=True)

        # Export Button
        if DOCX_AVAILABLE:
            ok_results = {k: v for k, v in results.items() if v[0] != "ERROR"}
            report_order = [(k, k.replace("_", " ").title()) for k in modes_to_run if k in ok_results]
            report_bytes = build_report_bytes(s1, ok_results, report_order)
            
            if report_bytes:
                st.download_button(
                    "📄 Export Technical Sizing Report",
                    data=report_bytes,
                    file_name=make_report_filename(project_name),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        st.info("✅ Sizing Complete. Please proceed to AC Sizing or SLD generation via sidebar.")
