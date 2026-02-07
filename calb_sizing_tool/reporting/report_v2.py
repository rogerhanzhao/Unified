# -----------------------------------------------------------------------------
# Personal Open-Source Notice
#
# Copyright (c) 2026 Alex.Zhao. All rights reserved.
#
# This repository is released under the MIT License (see LICENSE file).
# Intended use: learning, evaluation, and engineering reference for Utility-scale
# BESS/ESS sizing and Reporting workflows.
#
# DISCLAIMER: This software is provided "AS IS", without warranty of any kind,
# express or implied. In no event shall the author(s) be liable for any claim,
# damages, or other liability arising from, out of, or in connection with the
# software or the use or other dealings in the software.
#
# NOTE: This is a personal project. It is not an official product or statement
# of any company or organization.
# -----------------------------------------------------------------------------

import datetime
import io
import re
from typing import Optional

import pandas as pd
from docx import Document
from docx.shared import Inches

from calb_sizing_tool.reporting.export_docx import _add_cover_page, _add_table, _doc_to_bytes, _setup_header, _setup_margins
from calb_sizing_tool.reporting.formatter import format_percent, format_value
from calb_sizing_tool.reporting.report_context import ReportContext

try:
    import matplotlib.pyplot as plt

    MATPLOTLIB_AVAILABLE = True
except Exception:
    MATPLOTLIB_AVAILABLE = False

try:
    import cairosvg

    CAIROSVG_AVAILABLE = True
except Exception:
    CAIROSVG_AVAILABLE = False


def _format_percent_with_fraction(value, input_is_fraction=None, fraction_decimals=4) -> str:
    if value is None:
        return ""
    try:
        numeric = float(value)
    except Exception:
        return str(value)

    if input_is_fraction is None:
        is_fraction = numeric <= 1.2
    else:
        is_fraction = bool(input_is_fraction)

    fraction = numeric if is_fraction else numeric / 100.0
    percent_text = format_percent(numeric, input_is_fraction=is_fraction)
    return f"{percent_text} ({fraction:.{fraction_decimals}f})"


def _default_formatter(value):
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass
    return str(value)


def _add_dataframe_table(doc: Document, df: Optional[pd.DataFrame], columns, headers_map, formatters):
    if df is None or df.empty:
        doc.add_paragraph("No data available.")
        return

    rows = []
    for _, row in df.iterrows():
        rows.append([formatters.get(col, _default_formatter)(row.get(col)) for col in columns])

    headers = [headers_map.get(col, col) for col in columns]
    _add_table(doc, rows, headers)


def _plot_poi_usable_png(df: pd.DataFrame, poi_target: float, title: str) -> Optional[io.BytesIO]:
    if not MATPLOTLIB_AVAILABLE:
        return None
    try:
        data = df.sort_values("Year_Index").copy()
        x = data["Year_Index"].astype(int).tolist()
        y = data["POI_Usable_Energy_MWh"].astype(float).tolist()

        fig = plt.figure(figsize=(7.0, 3.2))
        ax = fig.add_subplot(111)
        ax.bar(x, y, color="#5cc3e4")
        ax.axhline(poi_target, linewidth=2, color="#ff0000")
        ax.set_title(title)
        ax.set_xlabel("Year (from COD)")
        ax.set_ylabel("POI Usable Energy (MWh)")
        ax.set_xticks(x)
        ax.grid(True, axis="y", linestyle="--", alpha=0.35)

        buf = io.BytesIO()
        fig.tight_layout()
        fig.savefig(buf, format="png", dpi=150)
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception:
        return None


def _plot_dc_capacity_bar_png(
    bol_mwh: Optional[float],
    s3_df: Optional[pd.DataFrame],
    guarantee_year: int,
    title: str,
) -> Optional[io.BytesIO]:
    if not MATPLOTLIB_AVAILABLE:
        return None
    try:
        cod = None
        yx = None
        if s3_df is not None and not s3_df.empty:
            year0 = s3_df[s3_df["Year_Index"] == 0]
            if not year0.empty:
                cod = float(year0["POI_Usable_Energy_MWh"].iloc[0])
            g_row = s3_df[s3_df["Year_Index"] == int(guarantee_year)]
            if not g_row.empty:
                yx = float(g_row["POI_Usable_Energy_MWh"].iloc[0])

        labels = ["BOL", "COD", f"Y{int(guarantee_year)}"]
        values = [
            float(bol_mwh) if bol_mwh is not None else 0.0,
            float(cod) if cod is not None else 0.0,
            float(yx) if yx is not None else 0.0,
        ]

        fig = plt.figure(figsize=(6.6, 3.0))
        ax = fig.add_subplot(111)
        ax.bar(labels, values, color="#5cc3e4")
        ax.set_title(title)
        ax.set_xlabel("Stage")
        ax.set_ylabel("Energy (MWh)")
        ax.grid(True, axis="y", linestyle="--", alpha=0.35)

        buf = io.BytesIO()
        fig.tight_layout()
        fig.savefig(buf, format="png", dpi=150)
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception:
        return None


def _format_or_tbd(value, unit: str) -> str:
    if value is None:
        return "TBD"
    try:
        numeric = float(value)
    except Exception:
        return str(value) if value else "TBD"
    if numeric <= 0:
        return "TBD"
    return format_value(numeric, unit)


def _format_transformer_rating(value) -> str:
    if value is None:
        return "TBD"
    try:
        kva = float(value)
    except Exception:
        return "TBD"
    if kva <= 0:
        return "TBD"
    mva = kva / 1000.0
    return f"{mva:.2f} MVA ({kva:.0f} kVA)"


def _svg_bytes_to_png(svg_bytes: bytes, width_px: int = 900) -> Optional[bytes]:
    if not svg_bytes or not CAIROSVG_AVAILABLE:
        return None
    try:
        return cairosvg.svg2png(bytestring=svg_bytes, output_width=width_px)
    except Exception:
        return None


def _validate_efficiency_chain(ctx: ReportContext) -> list[str]:
    """Validate that efficiency chain data is complete and self-consistent.
    
    Ensures all values come from DC SIZING stage1 output and are internally consistent.
    Returns list of warning/error messages (does not block export, but logs issues).
    Caller should log these warnings.
    
    NOTE: This validation is advisory only. Efficiency values are sourced directly
    from DC SIZING stage1 output. If DC SIZING was not run or values are missing,
    the report will still be generated but this function will warn.
    """
    warnings = []
    
    # Verify efficiency values came from stage1 (DC SIZING output)
    if not isinstance(ctx.stage1, dict) or not ctx.stage1:
        warnings.append(
            "Cannot validate efficiency: stage1 (DC SIZING output) is missing or invalid. "
            "Ensure DC SIZING was completed before exporting."
        )
    
    # Check all components are present (not None/zero)
    components = [
        ("DC Cables", ctx.efficiency_components_frac.get("eff_dc_cables_frac")),
        ("PCS", ctx.efficiency_components_frac.get("eff_pcs_frac")),
        ("Transformer (MVT)", ctx.efficiency_components_frac.get("eff_mvt_frac")),
        ("RMU/Switchgear/AC Cables", ctx.efficiency_components_frac.get("eff_ac_cables_sw_rmu_frac")),
        ("HVT/Others", ctx.efficiency_components_frac.get("eff_hvt_others_frac")),
    ]
    
    has_all_components = True
    for name, value in components:
        if value is None or value <= 0:
            warnings.append(f"Efficiency component '{name}' is missing or zero: {value}. "
                           f"Efficiency values should come from DC SIZING stage1 output.")
            has_all_components = False
        elif value > 1.2:  # Allow some margin (e.g., 105% input)
            warnings.append(f"Efficiency component '{name}' exceeds 120%: {value}")
    
    # Check total efficiency
    if ctx.efficiency_chain_oneway_frac is None or ctx.efficiency_chain_oneway_frac <= 0:
        warnings.append(f"Total one-way efficiency chain is missing or zero: {ctx.efficiency_chain_oneway_frac}. "
                       f"Ensure DC SIZING was completed.")
    elif ctx.efficiency_chain_oneway_frac > 1.2:
        warnings.append(f"Total one-way efficiency exceeds 120%: {ctx.efficiency_chain_oneway_frac}")
    
    # If all components present, verify total is their product (with tolerance)
    if has_all_components and ctx.efficiency_chain_oneway_frac and ctx.efficiency_chain_oneway_frac > 0:
        product = 1.0
        for name, value in components:
            if value and value > 0:
                product *= value
        # Allow 2% tolerance for rounding and numerical precision
        relative_error = abs(product - ctx.efficiency_chain_oneway_frac) / ctx.efficiency_chain_oneway_frac if ctx.efficiency_chain_oneway_frac != 0 else 0
        if relative_error > 0.02:
            warnings.append(
                f"Total efficiency ({ctx.efficiency_chain_oneway_frac:.6f}) "
                f"does not match product of components ({product:.6f}). "
                f"Relative error: {relative_error*100:.2f}%. "
                f"Please verify DC SIZING stage1 calculation is complete and consistent."
            )
    
    # Warn if efficiency chain appears uninitialized (less than 0.1% = 0.001)
    if ctx.efficiency_chain_oneway_frac is not None and ctx.efficiency_chain_oneway_frac < 0.001:
        warnings.append("Efficiency chain appears to be zero or uninitialized; "
                       "ensure DC SIZING was completed before exporting.")
    
    return warnings


def _aggregate_ac_block_configs(ctx: ReportContext) -> list[dict]:
    """Aggregate AC Block configurations by signature (PCS count, rating, power per block).
    
    Returns list of dicts: [{"pcs_per_block": int, "pcs_kw": int, "ac_block_power_mw": float, "count": int}, ...]
    All blocks are typically identical, but function handles exceptions.
    """
    if ctx.ac_blocks_total == 0:
        return []
    
    # Build the single (or primary) configuration from context
    pcs_per_block = ctx.pcs_per_block
    
    # Try to get PCS rating from ac_output, fallback to computation
    pcs_kw = None
    if isinstance(ctx.ac_output, dict) and ctx.ac_output.get("pcs_kw"):
        pcs_kw = ctx.ac_output.get("pcs_kw")
    
    if pcs_kw is None and isinstance(ctx.ac_output, dict):
        pcs_kw = ctx.ac_output.get("pcs_power_kw")
    
    ac_block_power_mw = ctx.ac_block_size_mw
    
    # If still no PCS kW and we have block power, derive it
    if pcs_kw is None and ac_block_power_mw and pcs_per_block and pcs_per_block > 0:
        pcs_kw = int((ac_block_power_mw * 1000) / pcs_per_block)
    
    # For now, assume all AC blocks use the same configuration
    # (Future: parse pcs_count_by_block if blocks are heterogeneous)
    return [
        {
            "pcs_per_block": pcs_per_block,
            "pcs_kw": pcs_kw,
            "ac_block_power_mw": ac_block_power_mw,
            "count": ctx.ac_blocks_total,
        }
    ]


def _validate_report_consistency(ctx: ReportContext) -> list[str]:
    """Validate overall report consistency (power/energy/efficiency).
    
    Returns list of warning messages (does not block export, only for logging/QC).
    Warnings address: power balance, energy consistency, efficiency completeness,
    unit consistency, and contract compliance (guarantee year).
    """
    warnings = []
    
    # Efficiency chain validation (source and internal consistency)
    eff_warnings = _validate_efficiency_chain(ctx)
    warnings.extend(eff_warnings)
    
    # AC/DC counts consistency
    if ctx.ac_blocks_total > 0 and ctx.dc_blocks_total == 0:
        warnings.append("AC Blocks present but DC Blocks count is zero.")
    
    # PCS count consistency
    expected_pcs = ctx.ac_blocks_total * ctx.pcs_per_block
    if ctx.pcs_modules_total > 0 and ctx.pcs_modules_total != expected_pcs:
        warnings.append(
            f"PCS module count mismatch: expected {expected_pcs} "
            f"(AC blocks={ctx.ac_blocks_total} × PCS/block={ctx.pcs_per_block}), "
            f"got {ctx.pcs_modules_total}."
        )
    
    # AC power consistency (with tolerance for rounding and intentional overbuild)
    if ctx.ac_blocks_total > 0 and ctx.ac_block_size_mw and ctx.ac_block_size_mw > 0:
        total_ac_power = ctx.ac_blocks_total * ctx.ac_block_size_mw
        poi_requirement = ctx.poi_power_requirement_mw
        
        # Allow 10% tolerance for overbuild (common in BESS sizing)
        overage = total_ac_power - poi_requirement
        overage_pct = (overage / poi_requirement * 100) if poi_requirement > 0 else 0
        if overage > 0.5:  # At least 0.5 MW overbuild is notable
            if overage_pct > 10:
                warnings.append(
                    f"AC power overbuild is {overage_pct:.1f}% "
                    f"(total {total_ac_power:.2f} MW vs requirement {poi_requirement:.2f} MW). "
                    f"This may be intentional or may indicate AC undersizing in the ratio; "
                    f"verify DC-to-AC ratio selection."
                )
    
    # Energy consistency: DC energy should meet POI requirement (through degradation modeling)
    if ctx.dc_total_energy_mwh is not None and ctx.poi_energy_requirement_mwh is not None:
        if ctx.dc_total_energy_mwh < ctx.poi_energy_requirement_mwh:
            warnings.append(
                f"DC nameplate capacity ({ctx.dc_total_energy_mwh:.2f} MWh) is less than "
                f"POI requirement ({ctx.poi_energy_requirement_mwh:.2f} MWh). "
                f"This is expected; degradation modeling determines actual delivery."
            )
    
    # POI usable vs guarantee
    if (ctx.poi_usable_energy_mwh_at_guarantee_year is not None and 
        ctx.poi_energy_guarantee_mwh is not None):
        if ctx.poi_usable_energy_mwh_at_guarantee_year + 0.1 < ctx.poi_energy_guarantee_mwh:
            warnings.append(
                f"POI usable energy at guarantee year ({ctx.poi_usable_energy_mwh_at_guarantee_year:.2f} MWh) "
                f"is below guarantee target ({ctx.poi_energy_guarantee_mwh:.2f} MWh)."
            )
    
    # Guarantee year within project life
    if ctx.poi_guarantee_year > ctx.project_life_years:
        warnings.append(
            f"Guarantee year ({ctx.poi_guarantee_year}) exceeds project life ({ctx.project_life_years} years)."
        )
    
    return warnings


def export_report_v2_1(ctx: ReportContext, brand: dict | None = None) -> bytes:
    doc = Document()
    _setup_margins(doc)
    if brand:
        _setup_header(
            doc,
            title=brand.get("header_title", "Confidential Sizing Report (V2.1 Beta)"),
            logo_path=brand.get("logo_path"),
            header_lines=brand.get("header_lines"),
            footer_lines=brand.get("footer_lines"),
        )
        cover_title = brand.get(
            "cover_title", "CALB Utility-Scale ESS Sizing Report (V2.1 Beta)"
        )
        tool_version = brand.get("tool_version", "V2.1 Beta")
    else:
        _setup_header(doc, title="Confidential Sizing Report (V2.1 Beta)")
        cover_title = "CALB Utility-Scale ESS Sizing Report (V2.1 Beta)"
        tool_version = "V2.1 Beta"

    _add_cover_page(
        doc,
        cover_title,
        ctx.project_name,
        {"tool_version": tool_version},
    )

    doc.add_heading("Conventions & Units", level=2)
    doc.add_paragraph(
        "All efficiencies, losses, DoD, and RTE values are displayed as percentages. "
        "Fractions (0-1) appear only inside formula parentheses."
    )
    doc.add_paragraph("")

    doc.add_heading("Executive Summary", level=2)
    exec_rows = [
        ("POI Power Requirement (MW)", format_value(ctx.poi_power_requirement_mw, "MW")),
        ("POI Energy Requirement (MWh)", format_value(ctx.poi_energy_requirement_mwh, "MWh")),
        ("POI Energy Guarantee (MWh)", format_value(ctx.poi_energy_guarantee_mwh, "MWh")),
        ("Guarantee Year (from COD)", f"{ctx.poi_guarantee_year:d}"),
        ("POI Usable @ Guarantee Year (MWh)", format_value(ctx.poi_usable_energy_mwh_at_guarantee_year, "MWh")),
        ("DC Blocks Total", f"{ctx.dc_blocks_total:d}"),
        ("DC Nameplate @BOL (MWh)", format_value(ctx.dc_total_energy_mwh, "MWh")),
        ("DC Oversize Margin (MWh)", format_value(ctx.stage2.get("oversize_mwh"), "MWh")),
        ("AC Block Template", ctx.ac_block_template_id),
        ("AC Blocks Total", f"{ctx.ac_blocks_total:d}"),
        ("Total PCS Modules", f"{ctx.pcs_modules_total:d}"),
        (
            "Transformer Rating (MVA/kVA)",
            _format_transformer_rating(ctx.transformer_rating_kva),
        ),
        ("Avg DC Blocks per AC Block", f"{ctx.avg_dc_blocks_per_ac_block:.3f}" if ctx.avg_dc_blocks_per_ac_block is not None else ""),
    ]
    _add_table(doc, exec_rows, ["Metric", "Value"])

    if ctx.dc_blocks_allocation:
        doc.add_paragraph("")
        doc.add_paragraph("DC Blocks per AC Block Allocation:")
        alloc_rows = [
            (entry.get("dc_blocks_per_ac_block"), entry.get("ac_blocks_count"))
            for entry in ctx.dc_blocks_allocation
        ]
        _add_table(doc, alloc_rows, ["DC Blocks per AC Block", "Number of AC Blocks"])
    doc.add_paragraph("")

    doc.add_heading("Inputs & Assumptions", level=2)
    doc.add_heading("Site / POI", level=3)
    site_rows = [
        ("POI Power Requirement (MW)", format_value(ctx.poi_power_requirement_mw, "MW")),
        ("POI Energy Requirement (MWh)", format_value(ctx.poi_energy_requirement_mwh, "MWh")),
        ("POI Energy Guarantee (MWh)", format_value(ctx.poi_energy_guarantee_mwh, "MWh")),
        ("POI MV Voltage (kV)", format_value(ctx.grid_mv_voltage_kv_ac, "kV")),
        ("POI Frequency (Hz)", _format_or_tbd(ctx.project_inputs.get("poi_frequency_hz"), "Hz")),
        ("Grid Power Factor (PF)", format_value(ctx.grid_power_factor, "PF")),
    ]
    _add_table(doc, site_rows, ["Parameter", "Value"])
    doc.add_paragraph("")

    doc.add_heading("Stage 1: Energy Requirement", level=2)
    doc.add_paragraph(
        "DC energy capacity required (MWh) = POI energy requirement / "
        "((1 - SC loss) * DoD * sqrt(DC RTE) * One-way Efficiency (DC to POI))."
    )
    doc.add_paragraph(
        f"One-way Efficiency (DC to POI) = {_format_percent_with_fraction(ctx.efficiency_chain_oneway_frac, input_is_fraction=True)}"
    )
    doc.add_paragraph(
        f"SC loss = {_format_percent_with_fraction(ctx.stage1.get('sc_loss_frac') or 0.0, input_is_fraction=True)}; "
        f"DoD = {_format_percent_with_fraction(ctx.stage1.get('dod_frac') or 0.0, input_is_fraction=True)}; "
        f"DC RTE = {_format_percent_with_fraction(ctx.stage1.get('dc_round_trip_efficiency_frac') or 0.0, input_is_fraction=True)}"
    )
    doc.add_paragraph(
        f"RTE Curve Adjustment (Δpp): {float(ctx.stage1.get('rte_curve_adjust_pp') or 0.0):.1f}"
    )
    s1_rows = [
        ("DC Energy Capacity Required (MWh)", format_value(ctx.stage1.get("dc_energy_capacity_required_mwh") or 0.0, "MWh")),
        ("DC Power Required (MW)", format_value(ctx.stage1.get("dc_power_required_mw") or 0.0, "MW")),
    ]
    _add_table(doc, s1_rows, ["Metric", "Value"])
    
    # Efficiency Chain breakdown
    doc.add_paragraph("")
    doc.add_heading("Efficiency Chain (one-way)", level=3)
    doc.add_paragraph(
        "Note: Efficiency chain values (below) represent the one-way conversion path from DC side to AC/POI and "
        "do not include Auxiliary losses or station service loads. "
        "All efficiency and loss values are exclusive of Auxiliary loads. "
        "The product of all component efficiencies yields the total one-way chain efficiency."
    )
    
    eff_rows = [
        ("Total Efficiency (one-way)", format_percent(ctx.efficiency_chain_oneway_frac, input_is_fraction=True)),
        ("DC Cables", format_percent(ctx.efficiency_components_frac.get("eff_dc_cables_frac"), input_is_fraction=True)),
        ("PCS", format_percent(ctx.efficiency_components_frac.get("eff_pcs_frac"), input_is_fraction=True)),
        ("Transformer (MVT)", format_percent(ctx.efficiency_components_frac.get("eff_mvt_frac"), input_is_fraction=True)),
        ("RMU / Switchgear / AC Cables", format_percent(ctx.efficiency_components_frac.get("eff_ac_cables_sw_rmu_frac"), input_is_fraction=True)),
        ("HVT / Others", format_percent(ctx.efficiency_components_frac.get("eff_hvt_others_frac"), input_is_fraction=True)),
    ]
    _add_table(doc, eff_rows, ["Component", "Value"])
    doc.add_paragraph("")

    doc.add_heading("Stage 2: DC Configuration", level=2)
    dc_table = ctx.stage2.get("block_config_table") if isinstance(ctx.stage2, dict) else None
    def _format_mwh_3(value):
        try:
            return f"{float(value):.3f}"
        except Exception:
            return "" if value is None else str(value)

    if dc_table is not None and not dc_table.empty:
        drop_cols = {"Config Adjustment (%)", "Oversize (MWh)", "Busbars Needed (K=10)", "Busbars Needed"}
        dc_columns = [c for c in dc_table.columns if c not in drop_cols]
        headers_map = {c: c for c in dc_columns}
        formatters = {}
        for col in dc_columns:
            if col in ("Unit Capacity (MWh)", "Subtotal (MWh)", "Total DC Nameplate @BOL (MWh)"):
                formatters[col] = _format_mwh_3
            else:
                formatters[col] = lambda v: "" if v is None else str(v)
        _add_dataframe_table(doc, dc_table, dc_columns, headers_map, formatters)
    else:
        doc.add_paragraph("DC block configuration table unavailable.")
    doc.add_paragraph(
        f"DC total nameplate @BOL (MWh): {_format_mwh_3(ctx.dc_total_energy_mwh)}; "
        f"Oversize margin (MWh): {_format_mwh_3(ctx.stage2.get('oversize_mwh'))}."
    )
    doc.add_paragraph("")

    doc.add_heading("Stage 3: Degradation & Deliverable at POI", level=2)
    s3_meta = ctx.stage3_meta if isinstance(ctx.stage3_meta, dict) else {}
    if s3_meta:
        def _fmt_float(value, decimals=2, default=""):
            try:
                return f"{float(value):.{decimals}f}"
            except Exception:
                return default

        poi_power = s3_meta.get("poi_power_mw", ctx.poi_power_requirement_mw)
        dc_power = s3_meta.get("dc_power_mw")
        if dc_power is None and isinstance(ctx.stage1, dict):
            dc_power = ctx.stage1.get("dc_power_required_mw")
        eff_c_rate = s3_meta.get("effective_c_rate")
        soh_profile_id = s3_meta.get("soh_profile_id")
        chosen_soh_c_rate = s3_meta.get("chosen_soh_c_rate")
        chosen_cycles_per_year = s3_meta.get("chosen_soh_cycles_per_year")
        rte_profile_id = s3_meta.get("rte_profile_id")
        chosen_rte_c_rate = s3_meta.get("chosen_rte_c_rate")

        doc.add_paragraph(
            f"POI Power = {_fmt_float(poi_power, 2)} MW | "
            f"DC-equivalent Power = {_fmt_float(dc_power, 2)} MW | "
            f"Effective C-rate (DC-side) = {_fmt_float(eff_c_rate, 3)} C"
        )
        doc.add_paragraph(
            f"SOH profile ID = {soh_profile_id} "
            f"(C-rate \u2248 {chosen_soh_c_rate}, cycles/year = {chosen_cycles_per_year}); "
            f"RTE profile ID = {rte_profile_id} (C-rate \u2248 {chosen_rte_c_rate})."
        )
        doc.add_paragraph(
            f"Guarantee Year (from COD) = {ctx.poi_guarantee_year} | "
            f"POI Energy Target = {format_value(ctx.poi_energy_guarantee_mwh, 'MWh')} MWh"
        )
        doc.add_paragraph("")
    s3_df = ctx.stage3_df
    if s3_df is not None and not s3_df.empty:
        doc.add_paragraph("System RTE = DC RTE * (One-way Efficiency)^2. Note: One-way Efficiency refers to DC-to-POI efficiency.")
        rte_dc = s3_df["DC_RTE_Pct"].astype(float)
        rte_sys = s3_df["System_RTE_Pct"].astype(float)
        rte_dc_min, rte_dc_max = float(rte_dc.min()), float(rte_dc.max())
        rte_sys_min, rte_sys_max = float(rte_sys.min()), float(rte_sys.max())

        if abs(rte_dc_min - rte_dc_max) < 1e-6:
            doc.add_paragraph(f"DC RTE: {format_percent(rte_dc_min, input_is_fraction=False)}")
        else:
            doc.add_paragraph(
                f"DC RTE varies by year: {format_percent(rte_dc_min, input_is_fraction=False)} to "
                f"{format_percent(rte_dc_max, input_is_fraction=False)}"
            )
        if abs(rte_sys_min - rte_sys_max) < 1e-6:
            doc.add_paragraph(f"System RTE: {format_percent(rte_sys_min, input_is_fraction=False)}")
        else:
            doc.add_paragraph(
                f"System RTE varies by year: {format_percent(rte_sys_min, input_is_fraction=False)} to "
                f"{format_percent(rte_sys_max, input_is_fraction=False)}"
            )

        s3_df = s3_df.copy()
        s3_df["Meets_Guarantee"] = s3_df["POI_Usable_Energy_MWh"] >= float(ctx.poi_energy_guarantee_mwh or 0.0)
        doc.add_paragraph(
            f"Guarantee Year = {ctx.poi_guarantee_year}; pass/fail is evaluated at year {ctx.poi_guarantee_year}."
        )

        try:
            available_years = set(s3_df["Year_Index"].astype(int).tolist())
        except Exception:
            available_years = set()
        
        # Keep full data for chart (all years)
        s3_df_full = s3_df.copy()
        


        cap_chart = _plot_dc_capacity_bar_png(
            bol_mwh=ctx.dc_total_energy_mwh,
            s3_df=s3_df,
            guarantee_year=ctx.poi_guarantee_year,
            title="DC Block Energy (BOL/COD/Yx at POI)",
        )
        if cap_chart is not None and cap_chart.getbuffer().nbytes > 0:
            doc.add_paragraph("")
            doc.add_picture(cap_chart, width=Inches(6.7))

        chart = _plot_poi_usable_png(
            s3_df_full,
            poi_target=ctx.poi_energy_guarantee_mwh,
            title="POI Usable Energy vs Year",
        )
        if chart is not None and chart.getbuffer().nbytes > 0:
            doc.add_paragraph("")
            doc.add_picture(chart, width=Inches(6.7))
        else:
            # Show diagnostic if recompute failed
            err = None
            try:
                err = ctx.stage3_meta.get("error") if isinstance(ctx.stage3_meta, dict) else None
            except Exception:
                err = None
            if err:
                doc.add_paragraph(f"Stage 3 data unavailable: {err}")
            else:
                doc.add_paragraph("Stage 3 data unavailable.")
        # Add full year-by-year table after the chart
        doc.add_paragraph("Year-by-Year Degradation & Deliverable at POI:")
        s3_columns = [
            "Year_Index",
            "SOH_Display_Pct",
            "SOH_Absolute_Pct",
            "DC_Usable_MWh",
            "POI_Usable_Energy_MWh",
            "DC_RTE_Pct",
            "System_RTE_Pct",
        ]
        headers_map = {
            "Year_Index": "Year (From COD)",
            "SOH_Display_Pct": "SOH @ COD Baseline (%)",
            "SOH_Absolute_Pct": "SOH Vs FAT (%)",
            "DC_Usable_MWh": "DC Usable (MWh)",
            "POI_Usable_Energy_MWh": "POI Usable (MWh)",
            "DC_RTE_Pct": "DC RTE (%)",
            "System_RTE_Pct": "System RTE (%)",
        }
        formatters = {
            "Year_Index": lambda v: f"{int(v)}",
            "SOH_Display_Pct": lambda v: format_percent(v, input_is_fraction=False),
            "SOH_Absolute_Pct": lambda v: format_percent(v, input_is_fraction=False),
            "DC_Usable_MWh": lambda v: format_value(v, "MWh"),
            "POI_Usable_Energy_MWh": lambda v: format_value(v, "MWh"),
            "DC_RTE_Pct": lambda v: format_percent(v, input_is_fraction=False),
            "System_RTE_Pct": lambda v: format_percent(v, input_is_fraction=False),
        }
        _add_dataframe_table(doc, s3_df_full, s3_columns, headers_map, formatters)
        doc.add_paragraph("")

    doc.add_heading("Stage 4: AC Block Sizing", level=2)
    
    # Extract AC configuration data from ac_output and context
    ac_ratio = ctx.ac_output.get("selected_ratio") if isinstance(ctx.ac_output, dict) else None
    ac_pcs_per_block = ctx.ac_output.get("pcs_per_block") if isinstance(ctx.ac_output, dict) else ctx.pcs_per_block
    # PCS kW: try ac_output.pcs_kw first, fallback to pcs_power_kw, then compute from AC block size
    ac_pcs_kw = ctx.ac_output.get("pcs_kw") if isinstance(ctx.ac_output, dict) else None
    if ac_pcs_kw is None and isinstance(ctx.ac_output, dict):
        ac_pcs_kw = ctx.ac_output.get("pcs_power_kw")
    dc_blocks_per_ac = ctx.ac_output.get("dc_blocks_per_ac") if isinstance(ctx.ac_output, dict) else None
    
    if ac_ratio:
        doc.add_paragraph(f"AC:DC Ratio: {ac_ratio} (1 AC Block per specified DC Blocks)")
        doc.add_paragraph("")
    
    # Show AC Block Configuration Summary if we have key data
    if (ac_pcs_per_block and ctx.ac_blocks_total > 0) and (ac_pcs_kw or ctx.ac_block_size_mw):
        doc.add_heading("AC Block Configuration Summary", level=3)
        
        # Use pcs_kw if available, otherwise derive from AC block power
        pcs_rating = ac_pcs_kw
        if pcs_rating is None and ctx.ac_block_size_mw and ac_pcs_per_block and ac_pcs_per_block > 0:
            pcs_rating = (ctx.ac_block_size_mw * 1000) / ac_pcs_per_block
        
        ac_config_rows = [
            ("PCS per AC Block", f"{ac_pcs_per_block}"),
        ]
        if pcs_rating:
            ac_config_rows.append(("PCS Rating", f"{pcs_rating:.0f} kW"))
            ac_config_rows.append(("AC Block Power per Block", f"{ac_pcs_per_block * pcs_rating / 1000:.2f} MW"))
        ac_config_rows.append(("Total AC Blocks", f"{ctx.ac_blocks_total}"))
        
        if ac_ratio:
            ac_config_rows.insert(0, ("AC:DC Ratio", ac_ratio))
        
        _add_table(doc, ac_config_rows, ["Parameter", "Value"])
        doc.add_paragraph("")
    
    transformer_mva = None
    if ctx.grid_power_factor and ctx.grid_power_factor > 0 and ctx.ac_block_size_mw:
        transformer_mva = ctx.ac_block_size_mw / ctx.grid_power_factor
    transformer_formula = "n/a"
    if transformer_mva is not None and ctx.ac_block_size_mw and ctx.grid_power_factor:
        transformer_formula = (
            f"{format_value(ctx.ac_block_size_mw, 'MW')} / {format_value(ctx.grid_power_factor, 'PF')} = "
            f"{format_value(transformer_mva, 'MVA')}"
        )
    pcs_count_by_block = ctx.ac_output.get("pcs_count_by_block") if isinstance(ctx.ac_output, dict) else None
    pcs_by_block_text = ""
    if isinstance(pcs_count_by_block, list) and pcs_count_by_block:
        pcs_by_block_text = ", ".join(
            f"B{idx + 1}={int(value)}" for idx, value in enumerate(pcs_count_by_block)
        )
    s4_rows = [
        ("AC Block Template", ctx.ac_block_template_id),
        ("AC Block Size (MW)", format_value(ctx.ac_block_size_mw, "MW")),
        ("PCS per AC Block", f"{ctx.pcs_per_block:d}"),
        ("Feeders per AC Block", f"{ctx.feeders_per_block:d}"),
        ("PCS LV Voltage (V_LL)", format_value(ctx.pcs_lv_voltage_v_ll_rms_ac, "V")),
        ("Total PCS Modules", f"{ctx.pcs_modules_total:d}"),
        ("Transformer Rating (kVA)", format_value(ctx.transformer_rating_kva, "kVA")),
        ("Transformer Formula (MVA)", transformer_formula),
    ]
    if pcs_by_block_text:
        s4_rows.insert(3, ("PCS per AC Block (by block)", pcs_by_block_text))
    _add_table(doc, s4_rows, ["Metric", "Value"])
    doc.add_paragraph("")

    doc.add_heading("Integrated Configuration Summary", level=2)
    combined_rows = [
        ("DC Blocks Total", f"{ctx.dc_blocks_total:d}", ""),
        ("AC Blocks Total", "", f"{ctx.ac_blocks_total:d}"),
        ("Total PCS Modules", "", f"{ctx.pcs_modules_total:d}"),
        ("Transformer Rating (MVA/kVA)", "", _format_transformer_rating(ctx.transformer_rating_kva)),
        ("Grid MV Voltage (kV)", "", format_value(ctx.grid_mv_voltage_kv_ac, "kV")),
        ("Guarantee Year (from COD)", f"{ctx.poi_guarantee_year:d}", ""),
    ]
    _add_table(doc, combined_rows, ["Metric", "DC", "AC"])
    doc.add_paragraph("")

    figure_index = 1

    doc.add_heading("Single Line Diagram (1 AC Block group)", level=2)
    doc.add_paragraph("SLD output represents a single MV node chain (RMU -> TR -> 1 AC block group).")
    if ctx.sld_snapshot_hash:
        generated_at = ctx.sld_generated_at or "unknown time"
        doc.add_paragraph(
            f"SLD snapshot hash: {ctx.sld_snapshot_hash} (generated at {generated_at})."
        )
        if ctx.sld_group_index:
            doc.add_paragraph(f"SLD preview group index: {ctx.sld_group_index}.")
    sld_embedded = False
    if ctx.sld_pro_png_bytes:
        doc.add_picture(io.BytesIO(ctx.sld_pro_png_bytes), width=Inches(6.7))
        doc.add_paragraph(f"Figure {figure_index} - Single Line Diagram (auto-generated)")
        figure_index += 1
        sld_embedded = True
    elif ctx.sld_preview_svg_bytes:
        png_bytes = _svg_bytes_to_png(ctx.sld_preview_svg_bytes)
        if png_bytes:
            doc.add_picture(io.BytesIO(png_bytes), width=Inches(6.7))
            doc.add_paragraph(f"Figure {figure_index} - Single Line Diagram (auto-generated)")
            figure_index += 1
            sld_embedded = True

    if not sld_embedded:
        doc.add_paragraph("SLD not generated. Please generate in Single Line Diagram page.")
    doc.add_paragraph("")

    doc.add_heading("Block Layout (template view)", level=2)
    layout_png_bytes = ctx.layout_png_bytes
    if not layout_png_bytes and ctx.layout_svg_bytes:
        layout_png_bytes = _svg_bytes_to_png(ctx.layout_svg_bytes)
    if layout_png_bytes:
        doc.add_picture(io.BytesIO(layout_png_bytes), width=Inches(6.7))
        doc.add_paragraph(f"Figure {figure_index} - Block Layout (auto-generated)")
        figure_index += 1
    else:
        doc.add_paragraph("Layout not generated. Please generate in Site Layout page.")
    doc.add_paragraph("")

    qc_checks = list(ctx.qc_checks)
    
    # Add consistency validation warnings
    consistency_warnings = _validate_report_consistency(ctx)
    qc_checks.extend(consistency_warnings)
    
    percent_pairs = [
        (ctx.efficiency_chain_oneway_frac, format_percent(ctx.efficiency_chain_oneway_frac, input_is_fraction=True)),
        (ctx.stage1.get("dod_frac"), format_percent(ctx.stage1.get("dod_frac"), input_is_fraction=True)),
        (ctx.stage1.get("sc_loss_frac"), format_percent(ctx.stage1.get("sc_loss_frac"), input_is_fraction=True)),
        (
            ctx.stage1.get("dc_round_trip_efficiency_frac"),
            format_percent(ctx.stage1.get("dc_round_trip_efficiency_frac"), input_is_fraction=True),
        ),
    ]
    for raw_value, text in percent_pairs:
        if raw_value is None:
            continue
        try:
            raw_value = float(raw_value)
            displayed = float(text.replace("%", ""))
        except Exception:
            continue
        if raw_value > 0.1 and displayed < 1.0:
            qc_checks.append("Percent formatting appears to show a fraction (0.xxx%) instead of a percent.")
            break

    doc.add_heading("QC / Warnings", level=2)
    if qc_checks:
        for item in qc_checks:
            doc.add_paragraph(f"- {item}")
    else:
        doc.add_paragraph("No QC warnings.")

    return _doc_to_bytes(doc)


export_report_v2 = export_report_v2_1
