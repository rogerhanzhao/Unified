import datetime
import io
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from calb_sizing_tool.config import AC_DATA_PATH, DC_DATA_PATH, PROJECT_ROOT

# ----------------------------------------
# Shared DOCX helpers (match DC report style)
# ----------------------------------------


def _resolve_logo_path() -> Path | None:
    candidates = [
        PROJECT_ROOT / "calb_assets" / "logo" / "calb_logo.png",
        PROJECT_ROOT / "calb_logo.png",
    ]
    for candidate in candidates:
        try:
            if candidate.exists():
                return candidate
        except Exception:
            continue
    return None


def _setup_margins(doc: Document):
    section = doc.sections[0]
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)


def add_header_logo(document: Document, logo_path_or_bytes) -> list:
    tables = []
    for section in document.sections:
        header = section.header
        header.is_linked_to_previous = False
        for para in list(header.paragraphs):
            para._element.getparent().remove(para._element)
        header_table = header.add_table(rows=1, cols=2, width=Inches(6.9))
        tables.append(header_table)
        if logo_path_or_bytes:
            p_logo = header_table.rows[0].cells[0].paragraphs[0]
            run_logo = p_logo.add_run()
            if isinstance(logo_path_or_bytes, (str, Path)):
                run_logo.add_picture(str(logo_path_or_bytes), width=Inches(1.2))
            else:
                run_logo.add_picture(io.BytesIO(logo_path_or_bytes), width=Inches(1.2))
    return tables


def _setup_header(doc: Document, title: str = "Confidential Sizing Report"):
    logo_path = _resolve_logo_path()
    header_tables = add_header_logo(doc, logo_path)

    for header_table in header_tables:
        hdr_cells = header_table.rows[0].cells
        p_info = hdr_cells[1].paragraphs[0]
        p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_info = p_info.add_run(
            "CALB Group Co., Ltd.\n"
            "Utility-Scale Energy Storage Systems\n"
            f"{title}"
        )
        run_info.font.size = Pt(9)
        run_info.font.bold = True


def _format_float(value, decimals):
    if value is None:
        return ""
    try:
        return f"{float(value):.{decimals}f}"
    except Exception:
        return str(value)


def _add_table(doc: Document, rows, headers):
    if not rows:
        doc.add_paragraph("No data available.")
        return

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"

    hdr = tbl.rows[0].cells
    for j, col in enumerate(headers):
        hdr[j].text = str(col)
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for row in rows:
        rc = tbl.add_row().cells
        for j, val in enumerate(row):
            rc[j].text = "" if val is None else str(val)


def _doc_to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _svg_bytes_to_png(svg_bytes: bytes) -> bytes | None:
    if not svg_bytes:
        return None
    try:
        import cairosvg
    except Exception:
        return None
    try:
        return cairosvg.svg2png(bytestring=svg_bytes)
    except Exception:
        return None


def _resolve_diagram_bytes(ctx: dict | None, png_key: str, svg_key: str) -> bytes | None:
    if not ctx:
        return None
    png_bytes = ctx.get(png_key)
    if png_bytes:
        return png_bytes
    svg_bytes = ctx.get(svg_key)
    if svg_bytes:
        return _svg_bytes_to_png(svg_bytes)
    return None


# ----------------------------------------
# Cover + Appendix helpers
# ----------------------------------------


def _get_commit_hash():
    try:
        head_path = PROJECT_ROOT / ".git" / "HEAD"
        if not head_path.exists():
            return "unknown"
        head = head_path.read_text(encoding="utf-8").strip()
        if head.startswith("ref:"):
            ref_path = head.split(" ", 1)[1].strip()
            ref_file = PROJECT_ROOT / ".git" / ref_path
            if ref_file.exists():
                return ref_file.read_text(encoding="utf-8").strip()[:7]
        return head[:7]
    except Exception:
        return "unknown"


def _resolve_tool_version(ctx: dict):
    if ctx and ctx.get("tool_version"):
        return str(ctx.get("tool_version"))
    return "V1.0"


def _add_cover_page(doc: Document, title: str, project_name: str, ctx: dict):
    tool_version = _resolve_tool_version(ctx)
    commit_hash = ctx.get("commit_hash") if ctx else None
    if not commit_hash:
        commit_hash = _get_commit_hash()

    doc.add_paragraph("")
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(
        f"Project: {project_name}\n"
        f"Date: {datetime.datetime.now().strftime('%Y-%m-%d')}\n"
        f"Tool Version: {tool_version}\n"
        f"Commit: {commit_hash}"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def _add_appendix(doc: Document, ctx: dict):
    return


# ----------------------------------------
# DC dictionary extraction
# ----------------------------------------


def _match_pack_from_block_name(packs: pd.DataFrame, block_name: str, block_code: str):
    if packs is None or packs.empty:
        return None

    candidates = packs.copy()
    if "Is_Active" in candidates.columns:
        active = candidates[candidates["Is_Active"] == 1]
        if not active.empty:
            candidates = active

    search = f"{block_name} {block_code}".lower()
    for _, row in candidates.iterrows():
        model = str(row.get("Pack_Model", "")).lower()
        if model and model in search:
            return row
        tail = model.split("calb_")[-1] if "calb_" in model else model
        if tail and tail in search:
            return row

    return candidates.iloc[0] if not candidates.empty else None


def _voltage_range_from_row(row: pd.Series):
    if row is None:
        return None

    min_col = None
    max_col = None
    for col in row.index:
        label = str(col).lower()
        if "voltage" in label and "min" in label:
            min_col = col
        if "voltage" in label and "max" in label:
            max_col = col

    if min_col and max_col:
        try:
            v_min = float(row[min_col])
            v_max = float(row[max_col])
            if v_min > 0 and v_max > 0:
                return f"{v_min:.0f}-{v_max:.0f} V DC"
        except Exception:
            return None

    return None


def extract_dc_equipment_spec(block_code=None, block_name=None, data_path=None):
    data_path = Path(data_path) if data_path else Path(DC_DATA_PATH)
    if not data_path.exists():
        raise FileNotFoundError(f"DC dictionary not found at {data_path}")

    xls = pd.ExcelFile(data_path)
    blocks = pd.read_excel(xls, "dc_block_template_314_data")
    racks = pd.read_excel(xls, "rack_type_314_data")
    packs = pd.read_excel(xls, "pack_type_314_data")
    cells = pd.read_excel(xls, "battery_cell_type_314_data")

    if "Is_Active" in blocks.columns:
        blocks = blocks[blocks["Is_Active"] == 1]

    match = pd.DataFrame()
    if block_code and "Dc_Block_Code" in blocks.columns:
        match = blocks[blocks["Dc_Block_Code"] == block_code]
    if match.empty and block_name and "Dc_Block_Name" in blocks.columns:
        match = blocks[blocks["Dc_Block_Name"] == block_name]
    if match.empty:
        default = blocks[blocks.get("Is_Default_Option", 0) == 1]
        if not default.empty:
            container = default[default["Block_Form"].astype(str).str.lower() == "container"]
            match = container if not container.empty else default
    if match.empty:
        match = blocks.head(1)
    if match.empty:
        return {
            "container_model": "",
            "cell_type": "",
            "configuration": "",
            "unit_capacity_mwh": None,
            "nominal_voltage_v": None,
            "voltage_range_v": None,
        }

    block_row = match.iloc[0]
    rack_row = None
    pack_row = None

    rack_type_id = block_row.get("Rack_Type_Id")
    if pd.notna(rack_type_id) and "Rack_Type_Id" in racks.columns:
        rack = racks[racks["Rack_Type_Id"] == rack_type_id]
        if not rack.empty:
            rack_row = rack.iloc[0]
            pack_type_id = rack_row.get("Pack_Type_Id")
            if pd.notna(pack_type_id) and "Pack_Type_Id" in packs.columns:
                pack = packs[packs["Pack_Type_Id"] == pack_type_id]
                if not pack.empty:
                    pack_row = pack.iloc[0]

    if pack_row is None:
        pack_row = _match_pack_from_block_name(
            packs,
            str(block_row.get("Dc_Block_Name", "")),
            str(block_row.get("Dc_Block_Code", "")),
        )

    cell_row = None
    if pack_row is not None and "Cell_Type_Id" in packs.columns:
        cell_type_id = pack_row.get("Cell_Type_Id")
        if pd.notna(cell_type_id) and "Cell_Type_Id" in cells.columns:
            cell = cells[cells["Cell_Type_Id"] == cell_type_id]
            if not cell.empty:
                cell_row = cell.iloc[0]

    container_model = str(block_row.get("Dc_Block_Name") or block_row.get("Dc_Block_Code") or "")
    cell_type = str(cell_row.get("Cell_Model")) if cell_row is not None else ""

    configuration = ""
    racks_per_block = block_row.get("Racks_Per_Block")
    packs_per_block = block_row.get("Packs_Per_Block")
    if pd.notna(racks_per_block):
        configuration = f"{int(racks_per_block)} Racks"
    elif pd.notna(packs_per_block):
        configuration = f"{int(packs_per_block)} Packs"

    unit_capacity_mwh = None
    if pd.notna(block_row.get("Block_Nameplate_Capacity_Mwh")):
        unit_capacity_mwh = float(block_row.get("Block_Nameplate_Capacity_Mwh"))

    nominal_voltage_v = None
    pack_nominal = None
    if pack_row is not None and pd.notna(pack_row.get("Pack_Nominal_Voltage_V")):
        pack_nominal = float(pack_row.get("Pack_Nominal_Voltage_V"))

    if pack_nominal is not None:
        packs_per_rack = rack_row.get("Packs_Per_Rack") if rack_row is not None else None
        if pd.notna(packs_per_rack):
            nominal_voltage_v = pack_nominal * float(packs_per_rack)
        elif pd.notna(packs_per_block):
            nominal_voltage_v = pack_nominal * float(packs_per_block)
        else:
            nominal_voltage_v = pack_nominal

    voltage_range = _voltage_range_from_row(block_row)
    if voltage_range is None and pack_row is not None:
        voltage_range = _voltage_range_from_row(pack_row)

    return {
        "container_model": container_model,
        "cell_type": cell_type,
        "configuration": configuration,
        "unit_capacity_mwh": unit_capacity_mwh,
        "nominal_voltage_v": nominal_voltage_v,
        "voltage_range_v": voltage_range,
        "block_code": str(block_row.get("Dc_Block_Code", "")),
    }


# ----------------------------------------
# DC report helper (reuse DC content)
# ----------------------------------------


def _resolve_dc_report_data(dc_output: dict):
    try:
        from calb_sizing_tool.ui import dc_view
    except Exception:
        return None, None, None

    stage1 = dc_output.get("stage1", {}) if dc_output else {}
    if not stage1:
        return None, None, None

    if "results_dict" in dc_output and "report_order" in dc_output:
        return stage1, dc_output.get("results_dict"), dc_output.get("report_order")

    selected = dc_output.get("selected_scenario", stage1.get("selected_scenario", "container_only"))
    try:
        _, df_blocks, df_soh_profile, df_soh_curve, df_rte_profile, df_rte_curve = dc_view.load_data(DC_DATA_PATH)
        s2, s3_df, s3_meta, iter_count, poi_g, converged = dc_view.size_with_guarantee(
            stage1,
            selected,
            df_blocks,
            df_soh_profile,
            df_soh_curve,
            df_rte_profile,
            df_rte_curve,
            k_max=dc_view.K_MAX_FIXED,
        )
    except Exception:
        return stage1, None, None

    results_dict = {
        selected: (s2, s3_df, s3_meta, iter_count, poi_g, converged)
    }
    report_order = [(selected, selected.replace("_", " ").title())]
    return stage1, results_dict, report_order


def _dc_section_heading(chapter_prefix: str, index: int, title: str):
    if str(chapter_prefix) == "1":
        return f"{index}. {title}"
    return f"{chapter_prefix}.{index} {title}"


def _append_dc_report_sections(doc: Document, dc_output: dict, ctx: dict, chapter_prefix: str = "3"):
    stage1, results_dict, report_order = _resolve_dc_report_data(dc_output or {})
    if stage1 is None or results_dict is None or report_order is None:
        doc.add_paragraph("DC report section unavailable.")
        return False

    try:
        from calb_sizing_tool.ui import dc_view
    except Exception:
        doc.add_paragraph("DC report section unavailable.")
        return False

    heading_level = 2 if str(chapter_prefix) == "1" else 3

    doc.add_heading(
        _dc_section_heading(chapter_prefix, 1, "Project Summary"),
        level=heading_level,
    )
    p = doc.add_paragraph()
    p.add_run(f"Project life: {int(stage1['project_life_years'])} years\n")
    p.add_run(f"POI guarantee year: {int(stage1.get('poi_guarantee_year', 0))}\n")
    p.add_run(f"Cycles per year (assumed): {int(stage1['cycles_per_year'])}\n")
    p.add_run(
        f"S&C time from FAT to COD: {int(round(stage1.get('sc_time_months', 0)))} months\n"
    )
    p.add_run(
        f"DC\u2192POI efficiency chain (one-way): {stage1.get('eff_dc_to_poi_frac', 0.0)*100:.2f}%\n"
    )
    p.add_run(f"POI\u2192DC equivalent power: {stage1.get('dc_power_required_mw', 0.0):.2f} MW")

    doc.add_paragraph(
        "This sizing report is based on the 314 Ah cell database and the internal "
        "CALB SOH/RTE profiles for the selected operating conditions."
    )

    doc.add_heading(
        _dc_section_heading(chapter_prefix, 2, "Equipment Summary (DC Blocks)"),
        level=heading_level,
    )

    for key, title in report_order:
        if key not in results_dict:
            continue
        s2, _, _, iter_count, poi_g, converged = results_dict[key]
        doc.add_paragraph(title, style=None)
        dc_view._docx_add_config_table(doc, s2.get("block_config_table"))
        doc.add_paragraph(f"Iterations: {iter_count} | Guarantee met: {bool(converged)}")
        if poi_g is not None:
            doc.add_paragraph(f"POI usable energy @ guarantee year: {poi_g:.2f} MWh")
        doc.add_paragraph("")

    doc.add_heading(
        _dc_section_heading(
            chapter_prefix,
            3,
            "Lifetime POI Usable Energy & SOH (Per Configuration)",
        ),
        level=heading_level,
    )

    poi_target = float(stage1["poi_energy_req_mwh"])
    guarantee_year = int(stage1.get("poi_guarantee_year", 0))

    for key, title in report_order:
        if key not in results_dict:
            continue
        _, s3_df, s3_meta, _, _, _ = results_dict[key]

        doc.add_paragraph(title, style=None)
        doc.add_paragraph(
            f"POI Power = {s3_meta.get('poi_power_mw', 0.0):.2f} MW | "
            f"DC-equivalent Power = {s3_meta.get('dc_power_mw', 0.0):.2f} MW | "
            f"Effective C-rate (DC-side) = {s3_meta.get('effective_c_rate', 0.0):.3f} C"
        )
        doc.add_paragraph(
            f"SOH profile ID = {s3_meta.get('soh_profile_id')} "
            f"(C-rate \u2248 {s3_meta.get('chosen_soh_c_rate')}, cycles/year = {s3_meta.get('chosen_soh_cycles_per_year')}); "
            f"RTE profile ID = {s3_meta.get('rte_profile_id')} (C-rate \u2248 {s3_meta.get('chosen_rte_c_rate')})."
        )
        doc.add_paragraph(
            f"Guarantee Year (from COD) = {guarantee_year} | POI Energy Target = {poi_target:.2f} MWh"
        )

        if dc_view.MATPLOTLIB_AVAILABLE:
            try:
                cap_png = dc_view._plot_dc_capacity_bar_png(
                    s2=s2,
                    s3_df=s3_df,
                    guarantee_year=guarantee_year,
                    title="DC Block Energy (BOL/COD/Yx at POI)",
                )
                if cap_png and cap_png.getbuffer().nbytes > 0:
                    doc.add_picture(cap_png, width=Inches(6.7))
                png = dc_view._plot_poi_usable_png(
                    s3_df=s3_df,
                    poi_target=poi_target,
                    title=f"POI Usable Energy vs Year \u2013 {key}",
                )
                if png and png.getbuffer().nbytes > 0:
                    doc.add_picture(png, width=Inches(6.7))
            except Exception:
                doc.add_paragraph("Chart export skipped due to plotting error.")
        else:
            doc.add_paragraph("Chart export skipped (matplotlib not available).")

        dc_view._docx_add_lifetime_table(doc, s3_df)
        doc.add_paragraph("")

    p_final = doc.add_paragraph()
    p_fmt = p_final.paragraph_format
    p_fmt.space_before = Pt(0)
    p_fmt.space_after = Pt(0)
    p_fmt.line_spacing = Pt(0)
    run_final = p_final.add_run()
    run_final.font.size = Pt(1)
    return True


# ----------------------------------------
# AC report helpers
# ----------------------------------------


def _ac_section_heading(chapter_prefix: str, index: int, title: str):
    if str(chapter_prefix) == "1":
        return f"{index}. {title}"
    return f"{chapter_prefix}.{index} {title}"


def _build_ac_exec_summary_rows(ac_output: dict):
    rows = [
        ("POI Power (MW)", _format_float(ac_output.get("poi_power_mw"), 2)),
    ]
    if ac_output.get("poi_energy_mwh") is not None:
        rows.append(("POI Energy (MWh)", _format_float(ac_output.get("poi_energy_mwh"), 2)))
    rows.extend(
        [
            ("Total AC Blocks", _format_float(ac_output.get("num_blocks"), 0)),
            ("Total PCS Modules", _format_float(ac_output.get("total_pcs"), 0)),
            ("Transformer Count", _format_float(ac_output.get("transformer_count"), 0)),
            ("MV Level (kV)", _format_float(ac_output.get("grid_kv"), 1)),
        ]
    )
    return rows


def _build_ac_inputs_rows(ac_output: dict, ctx: dict):
    rows = []
    if ctx and ctx.get("inputs"):
        for key, value in ctx.get("inputs").items():
            rows.append((key, value))
    else:
        rows = [
            ("Project Name", ac_output.get("project_name", "")),
            ("POI Power Requirement (MW)", _format_float(ac_output.get("poi_power_mw"), 2)),
            ("POI Energy Requirement (MWh)", _format_float(ac_output.get("poi_energy_mwh"), 2)),
            ("Grid Voltage (kV)", _format_float(ac_output.get("grid_kv"), 1)),
            ("PCS AC Output Voltage (V_LL,rms)", _format_float(ac_output.get("inverter_lv_v"), 0)),
            ("Standard AC Block Size (MW)", _format_float(ac_output.get("block_size_mw"), 2)),
        ]
    return rows


def _build_ac_results_rows(ac_output: dict):
    rows = [
        ("AC Blocks", _format_float(ac_output.get("num_blocks"), 0)),
        ("Total AC Power (MW)", _format_float(ac_output.get("total_ac_mw"), 2)),
        ("Overhead Margin (MW)", _format_float(ac_output.get("overhead_mw"), 2)),
        ("Total PCS Modules", _format_float(ac_output.get("total_pcs"), 0)),
        ("Transformer Count", _format_float(ac_output.get("transformer_count"), 0)),
    ]
    if ac_output.get("dc_blocks_per_ac") is not None:
        rows.append(("DC Blocks per AC Block", _format_float(ac_output.get("dc_blocks_per_ac"), 0)))
    return rows


def _build_ac_equipment_rows(ac_output: dict):
    return [
        ("Standard AC Block Size (MW)", _format_float(ac_output.get("block_size_mw"), 3)),
        ("PCS Rating (kW)", _format_float(ac_output.get("pcs_power_kw"), 0)),
        ("PCS per AC Block", _format_float(ac_output.get("pcs_per_block"), 0)),
        ("Transformer Rating (kVA)", _format_float(ac_output.get("transformer_kva"), 0)),
        ("MV Voltage Level (kV)", _format_float(ac_output.get("grid_kv"), 1)),
        ("LV Voltage Level (V)", _format_float(ac_output.get("inverter_lv_v"), 0)),
    ]


def _append_ac_report_sections(doc: Document, ac_output: dict, ctx: dict, chapter_prefix: str = "1"):
    heading_level = 2 if str(chapter_prefix) == "1" else 3

    doc.add_heading(
        _ac_section_heading(chapter_prefix, 1, "Executive Summary"),
        level=heading_level,
    )
    _add_table(doc, _build_ac_exec_summary_rows(ac_output), ["Metric", "Value"])
    doc.add_paragraph("")

    doc.add_heading(
        _ac_section_heading(chapter_prefix, 2, "Project Inputs & Assumptions"),
        level=heading_level,
    )
    _add_table(doc, _build_ac_inputs_rows(ac_output, ctx), ["Parameter", "Value"])
    doc.add_paragraph("")

    doc.add_heading(
        _ac_section_heading(chapter_prefix, 3, "AC Sizing Results"),
        level=heading_level,
    )
    _add_table(doc, _build_ac_results_rows(ac_output), ["Metric", "Value"])
    doc.add_paragraph("")

    doc.add_heading(
        _ac_section_heading(chapter_prefix, 4, "AC Equipment Specification"),
        level=heading_level,
    )
    _add_table(doc, _build_ac_equipment_rows(ac_output), ["Item", "Specification"])


# ----------------------------------------
# Public report generators
# ----------------------------------------


def make_report_filename(proj_name, suffix="Report"):
    safe = "".join(c if c.isalnum() else "_" for c in proj_name)
    return f"{safe}_{suffix}.docx"


def sanitize_filename(text: str, max_length: int = 80) -> str:
    cleaned = re.sub(r'[\\/:*?\"<>|\x00-\x1f]', '', text or '')
    cleaned = re.sub(r'\s+', '_', cleaned.strip())
    cleaned = cleaned.strip('._')
    if max_length and len(cleaned) > max_length:
        cleaned = cleaned[:max_length].rstrip('_')
    return cleaned


def make_proposal_filename(project_name: str | None) -> str:
    stamp = datetime.date.today().strftime('%Y%m%d')
    safe_project = sanitize_filename(project_name or '')
    if safe_project:
        return f'CALB_BESS_Proposal_{safe_project}_{stamp}.docx'
    return f'CALB_BESS_Proposal_{stamp}.docx'


def create_dc_report(dc_output: dict, ctx: dict) -> bytes:
    """Internal DC-only generator for regression testing against dc_view."""
    doc = Document()
    _setup_margins(doc)
    _setup_header(doc)

    stage1 = dc_output.get("stage1", {}) if dc_output else {}
    project_name = stage1.get("project_name", "CALB ESS Project")

    doc.add_heading("CALB Utility-Scale ESS Sizing Report", level=1)
    doc.add_paragraph(f"Project: {project_name}")
    doc.add_paragraph(
        f"POI Requirement: {stage1['poi_power_req_mw']:.2f} MW / "
        f"{stage1['poi_energy_req_mwh']:.2f} MWh"
    )
    doc.add_paragraph("")

    _append_dc_report_sections(doc, dc_output, ctx, chapter_prefix="1")
    return _doc_to_bytes(doc)


def create_ac_report(ac_output: dict, ctx: dict) -> bytes:
    doc = Document()
    _setup_margins(doc)
    _setup_header(doc)

    project_name = ac_output.get("project_name") or (ctx or {}).get("project_name") or "CALB ESS Project"
    _add_cover_page(doc, "CALB Utility-Scale ESS AC Sizing Report", project_name, ctx or {})

    _append_ac_report_sections(doc, ac_output, ctx or {}, chapter_prefix="1")
    _add_appendix(doc, ctx or {})

    return _doc_to_bytes(doc)


def create_combined_report(dc_output: dict, ac_output: dict, ctx: dict) -> bytes:
    doc = Document()
    _setup_margins(doc)
    _setup_header(doc)

    project_name = ac_output.get("project_name") or (ctx or {}).get("project_name") or "CALB ESS Project"
    _add_cover_page(doc, "CALB Utility-Scale ESS Combined Sizing Report", project_name, ctx or {})

    doc.add_heading("1. Executive Summary", level=2)
    dc_block_count = dc_output.get("dc_block_total_qty")
    if dc_block_count is None:
        dc_block_count = dc_output.get("container_count")
    combined_rows = [
        ("POI Power (MW)", _format_float(ac_output.get("poi_power_mw"), 2)),
    ]
    if ac_output.get("poi_energy_mwh") is not None:
        combined_rows.append(("POI Energy (MWh)", _format_float(ac_output.get("poi_energy_mwh"), 2)))
    combined_rows.extend(
        [
            ("DC Blocks", _format_float(dc_block_count, 0)),
            ("AC Blocks", _format_float(ac_output.get("num_blocks"), 0)),
            ("Total PCS Modules", _format_float(ac_output.get("total_pcs"), 0)),
            ("Transformer Count", _format_float(ac_output.get("transformer_count"), 0)),
            ("MV Level (kV)", _format_float(ac_output.get("grid_kv"), 1)),
        ]
    )
    _add_table(doc, combined_rows, ["Metric", "Value"])
    doc.add_paragraph("")

    doc.add_heading("2. Project Inputs & Assumptions", level=2)
    input_rows = []
    if ctx and ctx.get("inputs"):
        for key, value in ctx.get("inputs").items():
            input_rows.append((key, value))
    if not input_rows:
        input_rows = _build_ac_inputs_rows(ac_output, ctx)
    _add_table(doc, input_rows, ["Parameter", "Value"])
    doc.add_paragraph("")

    doc.add_heading("3. DC Sizing Results", level=2)
    _append_dc_report_sections(doc, dc_output or {}, ctx or {}, chapter_prefix="3")
    doc.add_paragraph("")

    doc.add_heading("4. AC Sizing Results", level=2)
    _append_ac_report_sections(doc, ac_output, ctx or {}, chapter_prefix="4")
    doc.add_paragraph("")

    doc.add_heading("5. Combined Configuration Summary", level=2)
    dc_spec = extract_dc_equipment_spec(
        block_code=dc_output.get("block_code") if dc_output else None,
        block_name=dc_output.get("block_name") if dc_output else None,
    )

    dc_total_qty = dc_output.get("dc_block_total_qty") if dc_output else None
    if dc_total_qty is None:
        dc_total_qty = dc_output.get("container_count") if dc_output else None
    dc_total_qty = int(dc_total_qty) if dc_total_qty is not None else None

    dc_unit = dc_spec.get("unit_capacity_mwh")
    dc_total_energy = None
    if dc_total_qty is not None and dc_unit is not None:
        dc_total_energy = dc_total_qty * dc_unit

    dc_voltage = dc_spec.get("voltage_range_v")
    if not dc_voltage and dc_spec.get("nominal_voltage_v") is not None:
        dc_voltage = f"{dc_spec.get('nominal_voltage_v'):.0f} V DC"

    combined_config_rows = [
        ("DC Container Model", dc_spec.get("container_model", ""), ""),
        ("DC Configuration", dc_spec.get("configuration", ""), ""),
        ("DC Unit Capacity (MWh)", _format_float(dc_unit, 3), ""),
        ("DC Total Quantity", dc_total_qty if dc_total_qty is not None else "", ""),
        ("DC Total Energy (MWh)", _format_float(dc_total_energy, 3), ""),
        ("DC Voltage", dc_voltage or "", ""),
        ("AC Block Size (MW)", "", _format_float(ac_output.get("block_size_mw"), 3)),
        ("AC Blocks", "", _format_float(ac_output.get("num_blocks"), 0)),
        ("Total AC Power (MW)", "", _format_float(ac_output.get("total_ac_mw"), 2)),
        ("Total PCS Modules", "", _format_float(ac_output.get("total_pcs"), 0)),
        ("Transformer Rating (kVA)", "", _format_float(ac_output.get("transformer_kva"), 0)),
        ("MV Level (kV)", "", _format_float(ac_output.get("grid_kv"), 1)),
    ]
    _add_table(doc, combined_config_rows, ["Metric", "DC", "AC"])
    doc.add_paragraph("")

    figure_index = 1

    doc.add_heading("6. Single Line Diagram", level=2)
    sld_png = _resolve_diagram_bytes(ctx or {}, "sld_png_bytes", "sld_svg_bytes")
    if sld_png:
        doc.add_picture(io.BytesIO(sld_png), width=Inches(6.7))
        doc.add_paragraph(f"Figure {figure_index} - Single Line Diagram (auto-generated)")
        figure_index += 1
    else:
        doc.add_paragraph("SLD not generated. Please generate in Single Line Diagram page.")
    doc.add_paragraph("")

    doc.add_heading("7. Site Layout", level=2)
    layout_png = _resolve_diagram_bytes(ctx or {}, "layout_png_bytes", "layout_svg_bytes")
    if layout_png:
        doc.add_picture(io.BytesIO(layout_png), width=Inches(6.7))
        doc.add_paragraph(f"Figure {figure_index} - Block Layout (auto-generated)")
        figure_index += 1
    else:
        doc.add_paragraph("Layout not generated. Please generate in Site Layout page.")
    doc.add_paragraph("")

    _add_appendix(doc, ctx or {})
    return _doc_to_bytes(doc)
