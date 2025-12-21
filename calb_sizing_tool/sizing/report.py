from docx import Document

def build_quality_report(results: dict, filepath: str):
    doc = Document()
    doc.add_heading("ESS Sizing Report", 0)
    for key, value in results.items():
        doc.add_paragraph(f"{key}: {value}")
    doc.save(filepath)
