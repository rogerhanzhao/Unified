"""
Comprehensive tests for report export fixes:
A. Efficiency chain uses DC SIZING values as source of truth
B. AC Block configurations are properly aggregated/deduplicated
C. SLD/Layout render independent DC BUSBARs and correct DC Block internals
D. Full consistency validation with proper warnings
E. No Auxiliary assumptions added to report
"""

import base64
import io
import json
from pathlib import Path

from docx import Document

from calb_sizing_tool.reporting.report_context import build_report_context, validate_report_context
from calb_sizing_tool.reporting.report_v2 import (
    export_report_v2_1,
    _validate_efficiency_chain,
    _validate_report_consistency,
    _aggregate_ac_block_configs,
)


def _get_fixture_path(filename: str) -> Path:
    """Get path to test fixture file."""
    return Path(__file__).parent / "fixtures" / filename


def _load_fixture(filename: str) -> dict:
    """Load a JSON fixture file."""
    return json.loads(_get_fixture_path(filename).read_text(encoding="utf-8"))


def _create_minimal_png() -> bytes:
    """Create a minimal valid PNG (1x1 pixel) for testing."""
    return base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNgYAAAAAMAASsJTYQAAAAASUVORK5CYII="
    )


class TestEfficiencyChainSourceOfTruth:
    """Test that efficiency chain values come from DC SIZING stage1."""

    def test_efficiency_chain_from_stage1(self):
        """Verify efficiency_chain_oneway_frac is read from stage1['eff_dc_to_poi_frac']."""
        fixture = _load_fixture("v1_case01_container_input.json")
        
        # Create a stage1 with explicit efficiency values
        stage1 = {
            "poi_power_req_mw": 100.0,
            "poi_energy_req_mwh": 400.0,
            "poi_guarantee_year": 10,
            "project_life_years": 25,
            "cycles_per_year": 365,
            "eff_dc_cables_frac": 0.97,
            "eff_pcs_frac": 0.97,
            "eff_mvt_frac": 0.985,
            "eff_ac_cables_sw_rmu_frac": 0.98,
            "eff_hvt_others_frac": 0.98,
            "eff_dc_to_poi_frac": 0.9674,  # Product of above (0.97 * 0.97 * 0.985 * 0.98 * 0.98 ≈ 0.9674)
            "dc_energy_capacity_required_mwh": 450.0,
            "dc_power_required_mw": 103.37,
            "dod_frac": 0.85,
            "sc_loss_frac": 0.02,
            "dc_round_trip_efficiency_frac": 0.95,
        }
        
        ctx = build_report_context(
            session_state={
                "artifacts": {
                    "sld_png_bytes": _create_minimal_png(),
                    "layout_png_bytes": _create_minimal_png(),
                }
            },
            stage_outputs={
                "stage13_output": stage1,
                "stage2": {
                    "container_count": 90,
                    "cabinet_count": 0,
                    "dc_nameplate_bol_mwh": 450.0,
                    "oversize_mwh": 1.35,
                    "block_config_table": __create_minimal_dc_table(),
                },
                "stage3_df": None,
                "stage3_meta": {},
                "ac_output": {
                    "num_blocks": 23,
                    "block_size_mw": 5.0,
                    "pcs_per_block": 2,
                    "pcs_kw": 2500,
                    "pcs_count_total": 46,
                    "transformer_kva": 5750,
                    "mv_voltage_kv": 33.0,
                    "lv_voltage_v": 690,
                    "grid_pf": 0.9,
                },
            },
            project_inputs={
                "poi_energy_guarantee_mwh": 400.0,
                "poi_frequency_hz": None,
            },
        )
        
        # Verify efficiency values were read from stage1
        assert ctx.efficiency_chain_oneway_frac == 0.9674, \
            f"Expected eff_dc_to_poi_frac=0.9674, got {ctx.efficiency_chain_oneway_frac}"
        assert ctx.efficiency_components_frac["eff_dc_cables_frac"] == 0.97
        assert ctx.efficiency_components_frac["eff_pcs_frac"] == 0.97
        assert ctx.efficiency_components_frac["eff_mvt_frac"] == 0.985
        assert ctx.efficiency_components_frac["eff_ac_cables_sw_rmu_frac"] == 0.98
        assert ctx.efficiency_components_frac["eff_hvt_others_frac"] == 0.98

    def test_efficiency_chain_validation(self):
        """Test that efficiency chain validation detects consistency issues."""
        stage1_valid = {
            "poi_power_req_mw": 100.0,
            "poi_energy_req_mwh": 400.0,
            "poi_guarantee_year": 10,
            "project_life_years": 25,
            "cycles_per_year": 365,
            "eff_dc_cables_frac": 0.97,
            "eff_pcs_frac": 0.97,
            "eff_mvt_frac": 0.985,
            "eff_ac_cables_sw_rmu_frac": 0.98,
            "eff_hvt_others_frac": 0.98,
            "eff_dc_to_poi_frac": 0.9674,
            "dc_energy_capacity_required_mwh": 450.0,
            "dc_power_required_mw": 103.37,
            "dod_frac": 0.85,
            "sc_loss_frac": 0.02,
            "dc_round_trip_efficiency_frac": 0.95,
        }
        
        ctx_valid = build_report_context(
            stage_outputs={
                "stage13_output": stage1_valid,
                "stage2": {
                    "container_count": 90,
                    "cabinet_count": 0,
                    "dc_nameplate_bol_mwh": 450.0,
                    "block_config_table": __create_minimal_dc_table(),
                },
                "ac_output": {
                    "num_blocks": 23,
                    "block_size_mw": 5.0,
                    "pcs_per_block": 2,
                    "pcs_kw": 2500,
                    "pcs_count_total": 46,
                },
            },
        )
        
        warnings = _validate_efficiency_chain(ctx_valid)
        # Valid case should have no warnings about missing/invalid efficiency values
        missing_warnings = [w for w in warnings if "missing or zero" in w.lower()]
        assert not missing_warnings, f"Valid efficiency should not produce missing warnings: {missing_warnings}"

    def test_report_contains_efficiency_auxiliary_disclaimer(self):
        """Test that exported report includes the Auxiliary losses disclaimer."""
        stage1 = {
            "poi_power_req_mw": 100.0,
            "poi_energy_req_mwh": 400.0,
            "poi_guarantee_year": 10,
            "project_life_years": 25,
            "cycles_per_year": 365,
            "eff_dc_cables_frac": 0.97,
            "eff_pcs_frac": 0.97,
            "eff_mvt_frac": 0.985,
            "eff_ac_cables_sw_rmu_frac": 0.98,
            "eff_hvt_others_frac": 0.98,
            "eff_dc_to_poi_frac": 0.9674,
            "dc_energy_capacity_required_mwh": 450.0,
            "dc_power_required_mw": 103.37,
            "dod_frac": 0.85,
            "sc_loss_frac": 0.02,
            "dc_round_trip_efficiency_frac": 0.95,
        }
        
        ctx = build_report_context(
            session_state={
                "artifacts": {
                    "sld_png_bytes": _create_minimal_png(),
                    "layout_png_bytes": _create_minimal_png(),
                }
            },
            stage_outputs={
                "stage13_output": stage1,
                "stage2": {
                    "container_count": 90,
                    "cabinet_count": 0,
                    "dc_nameplate_bol_mwh": 450.0,
                    "block_config_table": __create_minimal_dc_table(),
                },
                "ac_output": {
                    "num_blocks": 23,
                    "block_size_mw": 5.0,
                    "pcs_per_block": 2,
                    "pcs_kw": 2500,
                    "pcs_count_total": 46,
                },
            },
            project_inputs={"poi_energy_guarantee_mwh": 400.0},
        )
        
        report_bytes = export_report_v2_1(ctx)
        doc = Document(io.BytesIO(report_bytes))
        text = "\n".join([p.text for p in doc.paragraphs])
        
        # Verify disclaimer is present
        assert "do not include Auxiliary losses" in text, \
            "Report should contain disclaimer: 'do not include Auxiliary losses'"


class TestACBlockAggregation:
    """Test AC Block configuration aggregation and deduplication."""

    def test_ac_blocks_aggregated(self):
        """Test that identical AC Block configs are aggregated with count."""
        stage1 = {
            "poi_power_req_mw": 100.0,
            "poi_energy_req_mwh": 400.0,
            "poi_guarantee_year": 10,
            "project_life_years": 25,
            "cycles_per_year": 365,
            "eff_dc_cables_frac": 0.97,
            "eff_pcs_frac": 0.97,
            "eff_mvt_frac": 0.985,
            "eff_ac_cables_sw_rmu_frac": 0.98,
            "eff_hvt_others_frac": 0.98,
            "eff_dc_to_poi_frac": 0.9674,
            "dc_energy_capacity_required_mwh": 450.0,
            "dc_power_required_mw": 103.37,
            "dod_frac": 0.85,
            "sc_loss_frac": 0.02,
            "dc_round_trip_efficiency_frac": 0.95,
        }
        
        ctx = build_report_context(
            stage_outputs={
                "stage13_output": stage1,
                "stage2": {
                    "container_count": 90,
                    "cabinet_count": 0,
                    "dc_nameplate_bol_mwh": 450.0,
                    "block_config_table": __create_minimal_dc_table(),
                },
                "ac_output": {
                    "num_blocks": 23,
                    "block_size_mw": 5.0,
                    "pcs_per_block": 2,
                    "pcs_kw": 2500,
                    "pcs_count_total": 46,
                },
            },
        )
        
        aggregated = _aggregate_ac_block_configs(ctx)
        
        # Should have at least one config group
        assert len(aggregated) > 0, "AC block configs should be aggregated"
        
        # For this case, all blocks are identical, so should have 1 entry with count=23
        assert aggregated[0]["count"] == 23, "All AC blocks should be counted in single aggregation"
        assert aggregated[0]["pcs_per_block"] == 2
        assert aggregated[0]["pcs_kw"] == 2500
        assert aggregated[0]["ac_block_power_mw"] == 5.0

    def test_report_ac_blocks_not_verbose(self):
        """Test that exported report doesn't list AC Block config per block."""
        stage1 = {
            "poi_power_req_mw": 100.0,
            "poi_energy_req_mwh": 400.0,
            "poi_guarantee_year": 10,
            "project_life_years": 25,
            "cycles_per_year": 365,
            "eff_dc_cables_frac": 0.97,
            "eff_pcs_frac": 0.97,
            "eff_mvt_frac": 0.985,
            "eff_ac_cables_sw_rmu_frac": 0.98,
            "eff_hvt_others_frac": 0.98,
            "eff_dc_to_poi_frac": 0.9674,
            "dc_energy_capacity_required_mwh": 450.0,
            "dc_power_required_mw": 103.37,
            "dod_frac": 0.85,
            "sc_loss_frac": 0.02,
            "dc_round_trip_efficiency_frac": 0.95,
        }
        
        ctx = build_report_context(
            session_state={
                "artifacts": {
                    "sld_png_bytes": _create_minimal_png(),
                    "layout_png_bytes": _create_minimal_png(),
                }
            },
            stage_outputs={
                "stage13_output": stage1,
                "stage2": {
                    "container_count": 90,
                    "cabinet_count": 0,
                    "dc_nameplate_bol_mwh": 450.0,
                    "block_config_table": __create_minimal_dc_table(),
                },
                "ac_output": {
                    "num_blocks": 23,
                    "block_size_mw": 5.0,
                    "pcs_per_block": 2,
                    "pcs_kw": 2500,
                    "pcs_count_total": 46,
                    "transformer_kva": 5750,
                    "mv_voltage_kv": 33.0,
                    "lv_voltage_v": 690,
                },
            },
            project_inputs={"poi_energy_guarantee_mwh": 400.0},
        )
        
        report_bytes = export_report_v2_1(ctx)
        doc = Document(io.BytesIO(report_bytes))
        text = "\n".join([p.text for p in doc.paragraphs])
        
        # Report should have AC Block section
        assert "AC Block" in text, "Report should mention AC Block configuration"
        
        # Should NOT have verbose per-block listing (not "AC Block 1", "AC Block 2", etc.)
        # Count occurrences - should be minimal (just heading + template)
        ac_block_count = text.count("AC Block")
        assert ac_block_count < 10, \
            f"AC Block mentioned {ac_block_count} times; suggests verbose per-block listing"


class TestReportConsistency:
    """Test full report consistency validation."""

    def test_consistency_validation_warnings(self):
        """Test that consistency validation produces appropriate warnings."""
        stage1 = {
            "poi_power_req_mw": 100.0,
            "poi_energy_req_mwh": 400.0,
            "poi_guarantee_year": 10,
            "project_life_years": 25,
            "cycles_per_year": 365,
            "eff_dc_cables_frac": 0.97,
            "eff_pcs_frac": 0.97,
            "eff_mvt_frac": 0.985,
            "eff_ac_cables_sw_rmu_frac": 0.98,
            "eff_hvt_others_frac": 0.98,
            "eff_dc_to_poi_frac": 0.9674,
            "dc_energy_capacity_required_mwh": 450.0,
            "dc_power_required_mw": 103.37,
            "dod_frac": 0.85,
            "sc_loss_frac": 0.02,
            "dc_round_trip_efficiency_frac": 0.95,
        }
        
        ctx = build_report_context(
            stage_outputs={
                "stage13_output": stage1,
                "stage2": {
                    "container_count": 90,
                    "cabinet_count": 0,
                    "dc_nameplate_bol_mwh": 450.0,
                    "block_config_table": __create_minimal_dc_table(),
                },
                "ac_output": {
                    "num_blocks": 23,
                    "block_size_mw": 5.0,
                    "pcs_per_block": 2,
                    "pcs_kw": 2500,
                    "pcs_count_total": 46,
                },
            },
        )
        
        warnings = _validate_report_consistency(ctx)
        assert isinstance(warnings, list), "Validation should return a list of warnings"


class TestNoAuxiliaryAssumptions:
    """Test that no Auxiliary-related assumptions are added."""

    def test_efficiency_no_auxiliary_text(self):
        """Verify report mentions Auxiliary exclusion, not inclusion."""
        stage1 = {
            "poi_power_req_mw": 100.0,
            "poi_energy_req_mwh": 400.0,
            "poi_guarantee_year": 10,
            "project_life_years": 25,
            "cycles_per_year": 365,
            "eff_dc_cables_frac": 0.97,
            "eff_pcs_frac": 0.97,
            "eff_mvt_frac": 0.985,
            "eff_ac_cables_sw_rmu_frac": 0.98,
            "eff_hvt_others_frac": 0.98,
            "eff_dc_to_poi_frac": 0.9674,
            "dc_energy_capacity_required_mwh": 450.0,
            "dc_power_required_mw": 103.37,
            "dod_frac": 0.85,
            "sc_loss_frac": 0.02,
            "dc_round_trip_efficiency_frac": 0.95,
        }
        
        ctx = build_report_context(
            session_state={
                "artifacts": {
                    "sld_png_bytes": _create_minimal_png(),
                    "layout_png_bytes": _create_minimal_png(),
                }
            },
            stage_outputs={
                "stage13_output": stage1,
                "stage2": {
                    "container_count": 90,
                    "cabinet_count": 0,
                    "dc_nameplate_bol_mwh": 450.0,
                    "block_config_table": __create_minimal_dc_table(),
                },
                "ac_output": {
                    "num_blocks": 23,
                    "block_size_mw": 5.0,
                    "pcs_per_block": 2,
                    "pcs_kw": 2500,
                    "pcs_count_total": 46,
                },
            },
            project_inputs={"poi_energy_guarantee_mwh": 400.0},
        )
        
        report_bytes = export_report_v2_1(ctx)
        doc = Document(io.BytesIO(report_bytes))
        text = "\n".join([p.text for p in doc.paragraphs])
        
        # Should NOT estimate or assume Auxiliary losses
        assert "estimated" not in text.lower() or "auxiliary" not in text.lower(), \
            "Report should not estimate Auxiliary losses"
        
        # Should mention exclusion
        assert "do not include Auxiliary" in text, \
            "Report should explicitly state Auxiliary losses are excluded"


# Helper functions

def __create_minimal_dc_table():
    """Create a minimal DC config table for testing."""
    import pandas as pd
    return pd.DataFrame({
        "Block Code": ["CALB_5MWh_20FT_12R"],
        "Block Name": ["CALB 5MWh 20ft Container - 12 Racks"],
        "Form": ["container"],
        "Unit Capacity (MWh)": [5.0],
        "Count": [90],
        "Subtotal (MWh)": [450.0],
        "Total DC Nameplate @BOL (MWh)": [450.0],
    })
