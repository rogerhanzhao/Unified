import base64
import io
import json
from pathlib import Path

from docx import Document

from calb_sizing_tool.reporting.report_context import build_report_context
from calb_sizing_tool.reporting.report_v2 import export_report_v2_1, _validate_report_consistency
from tools.regress_export import run_ac_sizing, run_dc_sizing


def test_efficiency_chain_uses_dc_sizing_values():
    """Verify that exported efficiency chain uses actual DC SIZING values, not defaults."""
    fixture_path = Path(__file__).parent / "fixtures" / "v1_case01_container_input.json"
    fixture = json.loads(fixture_path.read_text(encoding="utf-8"))

    dc_results = run_dc_sizing(fixture)
    ac_output = run_ac_sizing(fixture, dc_results["stage1"], dc_results["stage2"])

    png_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNgYAAAAAMAASsJTYQAAAAASUVORK5CYII="
    )

    ctx = build_report_context(
        session_state={
            "artifacts": {
                "sld_png_bytes": png_bytes,
                "layout_png_bytes": png_bytes,
            }
        },
        stage_outputs={
            "stage13_output": dc_results["stage1"],
            "stage2": dc_results["stage2"],
            "stage3_df": dc_results["stage3_df"],
            "stage3_meta": dc_results["stage3_meta"],
            "ac_output": ac_output,
        },
        project_inputs={"poi_energy_guarantee_mwh": fixture["poi_energy_req_mwh"]},
        scenario_ids=fixture["scenario_id"],
    )

    # Verify efficiency values were read from DC SIZING, not defaults
    assert ctx.efficiency_chain_oneway_frac > 0, "Total efficiency should be positive"
    assert ctx.efficiency_chain_oneway_frac <= 1.0, "Total efficiency should not exceed 100%"
    
    # All component efficiencies should be present
    assert ctx.efficiency_components_frac.get("eff_dc_cables_frac") is not None
    assert ctx.efficiency_components_frac.get("eff_pcs_frac") is not None
    assert ctx.efficiency_components_frac.get("eff_mvt_frac") is not None
    assert ctx.efficiency_components_frac.get("eff_ac_cables_sw_rmu_frac") is not None
    assert ctx.efficiency_components_frac.get("eff_hvt_others_frac") is not None
    
    # Export and verify report contains actual values
    report_bytes = export_report_v2_1(ctx)
    doc = Document(io.BytesIO(report_bytes))
    texts = [p.text for p in doc.paragraphs]
    joined = "\n".join(texts)

    # Should NOT contain defaults (e.g., "97.00%" for PCS which is a common default)
    # Should contain the efficiency disclaimer about not including Auxiliary
    assert "do not include Auxiliary losses" in joined, "Report should state that efficiencies exclude Auxiliary"
    
    # Verify Efficiency Chain section exists
    assert "Efficiency Chain (one-way)" in joined


def test_ac_block_config_not_verbose():
    """Verify AC Block configuration doesn't list every block when they're identical."""
    fixture_path = Path(__file__).parent / "fixtures" / "v1_case01_container_input.json"
    fixture = json.loads(fixture_path.read_text(encoding="utf-8"))

    dc_results = run_dc_sizing(fixture)
    ac_output = run_ac_sizing(fixture, dc_results["stage1"], dc_results["stage2"])

    png_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNgYAAAAAMAASsJTYQAAAAASUVORK5CYII="
    )

    ctx = build_report_context(
        session_state={
            "artifacts": {
                "sld_png_bytes": png_bytes,
                "layout_png_bytes": png_bytes,
            }
        },
        stage_outputs={
            "stage13_output": dc_results["stage1"],
            "stage2": dc_results["stage2"],
            "stage3_df": dc_results["stage3_df"],
            "stage3_meta": dc_results["stage3_meta"],
            "ac_output": ac_output,
        },
        project_inputs={"poi_energy_guarantee_mwh": fixture["poi_energy_req_mwh"]},
        scenario_ids=fixture["scenario_id"],
    )

    # Export and verify report
    report_bytes = export_report_v2_1(ctx)
    doc = Document(io.BytesIO(report_bytes))
    texts = [p.text for p in doc.paragraphs]
    joined = "\n".join(texts)

    # Should have summary section
    assert "AC Block Configuration Summary" in joined or "AC:DC Ratio" in joined

    # Should NOT have verbose per-block listing (the old "AC Block 1", "AC Block 2", etc. lines)
    # Count how many times "AC Block" appears - should be minimal (just heading), not per-block
    ac_block_count = joined.count("AC Block")
    # Should be much less than actual AC blocks (e.g., 2-3 times for heading + maybe template ID)
    assert ac_block_count < 5, f"AC Block appears {ac_block_count} times, suggests verbose per-block listing"


def test_report_consistency_validation():
    """Verify consistency validation function works."""
    fixture_path = Path(__file__).parent / "fixtures" / "v1_case01_container_input.json"
    fixture = json.loads(fixture_path.read_text(encoding="utf-8"))

    dc_results = run_dc_sizing(fixture)
    ac_output = run_ac_sizing(fixture, dc_results["stage1"], dc_results["stage2"])

    png_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNgYAAAAAMAASsJTYQAAAAASUVORK5CYII="
    )

    ctx = build_report_context(
        session_state={
            "artifacts": {
                "sld_png_bytes": png_bytes,
                "layout_png_bytes": png_bytes,
            }
        },
        stage_outputs={
            "stage13_output": dc_results["stage1"],
            "stage2": dc_results["stage2"],
            "stage3_df": dc_results["stage3_df"],
            "stage3_meta": dc_results["stage3_meta"],
            "ac_output": ac_output,
        },
        project_inputs={"poi_energy_guarantee_mwh": fixture["poi_energy_req_mwh"]},
        scenario_ids=fixture["scenario_id"],
    )

    # Run validation
    warnings = _validate_report_consistency(ctx)
    
    # Valid context should have no critical errors
    # (May have warnings, but should not fail completely)
    assert isinstance(warnings, list)
