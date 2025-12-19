"""Lightweight DOCX report builder for sizing results."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Iterable, Mapping, Sequence, Tuple

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError as exc:  # pragma: no cover - environment dependency
    raise RuntimeError(
        "python-docx is required to build reports. Install it with `pip install python-docx`."
    ) from exc


KeyMetric = Tuple[str, str]


def _candidate_mappings(results: Mapping[str, Any]) -> Iterable[Mapping[str, Any]]:
    """Yield mappings that may hold metrics, from most to least specific."""
    if not isinstance(results, Mapping):
        return []

    for key in ("summary", "project", "stage1", "stage2", "stage3", "stage4"):
        candidate = results.get(key)
        if isinstance(candidate, Mapping):
            yield candidate

    yield results


def _read_first(mapping_candidates: Iterable[Mapping[str, Any]], key: str) -> Any:
    """Return the first non-None value for ``key`` across candidate mappings."""
    for mapping in mapping_candidates:
        if key in mapping and mapping[key] is not None:
            return mapping[key]
    return None


def _format_value(value: Any, fmt: str) -> str:
    """Format a metric value using ``fmt`` while handling missing/None gracefully."""
    if value is None:
        return "—"
    try:
        return fmt.format(value)
    except Exception:
        return str(value)


def _collect_key_metrics(results: Mapping[str, Any]) -> Sequence[KeyMetric]:
    """Extract a curated set of sizing metrics for the report."""
    candidates = list(_candidate_mappings(results))
    fields: Sequence[Tuple[str, str, str]] = (
        ("poi_power_req_mw", "POI Power Requirement", "{:.2f} MW"),
        ("poi_energy_req_mwh", "POI Energy Requirement", "{:.2f} MWh"),
        ("dc_power_required_mw", "DC Power (Equivalent)", "{:.2f} MW"),
        ("dc_energy_capacity_required_mwh", "DC Energy Required", "{:.2f} MWh"),
        ("dc_nameplate_bol_mwh", "DC Nameplate @ BOL", "{:.2f} MWh"),
        ("oversize_mwh", "Oversize vs Requirement", "{:.2f} MWh"),
        ("eff_dc_to_poi_frac", "DC→POI Efficiency (one-way)", "{:.2%}"),
        ("dc_round_trip_efficiency_frac", "DC Round-Trip Efficiency", "{:.2%}"),
        ("sc_loss_frac", "Self-commissioning Loss", "{:.2%}"),
        ("container_count", "Container Blocks", "{:,.0f}"),
        ("cabinet_count", "Cabinet Blocks", "{:,.0f}"),
        ("busbars_needed", "Busbars Needed", "{:,.0f}"),
    )

    metrics: list[KeyMetric] = []
    for key, label, fmt in fields:
        value = _read_first(candidates, key)
        if value is not None:
            metrics.append((label, _format_value(value, fmt)))
    return metrics


def _collect_ac_metrics(results: Mapping[str, Any]) -> Sequence[KeyMetric]:
    """Collect AC block sizing metrics when present."""
    ac_mapping = None
    for key in ("ac_block", "ac_config", "stage4", "stage4_step1_result"):
        candidate = results.get(key)
        if isinstance(candidate, Mapping):
            ac_mapping = candidate
            break
    if ac_mapping is None:
        return []

    fields: Sequence[Tuple[str, str, str]] = (
        ("strategy", "Strategy", "{}"),
        ("ac_block_qty", "AC Block Quantity", "{:,.0f}"),
        ("ac_block_rated_mw", "AC Block Rating", "{:.2f} MW"),
        ("pcs_units", "PCS Units per Block", "{:,.0f}"),
        ("pcs_unit_kw", "PCS Rating per Unit", "{:,.0f} kW"),
        ("dc_blocks_per_block", "DC Blocks per AC Block", "{:,.0f}"),
        ("container_per_block", "Containers per AC Block", "{:,.0f}"),
        ("cabinet_per_block", "Cabinets per AC Block", "{:,.0f}"),
        ("dc_blocks_per_block_base", "Base DC Blocks per Block", "{:,.0f}"),
        ("dc_blocks_per_block_max", "Max DC Blocks per Block", "{:,.0f}"),
        ("total_ac_mw", "Total AC Rating", "{:.2f} MW"),
        ("oversize_mw", "AC Oversize vs POI", "{:.2f} MW"),
    )

    metrics: list[KeyMetric] = []
    for key, label, fmt in fields:
        if key in ac_mapping:
            metrics.append((label, _format_value(ac_mapping.get(key), fmt)))
    return metrics


def _add_table(doc: Document, title: str, metrics: Sequence[KeyMetric]) -> None:
    """Render a two-column metric table with a heading."""
    if not metrics:
        return

    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1 + len(metrics), cols=2)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Metric"
    hdr_cells[1].text = "Value"

    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True

    for i, (label, value) in enumerate(metrics, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = label
        row_cells[1].text = value

    doc.add_paragraph()  # spacer


def _apply_base_styles(doc: Document) -> None:
    """Set mild defaults so the report looks clean without heavy styling."""
    section = doc.sections[0]
    section.left_margin = section.right_margin = Pt(54)  # ~0.75"
    section.top_margin = section.bottom_margin = Pt(54)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Segoe UI"
    normal_style.font.size = Pt(10.5)


def build_report_docx(results: Mapping[str, Any], outpath: str) -> Path:
    """Create a DOCX report summarizing sizing ``results``.

    The function is intentionally tolerant of partial dictionaries. It looks for
    commonly used sizing keys (e.g., ``poi_power_req_mw``, ``container_count``,
    ``ac_block_qty``) in the provided mapping and renders them into key metric
    tables. Additional sections are added when AC block metadata is available.

    Args:
        results: Mapping containing sizing metadata.
        outpath: Target path for the DOCX file. The parent directory is created
            automatically if needed. If ``outpath`` lacks a ``.docx`` suffix it
            will be appended.

    Returns:
        Path to the saved DOCX report.

    Raises:
        ValueError: If ``results`` is empty or ``outpath`` is falsy.
        TypeError: If ``results`` is not a mapping.
        IsADirectoryError: If ``outpath`` resolves to an existing directory.
        RuntimeError: If ``python-docx`` is not available in the environment.
    """
    if not results:
        raise ValueError("results must contain sizing data")
    if not isinstance(results, Mapping):
        raise TypeError("results must be a mapping of sizing data")
    if not outpath:
        raise ValueError("outpath is required")

    output_path = Path(outpath)
    if output_path.suffix.lower() != ".docx":
        output_path = output_path.with_suffix(".docx")

    if output_path.exists() and output_path.is_dir():
        raise IsADirectoryError(f"Cannot write report to directory: {output_path}")

    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _apply_base_styles(doc)

    project_candidates = list(_candidate_mappings(results))
    project_name = _read_first(project_candidates, "project_name") or "CALB ESS Project"

    title = doc.add_heading("Energy Storage Sizing Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f"Project: {project_name}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(11)
    doc.add_paragraph(
        f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M %Z')}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer

    key_metrics = _collect_key_metrics(results)
    _add_table(doc, "Key Metrics", key_metrics)

    ac_metrics = _collect_ac_metrics(results)
    _add_table(doc, "AC Block Summary", ac_metrics)

    scenarios = results.get("scenarios") or results.get("cases")
    if isinstance(scenarios, list) and scenarios:
        doc.add_heading("Scenarios", level=2)
        for scenario in scenarios:
            if not isinstance(scenario, Mapping):
                continue
            name = scenario.get("name") or scenario.get("label") or "Scenario"
            doc.add_heading(str(name), level=3)
            _add_table(doc, "Scenario Metrics", _collect_key_metrics(scenario))

    notes = results.get("notes")
    if notes:
        doc.add_heading("Notes", level=2)
        doc.add_paragraph(str(notes))

    temp_path = output_path.with_suffix(output_path.suffix + ".tmp")
    doc.save(temp_path)
    temp_path.replace(output_path)
    return output_path
