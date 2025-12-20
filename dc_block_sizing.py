# dc_block_sizing.py
import streamlit as st
import pandas as pd
import numpy as np
import math
import altair as alt
import os
import io
from pathlib import Path

from sizing.dc_logic import load_dc_data, calculate_dc_energy, dc_rte_calculation
from sizing.validation import validate_input_data
from stage4_interface import pack_stage13_output

# ==========================================
# 0. SETUP & LIBRARY CHECK
# ==========================================
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Optional: Matplotlib for DOCX chart export
try:
    import matplotlib.pyplot as plt
    MATPLOTLIB_AVAILABLE = True
except Exception:
    MATPLOTLIB_AVAILABLE = False

# Design rule: max 418kWh cabinets per DC busbar
K_MAX_FIXED = 10

# ==========================================
# CALB VI COLORS
# ==========================================
if not st.session_state.get("_dc_page_configured"):
    # app.py sets the global page config; keep fallback for direct module runs
    st.set_page_config(page_title="CALB ESS Sizing Tool – Stage 1–3 (DC)", layout="wide")
    st.session_state["_dc_page_configured"] = True
CALB_SKY_BLUE   = "#5cc3e4"
CALB_DEEP_BLUE  = "#23496b"
CALB_BLACK      = "#1e1e1e"
CALB_GREY       = "#58595b"
CALB_LIGHT_GREY = "#bebec3"
CALB_WISE_GREY  = "#cedbea"
CALB_WHITE      = "#ffffff"

# ------------------------------------------
# Theme-aware colors
# ------------------------------------------
base_theme = st.get_option("theme.base") or "light"
is_dark = (base_theme == "dark")

if is_dark:
    BG_MAIN      = "#0f1116"
    TITLE_COLOR  = CALB_SKY_BLUE
    CARD_BG      = "#1b1e24"
    CARD_BORDER  = "#2c2f36"
    METRIC_LABEL = "#9ca3af"
    METRIC_VALUE = CALB_SKY_BLUE
    BUTTON_BG    = CALB_SKY_BLUE
    BUTTON_FG    = CALB_BLACK
    BUTTON_HOVER_BG = CALB_WHITE
    BUTTON_HOVER_FG = CALB_DEEP_BLUE
else:
    BG_MAIN      = CALB_WISE_GREY
    TITLE_COLOR  = CALB_DEEP_BLUE
    CARD_BG      = CALB_WHITE
    CARD_BORDER  = CALB_LIGHT_GREY
    METRIC_LABEL = CALB_GREY
    METRIC_VALUE = CALB_DEEP_BLUE
    BUTTON_BG    = CALB_SKY_BLUE
    BUTTON_FG    = CALB_WHITE
    BUTTON_HOVER_BG = CALB_DEEP_BLUE
    BUTTON_HOVER_FG = CALB_WHITE

CHART_TEXT_COLOR = CALB_DEEP_BLUE if not is_dark else CALB_SKY_BLUE

# ------------------------------------------
# Global CSS
# ------------------------------------------
st.markdown(
    f"""
<style>
.main {{
    background-color: {BG_MAIN} !important;
}}
.block-container {{
    padding-top: 0.5rem;
    padding-bottom: 2rem;
}}
.calb-page-title {{
    font-family: "Segoe UI", "Helvetica Neue", Arial, sans-serif;
    color: {TITLE_COLOR};
    margin-bottom: 0.0rem;
}}
.calb-card {{
    background-color: {CARD_BG};
    padding: 1.2rem 1.6rem;
    border-radius: 18px;
    border: 1px solid {CARD_BORDER};
    box-shadow: 0 6px 18px rgba(0, 0, 0, 0.15);
    margin-bottom: 1.4rem;
}}
.metric-label {{
    color: {METRIC_LABEL} !important;
    font-size: 0.85rem;
    text-transform: uppercase;
    letter-spacing: 0.04em;
    font-family: "Segoe UI", "Helvetica Neue", Arial, sans-serif;
}}
.metric-value {{
    font-size: 1.3rem;
    font-weight: 600;
    color: {METRIC_VALUE} !important;
    font-family: "Segoe UI", "Helvetica Neue", Arial, sans-serif;
}}
.stButton>button {{
    background-color: {BUTTON_BG};
    color: {BUTTON_FG};
    border-radius: 999px;
    border: none;
    padding: 0.7rem 2.3rem;
    font-weight: 650;
    font-size: 1.0rem;
    font-family: "Segoe UI", "Helvetica Neue", Arial, sans-serif;
    box-shadow: 0 4px 10px rgba(0,0,0,0.25);
}}
.stButton>button:hover {{
    background-color: {BUTTON_HOVER_BG};
    color: {BUTTON_HOVER_FG};
}}
</style>
""",
    unsafe_allow_html=True,
)

# ==========================================
# 1. DATA FILE PATH (N1 ADAPTED)
# ==========================================
DEFAULT_FILENAME = "ess_sizing_data_dictionary_v13_dc_autofit.xlsx"
DEFAULT_LOGO_NAME = "calb_logo.png"
SCRIPT_DIR = Path(__file__).resolve().parent


def find_repo_root(start: Path) -> Path:
    """Return the repository root (looks for a ``.git`` folder) or ``start``."""

    for candidate in [start] + list(start.parents):
        if (candidate / ".git").is_dir():
            return candidate
    return start


REPO_ROOT = find_repo_root(SCRIPT_DIR)
DATA_DIR = REPO_ROOT / "data"
DATA_FILE = None


def resolve_data_file(default_filename: str) -> str | None:
    """
    Find the data dictionary Excel file.
    Priority:
      1) Path specified inside a nearby data_path.txt
      2) Default filename in common locations (cwd, script dir, repo root, repo_root/data)
      3) Any matching workbook discovered by glob search under the usual roots
    """
    candidates = []

    config_files = [
        Path("data_path.txt"),
        SCRIPT_DIR / "data_path.txt",
        REPO_ROOT / "data_path.txt",
    ]
    for cfg in config_files:
        if cfg.is_file():
            try:
                raw = cfg.read_text().strip()
            except Exception:
                continue
            if raw:
                raw_path = Path(raw).expanduser()
                if not raw_path.is_absolute():
                    raw_path = (cfg.parent / raw_path).resolve()
                candidates.append(raw_path)

    search_dirs = [
        Path(DATA_DIR),
        SCRIPT_DIR / "data",
        Path.cwd(),
        SCRIPT_DIR,
        REPO_ROOT,
    ]
    for directory in search_dirs:
        candidates.append(directory / default_filename)

    seen = set()
    for cand in candidates:
        cand = cand.expanduser()
        try:
            resolved = cand.resolve()
        except FileNotFoundError:
            resolved = cand.absolute()
        if resolved in seen:
            continue
        seen.add(resolved)
        if resolved.is_file():
            return str(resolved)

    fallback_roots = [
        Path(DATA_DIR),
        SCRIPT_DIR / "data",
        REPO_ROOT,
        SCRIPT_DIR,
    ]
    fallback_patterns = [
        default_filename,
        "ess_sizing_data_dictionary*dc*.xlsx",
        "ess_sizing_data_dictionary*.xlsx",
    ]
    for root in fallback_roots:
        if not root.exists():
            continue
        for pattern in fallback_patterns:
            for match in root.rglob(pattern):
                if match.is_file():
                    resolved = match.resolve()
                    if resolved not in seen:
                        return str(resolved)
    return None


DATA_FILE = resolve_data_file(DEFAULT_FILENAME)

if not DATA_FILE:
    search_locations = [
        str(Path(DATA_DIR)),
        str(SCRIPT_DIR / "data"),
        str(Path.cwd()),
        str(SCRIPT_DIR),
        str(REPO_ROOT),
    ]
    pattern_hint = "ess_sizing_data_dictionary*dc*.xlsx"
    st.error(
        "❌ Data file "
        f"'{DEFAULT_FILENAME}' not found. Looked in:\n- " + "\n- ".join(search_locations)
        + f"\nAlso searched for files matching '{pattern_hint}' in those locations."
        + "\nYou can also provide a custom path inside a 'data_path.txt' file."
    )
    st.stop()

# ==========================================
# 2. HELPERS
# ==========================================
def to_float(x, default=0.0):
    try:
        if isinstance(x, str):
            x = x.replace("%", "").replace(",", "").strip()
        return float(x)
    except Exception:
        return default

def to_frac(x, default=1.0):
    v = to_float(x, default)
    if v > 1.5:
        return v / 100.0
    return v

def safe_div(a: float, b: float, default: float = 0.0) -> float:
    try:
        if b is None or abs(float(b)) < 1e-12:
            return default
        return float(a) / float(b)
    except Exception:
        return default

def calc_sc_loss_pct(sc_months: float) -> float:
    m = int(round(to_float(sc_months, 0.0)))
    if m <= 0:
        return 0.0

    mapping = {
        1: 2.0, 2: 2.0, 3: 2.0,
        4: 2.5, 5: 2.8, 6: 3.0,
        7: 3.2, 8: 3.5, 9: 3.8,
        10: 4.1, 11: 4.3, 12: 4.5,
    }

    if m in mapping:
        return mapping[m]

    if m > 12:
        base_loss = 4.5
        extra_months = sc_months - 12.0
        return base_loss + (extra_months * 0.05)

    lower = int(math.floor(sc_months))
    upper = int(math.ceil(sc_months))
    if lower <= 0:
        return mapping.get(upper, 2.0)
    if upper == lower:
        return mapping.get(lower, 2.0)
    v_low = mapping.get(lower, 2.0)
    v_up = mapping.get(upper, v_low)
    ratio = (sc_months - lower) / (upper - lower)
    return v_low + (v_up - v_low) * ratio

def st_dataframe_full_width(df: pd.DataFrame):
    try:
        st.dataframe(df, width="stretch")
    except TypeError:
        st.dataframe(df, use_container_width=True)

def first_success_key(results: dict, preferred_order: list):
    for k in preferred_order:
        v = results.get(k)
        if isinstance(v, tuple) and len(v) > 0 and v[0] != "ERROR":
            return k
    return None

# ==========================================
# 3. LOAD DATA FROM EXCEL
# ==========================================
@st.cache_data
def load_data(path: str):
    data_path = Path(path).expanduser()
    if not data_path.is_file():
        raise FileNotFoundError(f"Data file not found: {data_path}")

    try:
        xls = pd.ExcelFile(data_path)
    except Exception as exc:
        raise ValueError(f"Unable to open Excel file '{data_path}': {exc}") from exc

    required_sheets = {
        "ess_sizing_case",
        "dc_block_template_314_data",
        "soh_profile_314_data",
        "soh_curve_314_template",
        "rte_profile_314_data",
        "rte_curve_314_template",
    }
    missing_sheets = [s for s in required_sheets if s not in xls.sheet_names]
    if missing_sheets:
        raise ValueError(
            f"Excel file '{data_path.name}' is missing required sheets: {', '.join(sorted(missing_sheets))}"
        )

    def ensure_columns(df: pd.DataFrame, sheet_name: str, required_cols: list[str]):
        missing = [c for c in required_cols if c not in df.columns]
        if missing:
            raise ValueError(
                f"Sheet '{sheet_name}' in '{data_path.name}' is missing required columns: {', '.join(missing)}"
            )

    df_case = pd.read_excel(xls, "ess_sizing_case")
    ensure_columns(df_case, "ess_sizing_case", ["Field Name", "Default Value"])
    defaults = {}
    for _, row in df_case.iterrows():
        field_name = row.get("Field Name")
        if isinstance(field_name, str) and field_name.strip():
            defaults[field_name.strip()] = row.get("Default Value")

    df_blocks = pd.read_excel(xls, "dc_block_template_314_data")
    ensure_columns(
        df_blocks,
        "dc_block_template_314_data",
        ["Block_Form", "Is_Active", "Is_Default_Option", "Block_Nameplate_Capacity_Mwh", "Dc_Block_Code", "Dc_Block_Name"],
    )

    df_soh_profile = pd.read_excel(xls, "soh_profile_314_data")
    ensure_columns(df_soh_profile, "soh_profile_314_data", ["Profile_Id", "C_Rate", "Cycles_Per_Year"])

    df_soh_curve = pd.read_excel(xls, "soh_curve_314_template")
    ensure_columns(df_soh_curve, "soh_curve_314_template", ["Profile_Id", "Life_Year_Index", "Soh_Dc_Pct"])

    df_rte_profile = pd.read_excel(xls, "rte_profile_314_data")
    ensure_columns(df_rte_profile, "rte_profile_314_data", ["Profile_Id", "C_Rate"])

    df_rte_curve = pd.read_excel(xls, "rte_curve_314_template")
    ensure_columns(df_rte_curve, "rte_curve_314_template", ["Profile_Id", "Soh_Band_Min_Pct", "Rte_Dc_Pct"])

    return defaults, df_blocks, df_soh_profile, df_soh_curve, df_rte_profile, df_rte_curve

# ==========================================
# 4. CORE CALC LOGIC
# ==========================================
def run_stage1(inputs: dict, defaults: dict) -> dict:
    def get(name, fallback=None):
        if name in inputs and inputs[name] is not None:
            return inputs[name]
        if name in defaults:
            return defaults[name]
        return fallback

    project_name = str(get("project_name", "CALB ESS Project"))

    poi_mw = to_float(get("poi_power_req_mw", 100.0))
    poi_mwh = to_float(get("poi_energy_req_mwh", 400.0))
    project_life_years = int(to_float(get("project_life_years", 20)))
    cycles_per_year = int(to_float(get("cycles_per_year", 365)))
    poi_guarantee_year = int(to_float(get("poi_guarantee_year", 0)))

    eff_dc_cables = to_frac(get("eff_dc_cables", 0.995))
    eff_pcs       = to_frac(get("eff_pcs", 0.985))
    eff_mvt       = to_frac(get("eff_mvt", 0.995))
    eff_ac_sw     = to_frac(get("eff_ac_cables_sw_rmu", 0.992))
    eff_hvt       = to_frac(get("eff_hvt_others", 1.0))
    eff_chain = eff_dc_cables * eff_pcs * eff_mvt * eff_ac_sw * eff_hvt

    sc_val = to_float(get("sc_time_months", 3.0))
    if sc_val < 3.0:
        sc_val = 3.0
    sc_time_months = sc_val
    sc_loss_pct = calc_sc_loss_pct(sc_time_months)
    sc_loss_frac = sc_loss_pct / 100.0

    dod_frac    = to_frac(get("dod_pct", 97.0))
    dc_rte_frac = to_frac(get("dc_round_trip_efficiency_pct", 94.0))
    dc_one_way_eff = math.sqrt(dc_rte_frac) if dc_rte_frac >= 0 else 0.0
    dc_usable_bol_frac = dod_frac * dc_one_way_eff

    denom = (1.0 - sc_loss_frac) * dc_usable_bol_frac * eff_chain
    dc_energy_required = safe_div(poi_mwh, denom, default=0.0)

    dc_power_required_mw = safe_div(poi_mw, eff_chain, default=0.0) if eff_chain > 0 else 0.0

    return {
        "project_name": project_name,
        "poi_power_req_mw": poi_mw,
        "poi_energy_req_mwh": poi_mwh,
        "project_life_years": project_life_years,
        "cycles_per_year": cycles_per_year,
        "poi_guarantee_year": poi_guarantee_year,
        "eff_dc_cables_frac": eff_dc_cables,
        "eff_pcs_frac": eff_pcs,
        "eff_mvt_frac": eff_mvt,
        "eff_ac_cables_sw_rmu_frac": eff_ac_sw,
        "eff_hvt_others_frac": eff_hvt,
        "eff_dc_to_poi_frac": eff_chain,
        "sc_time_months": sc_time_months,
        "sc_loss_pct": sc_loss_pct,
        "sc_loss_frac": sc_loss_frac,
        "dod_frac": dod_frac,
        "dc_round_trip_efficiency_frac": dc_rte_frac,
        "dc_one_way_efficiency_frac": dc_one_way_eff,
        "dc_usable_bol_frac": dc_usable_bol_frac,
        "dc_energy_capacity_required_mwh": dc_energy_required,
        "dc_power_required_mw": dc_power_required_mw,
    }

def _pick_dc_block(df_blocks: pd.DataFrame, form: str):
    df = df_blocks.copy()
    df["Block_Form_L"] = df["Block_Form"].astype(str).str.lower()
    cand = df[(df["Block_Form_L"] == form.lower()) & (df["Is_Active"] == 1)]
    if cand.empty:
        return None

    pref = cand[cand["Is_Default_Option"] == 1]
    if pref.empty:
        pref = cand

    pref = pref.sort_values("Block_Nameplate_Capacity_Mwh", ascending=False)
    row = pref.iloc[0]
    return str(row["Dc_Block_Code"]), str(row["Dc_Block_Name"]), float(row["Block_Nameplate_Capacity_Mwh"])

def _make_config_table(rows: list):
    df = pd.DataFrame(rows)
    if df.empty:
        return df

    df["Subtotal (MWh)"] = df["Unit Capacity (MWh)"] * df["Count"]
    total = float(df["Subtotal (MWh)"].sum())
    df["Total DC Nameplate @BOL (MWh)"] = total
    return df

def build_config_container_only(required_dc_mwh: float, container_unit: float, container_code: str, container_name: str):
    cnt = int(math.ceil(required_dc_mwh / container_unit)) if required_dc_mwh > 0 else 0
    rows = [{
        "Block Code": container_code,
        "Block Name": container_name,
        "Form": "container",
        "Unit Capacity (MWh)": float(container_unit),
        "Count": int(cnt),
    }]
    df = _make_config_table(rows)
    total = float(df["Subtotal (MWh)"].sum()) if not df.empty else 0.0
    oversize = total - required_dc_mwh
    return {
        "mode": "container_only",
        "dc_nameplate_bol_mwh": total,
        "oversize_mwh": oversize,
        "config_adjustment_frac": (total / required_dc_mwh - 1.0) if required_dc_mwh > 0 else 0.0,
        "block_config_table": df,
        "container_count": cnt,
        "cabinet_count": 0,
        "busbars_needed": 0,
    }

def build_config_cabinet_only(required_dc_mwh: float, cab_unit: float, cab_code: str, cab_name: str, k_max: int = K_MAX_FIXED):
    cab = int(math.ceil(required_dc_mwh / cab_unit)) if required_dc_mwh > 0 else 0
    busbars = int(math.ceil(cab / k_max)) if cab > 0 else 0
    rows = [{
        "Block Code": cab_code,
        "Block Name": cab_name,
        "Form": "cabinet",
        "Unit Capacity (MWh)": float(cab_unit),
        "Count": int(cab),
    }]
    df = _make_config_table(rows)
    total = float(df["Subtotal (MWh)"].sum()) if not df.empty else 0.0
    oversize = total - required_dc_mwh
    return {
        "mode": "cabinet_only",
        "dc_nameplate_bol_mwh": total,
        "oversize_mwh": oversize,
        "config_adjustment_frac": (total / required_dc_mwh - 1.0) if required_dc_mwh > 0 else 0.0,
        "block_config_table": df,
        "container_count": 0,
        "cabinet_count": cab,
        "busbars_needed": busbars,
    }

def build_config_hybrid(required_dc_mwh: float,
                        container_unit: float, container_code: str, container_name: str,
                        cab_unit: float, cab_code: str, cab_name: str,
                        k_max: int = K_MAX_FIXED):
    if required_dc_mwh <= 0:
        cont = 0
        cab = 0
    else:
        cont = int(math.floor(required_dc_mwh / container_unit))
        remainder = required_dc_mwh - cont * container_unit
        if remainder <= 1e-9:
            cab = 0
        else:
            cab = int(math.ceil(remainder / cab_unit))
            if cab > k_max:
                cont += 1
                cab = 0

    rows = []
    if cont > 0:
        rows.append({
            "Block Code": container_code,
            "Block Name": container_name,
            "Form": "container",
            "Unit Capacity (MWh)": float(container_unit),
            "Count": int(cont),
        })
    if cab > 0:
        rows.append({
            "Block Code": cab_code,
            "Block Name": cab_name,
            "Form": "cabinet",
            "Unit Capacity (MWh)": float(cab_unit),
            "Count": int(cab),
        })

    df = _make_config_table(rows)
    total = float(df["Subtotal (MWh)"].sum()) if not df.empty else 0.0
    oversize = total - required_dc_mwh
    busbars = 1 if cab > 0 else 0

    return {
        "mode": "hybrid",
        "dc_nameplate_bol_mwh": total,
        "oversize_mwh": oversize,
        "config_adjustment_frac": (total / required_dc_mwh - 1.0) if required_dc_mwh > 0 else 0.0,
        "block_config_table": df,
        "container_count": cont,
        "cabinet_count": cab,
        "busbars_needed": busbars,
    }

def select_soh_profile(effective_c_rate: float, cycles_per_year: int, df_soh_profile: pd.DataFrame):
    df = df_soh_profile.copy()
    df["c_rate_diff"] = (df["C_Rate"] - effective_c_rate).abs()
    df["cycles_diff"] = (df["Cycles_Per_Year"] - cycles_per_year).abs()
    df["score"] = df["c_rate_diff"] * 10.0 + df["cycles_diff"] / 365.0
    best = df.sort_values("score").iloc[0]
    return int(best["Profile_Id"]), float(best["C_Rate"]), int(best["Cycles_Per_Year"])

def select_rte_profile(effective_c_rate: float, df_rte_profile: pd.DataFrame):
    df = df_rte_profile.copy()
    df["c_rate_diff"] = (df["C_Rate"] - effective_c_rate).abs()
    best = df.sort_values("c_rate_diff").iloc[0]
    return int(best["Profile_Id"]), float(best["C_Rate"])

def run_stage3(stage1: dict,
               stage2: dict,
               df_soh_profile: pd.DataFrame,
               df_soh_curve: pd.DataFrame,
               df_rte_profile: pd.DataFrame,
               df_rte_curve: pd.DataFrame):

    dc_nameplate_bol_mwh = stage2["dc_nameplate_bol_mwh"]
    if dc_nameplate_bol_mwh <= 0:
        raise ValueError("DC nameplate @BOL must be > 0 for Stage 3.")

    poi_mw = stage1["poi_power_req_mw"]
    poi_energy_mwh = stage1["poi_energy_req_mwh"]
    project_life_years = int(stage1["project_life_years"])
    cycles_per_year = int(stage1["cycles_per_year"])
    sc_loss_frac = stage1["sc_loss_frac"]
    dod_frac = stage1["dod_frac"]
    eff_chain = stage1["eff_dc_to_poi_frac"]
    guarantee_year = int(stage1.get("poi_guarantee_year", 0))

    dc_power_mw = safe_div(poi_mw, eff_chain, default=0.0) if eff_chain > 0 else 0.0
    effective_c_rate = safe_div(dc_power_mw, dc_nameplate_bol_mwh, default=0.0)

    soh_profile_id, chosen_soh_c_rate, chosen_cycles_per_year = select_soh_profile(
        effective_c_rate, cycles_per_year, df_soh_profile
    )
    rte_profile_id, chosen_rte_c_rate = select_rte_profile(effective_c_rate, df_rte_profile)

    soh_curve_sel = df_soh_curve[df_soh_curve["Profile_Id"] == soh_profile_id].copy()
    if "Soh_Dc_Pct" in soh_curve_sel.columns:
        soh_curve_sel["Soh_Dc_Pct"] = soh_curve_sel["Soh_Dc_Pct"].apply(lambda x: to_frac(x))
    soh_curve_sel = soh_curve_sel.sort_values("Life_Year_Index")

    rte_curve_sel = df_rte_curve[df_rte_curve["Profile_Id"] == rte_profile_id].copy()
    if "Soh_Band_Min_Pct" in rte_curve_sel.columns:
        rte_curve_sel["Soh_Band_Min_Pct"] = rte_curve_sel["Soh_Band_Min_Pct"].apply(lambda x: to_frac(x))
    if "Rte_Dc_Pct" in rte_curve_sel.columns:
        rte_curve_sel["Rte_Dc_Pct"] = rte_curve_sel["Rte_Dc_Pct"].apply(lambda x: to_frac(x))
    rte_curve_sel = rte_curve_sel.sort_values("Soh_Band_Min_Pct", ascending=False)

    records = []

    for y in range(0, project_life_years + 1):
        row = soh_curve_sel[soh_curve_sel["Life_Year_Index"] == y]
        if not row.empty:
            soh_rel = float(row["Soh_Dc_Pct"].iloc[0])
        else:
            soh_rel = float(soh_curve_sel["Soh_Dc_Pct"].iloc[-1])

        soh_abs = soh_rel * (1.0 - sc_loss_frac)

        rte_row = rte_curve_sel[rte_curve_sel["Soh_Band_Min_Pct"] <= soh_abs].head(1)
        if rte_row.empty:
            rte_row = rte_curve_sel.tail(1)

        raw_rte = float(rte_row["Rte_Dc_Pct"].iloc[0])
        dc_rte_frac_year = min(1.0, max(0.0, raw_rte))
        dc_one_way_eff_year = math.sqrt(dc_rte_frac_year)

        dc_gross_capacity_mwh_year = dc_nameplate_bol_mwh * soh_abs
        dc_usable_mwh_year = dc_gross_capacity_mwh_year * dod_frac * dc_one_way_eff_year
        poi_usable_mwh_year = max(dc_usable_mwh_year * eff_chain, 0.0)

        system_rte_frac_year = min(1.0, max(0.0, dc_rte_frac_year * (eff_chain ** 2)))
        meets_poi_energy = poi_usable_mwh_year >= poi_energy_mwh

        records.append(
            {
                "Year_Index": int(y),
                "SOH_Relative": soh_rel,
                "SOH_Absolute": soh_abs,
                "DC_Nameplate_BOL_MWh": dc_nameplate_bol_mwh,
                "DC_Gross_Capacity_MWh": dc_gross_capacity_mwh_year,
                "DC_Usable_MWh": dc_usable_mwh_year,
                "DC_RTE_Frac": dc_rte_frac_year,
                "System_RTE_Frac": system_rte_frac_year,
                "POI_Usable_Energy_MWh": poi_usable_mwh_year,
                "Meets_POI_Req": meets_poi_energy,
                "Is_Guarantee_Year": (y == guarantee_year),
            }
        )

    df_years = pd.DataFrame(records)
    df_years["SOH_Display_Pct"] = df_years["SOH_Relative"] * 100.0
    df_years["SOH_Absolute_Pct"] = df_years["SOH_Absolute"] * 100.0
    df_years["DC_RTE_Pct"] = df_years["DC_RTE_Frac"] * 100.0
    df_years["System_RTE_Pct"] = df_years["System_RTE_Frac"] * 100.0

    meta = {
        "poi_power_mw": poi_mw,
        "dc_power_mw": dc_power_mw,
        "effective_c_rate": effective_c_rate,
        "soh_profile_id": soh_profile_id,
        "rte_profile_id": rte_profile_id,
        "chosen_soh_c_rate": chosen_soh_c_rate,
        "chosen_soh_cycles_per_year": chosen_cycles_per_year,
        "chosen_rte_c_rate": chosen_rte_c_rate,
    }
    return df_years, meta

def size_with_guarantee(stage1: dict,
                        mode: str,
                        df_blocks: pd.DataFrame,
                        df_soh_profile: pd.DataFrame,
                        df_soh_curve: pd.DataFrame,
                        df_rte_profile: pd.DataFrame,
                        df_rte_curve: pd.DataFrame,
                        k_max: int = K_MAX_FIXED,
                        max_iter: int = 60):

    required_dc = float(stage1["dc_energy_capacity_required_mwh"])
    poi_energy_req = float(stage1["poi_energy_req_mwh"])
    guarantee_year = int(stage1.get("poi_guarantee_year", 0))
    project_life_years = int(stage1["project_life_years"])
    guarantee_year = max(0, min(guarantee_year, project_life_years))

    picked_container = _pick_dc_block(df_blocks, "container")
    picked_cabinet   = _pick_dc_block(df_blocks, "cabinet")

    if picked_container is None:
        raise ValueError("No active 'container' DC block template found in Excel.")
    if picked_cabinet is None:
        if mode != "container_only":
            raise ValueError("No active 'cabinet' DC block template found in Excel (needed for hybrid/cabinet-only).")

    cont_code, cont_name, cont_unit = picked_container
    if picked_cabinet:
        cab_code, cab_name, cab_unit = picked_cabinet
    else:
        cab_code, cab_name, cab_unit = "", "", 0.0

    if mode == "container_only":
        s2 = build_config_container_only(required_dc, cont_unit, cont_code, cont_name)
    elif mode == "hybrid":
        s2 = build_config_hybrid(required_dc, cont_unit, cont_code, cont_name, cab_unit, cab_code, cab_name, k_max=k_max)
    elif mode == "cabinet_only":
        s2 = build_config_cabinet_only(required_dc, cab_unit, cab_code, cab_name, k_max=k_max)
    else:
        raise ValueError(f"Unknown mode: {mode}")

    s3_df, s3_meta = run_stage3(stage1, s2, df_soh_profile, df_soh_curve, df_rte_profile, df_rte_curve)

    def poi_at_year(df, year):
        row = df[df["Year_Index"] == int(year)]
        if row.empty:
            return None
        return float(row["POI_Usable_Energy_MWh"].iloc[0])

    poi_g = poi_at_year(s3_df, guarantee_year)
    if poi_g is None:
        return s2, s3_df, s3_meta, 1, None, False

    iter_count = 1
    converged = (poi_g + 1e-6 >= poi_energy_req)

    while not converged and iter_count < max_iter:
        if mode == "container_only":
            s2["container_count"] += 1
        elif mode == "cabinet_only":
            s2["cabinet_count"] += 1
        elif mode == "hybrid":
            if s2["cabinet_count"] < k_max and s2["cabinet_count"] > 0:
                s2["cabinet_count"] += 1
            elif s2["cabinet_count"] == 0:
                s2["cabinet_count"] = 1
            else:
                s2["container_count"] += 1

        rows = []
        if s2["container_count"] > 0:
            rows.append({
                "Block Code": cont_code,
                "Block Name": cont_name,
                "Form": "container",
                "Unit Capacity (MWh)": float(cont_unit),
                "Count": int(s2["container_count"]),
            })
        if mode in ("hybrid", "cabinet_only") and s2["cabinet_count"] > 0:
            rows.append({
                "Block Code": cab_code,
                "Block Name": cab_name,
                "Form": "cabinet",
                "Unit Capacity (MWh)": float(cab_unit),
                "Count": int(s2["cabinet_count"]),
            })

        df_cfg = _make_config_table(rows)
        s2["block_config_table"] = df_cfg
        total = float(df_cfg["Subtotal (MWh)"].sum()) if not df_cfg.empty else 0.0
        s2["dc_nameplate_bol_mwh"] = total
        s2["oversize_mwh"] = total - required_dc
        s2["config_adjustment_frac"] = (total / required_dc - 1.0) if required_dc > 0 else 0.0

        if mode == "cabinet_only":
            s2["busbars_needed"] = int(math.ceil(s2["cabinet_count"] / k_max)) if s2["cabinet_count"] > 0 else 0
        elif mode == "hybrid":
            s2["busbars_needed"] = 1 if s2["cabinet_count"] > 0 else 0
        else:
            s2["busbars_needed"] = 0

        s3_df, s3_meta = run_stage3(stage1, s2, df_soh_profile, df_soh_curve, df_rte_profile, df_rte_curve)
        poi_g = poi_at_year(s3_df, guarantee_year)

        iter_count += 1
        if poi_g is not None and (poi_g + 1e-6 >= poi_energy_req):
            converged = True

    return s2, s3_df, s3_meta, iter_count, poi_g, converged

# ==========================================
# 5. REPORT EXPORT HELPERS
# ==========================================
def find_logo_for_report():
    preferred_files = [
        DATA_DIR / DEFAULT_LOGO_NAME,
        SCRIPT_DIR / DEFAULT_LOGO_NAME,
        REPO_ROOT / DEFAULT_LOGO_NAME,
    ]

    for candidate in preferred_files:
        try:
            if candidate.is_file():
                return str(candidate.resolve())
        except Exception:
            continue

    search_dirs: list[Path] = []

    try:
        search_dirs.append(Path(DATA_FILE).resolve().parent)
    except Exception:
        pass

    for extra in [SCRIPT_DIR, REPO_ROOT, Path.cwd()] + list(Path.cwd().parents):
        if extra is not None:
            search_dirs.append(Path(extra))

    seen = set()
    for folder in search_dirs:
        try:
            folder_resolved = folder.resolve()
        except Exception:
            continue
        if folder_resolved in seen or not folder_resolved.is_dir():
            continue
        seen.add(folder_resolved)

        for fname in folder_resolved.iterdir():
            lower = fname.name.lower()
            if lower.endswith((".png", ".jpg", ".jpeg")) and ("logo" in lower or "calb" in lower):
                return str(fname)

    return None

def make_report_filename(project_name: str) -> str:
    base = (project_name or "CALB_ESS_Project").strip()
    if not base:
        base = "CALB_ESS_Project"
    safe = "".join(c if c.isalnum() or c in (" ", "_", "-") else "_" for c in base)
    safe = "_".join(safe.split())
    return f"{safe}_ESS_Sizing_Report.docx"

def _docx_add_config_table(doc: Document, df_conf: pd.DataFrame):
    if df_conf is None or df_conf.empty:
        doc.add_paragraph("No DC block configuration selected.", style="Intense Quote")
        return

    drop_cols = [
        "Block Code",
        "Form",
        "Config Adjustment (%)",
        "Busbars Needed (K=10)",
        "Oversize (MWh)",
        "Oversize_mwh",
        "Busbars Needed",
    ]
    df_show = df_conf.copy()
    cols_to_drop = [c for c in drop_cols if c in df_show.columns]
    if cols_to_drop:
        df_show = df_show.drop(columns=cols_to_drop)

    TOTAL_COL = "Total DC Nameplate @BOL (MWh)"
    has_total_col = TOTAL_COL in df_show.columns

    tbl = doc.add_table(rows=1, cols=len(df_show.columns))
    tbl.style = "Table Grid"

    hdr = tbl.rows[0].cells
    for j, col in enumerate(df_show.columns):
        hdr[j].text = str(col)
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for i, (_, row) in enumerate(df_show.iterrows()):
        rc = tbl.add_row().cells
        for j, col in enumerate(df_show.columns):
            if has_total_col and col == TOTAL_COL and i > 0:
                rc[j].text = ""
                continue

            val = row[col]
            if isinstance(val, float):
                rc[j].text = f"{val:.3f}"
            else:
                rc[j].text = f"{val}"

def _plot_poi_usable_png(s3_df: pd.DataFrame, poi_target: float, title: str) -> io.BytesIO:
    df = s3_df.sort_values("Year_Index").copy()
    x = df["Year_Index"].astype(int).tolist()
    y = df["POI_Usable_Energy_MWh"].astype(float).tolist()

    fig = plt.figure(figsize=(7.0, 3.2))
    ax = fig.add_subplot(111)
    ax.bar(x, y, color=CALB_SKY_BLUE)
    ax.axhline(poi_target, linewidth=2, color="#ff0000")
    ax.set_title(title)
    ax.set_xlabel("Year (from COD)")
    ax.set_ylabel("POI Usable Energy (MWh)")
    ax.set_xticks(x)
    ax.grid(True, axis="y", linestyle="--", alpha=0.35)

    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf

def _docx_add_lifetime_table(doc: Document, s3_df: pd.DataFrame):
    cols = [
        "Year_Index",
        "SOH_Display_Pct",
        "SOH_Absolute_Pct",
        "DC_Usable_MWh",
        "POI_Usable_Energy_MWh",
        "DC_RTE_Pct",
        "System_RTE_Pct",
    ]
    for c in cols:
        if c not in s3_df.columns:
            s3_df[c] = np.nan

    headers_map = {
        "Year_Index": "Year (From COD)",
        "SOH_Display_Pct": "SOH @ COD Baseline (%)",
        "SOH_Absolute_Pct": "SOH Vs FAT (%)",
        "DC_Usable_MWh": "DC Usable (MWh)",
        "POI_Usable_Energy_MWh": "POI Usable (MWh)",
        "DC_RTE_Pct": "DC RTE (%)",
        "System_RTE_Pct": "System RTE (%)",
    }

    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for j, c in enumerate(cols):
        hdr[j].text = headers_map[c]
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    df_sorted = s3_df.sort_values("Year_Index")
    for _, r in df_sorted.iterrows():
        rc = tbl.add_row().cells
        rc[0].text = str(int(r["Year_Index"]))
        rc[1].text = f"{r['SOH_Display_Pct']:.1f}" if not pd.isna(r["SOH_Display_Pct"]) else ""
        rc[2].text = f"{r['SOH_Absolute_Pct']:.1f}" if not pd.isna(r["SOH_Absolute_Pct"]) else ""
        rc[3].text = f"{r['DC_Usable_MWh']:.2f}" if not pd.isna(r["DC_Usable_MWh"]) else ""
        rc[4].text = f"{r['POI_Usable_Energy_MWh']:.2f}" if not pd.isna(r["POI_Usable_Energy_MWh"]) else ""
        rc[5].text = f"{r['DC_RTE_Pct']:.1f}%" if not pd.isna(r["DC_RTE_Pct"]) else ""
        rc[6].text = f"{r['System_RTE_Pct']:.1f}%" if not pd.isna(r["System_RTE_Pct"]) else ""

def build_report_bytes(stage1: dict, results_dict: dict, report_order: list):
    if not DOCX_AVAILABLE:
        return None

    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    logo_path = find_logo_for_report()
    header = section.header
    header.is_linked_to_previous = False
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.9))
    hdr_cells = header_table.rows[0].cells

    if logo_path:
        p_logo = hdr_cells[0].paragraphs[0]
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(1.2))

    p_info = hdr_cells[1].paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_info = p_info.add_run(
        "CALB Group Co., Ltd.\n"
        "Utility-Scale Energy Storage Systems\n"
        "Confidential Sizing Report"
    )
    run_info.font.size = Pt(9)
    run_info.font.bold = True

    project_name = stage1.get("project_name", "CALB ESS Project")

    doc.add_heading("CALB Utility-Scale ESS Sizing Report", level=1)
    doc.add_paragraph(f"Project: {project_name}")
    doc.add_paragraph(
        f"POI Requirement: {stage1['poi_power_req_mw']:.2f} MW / "
        f"{stage1['poi_energy_req_mwh']:.2f} MWh"
    )
    doc.add_paragraph("")

    doc.add_heading("1. Project Summary", level=2)
    p = doc.add_paragraph()
    p.add_run(f"Project life: {int(stage1['project_life_years'])} years\n")
    p.add_run(f"POI guarantee year: {int(stage1.get('poi_guarantee_year', 0))}\n")
    p.add_run(f"Cycles per year (assumed): {int(stage1['cycles_per_year'])}\n")
    p.add_run(f"S&C time from FAT to COD: {int(round(stage1.get('sc_time_months', 0)))} months\n")
    p.add_run(f"DC→POI efficiency chain (one-way): {stage1.get('eff_dc_to_poi_frac', 0.0)*100:.2f}%\n")
    p.add_run(f"POI→DC equivalent power: {stage1.get('dc_power_required_mw', 0.0):.2f} MW")

    doc.add_paragraph(
        "This sizing report is based on the 314 Ah cell database and the internal "
        "CALB SOH/RTE profiles for the selected operating conditions."
    )

    doc.add_heading("2. Equipment Summary (DC Blocks)", level=2)

    for key, title in report_order:
        if key not in results_dict:
            continue
        s2, _, _, iter_count, poi_g, converged = results_dict[key]
        doc.add_paragraph(title, style=None)
        _docx_add_config_table(doc, s2.get("block_config_table"))
        doc.add_paragraph(f"Iterations: {iter_count} | Guarantee met: {bool(converged)}")
        if poi_g is not None:
            doc.add_paragraph(f"POI usable energy @ guarantee year: {poi_g:.2f} MWh")
        doc.add_paragraph("")

    doc.add_heading("3. Lifetime POI Usable Energy & SOH (Per Configuration)", level=2)

    poi_target = float(stage1["poi_energy_req_mwh"])
    guarantee_year = int(stage1.get("poi_guarantee_year", 0))

    for key, title in report_order:
        if key not in results_dict:
            continue
        _, s3_df, s3_meta, _, _, _ = results_dict[key]

        doc.add_paragraph(title, style=None)
        doc.add_paragraph(
            f"POI Power = {s3_meta.get('poi_power_mw', 0.0):.2f} MW | "
            f"DC-equivalent Power = {s3_meta.get('dc_power_mw', 0.0):.2f} MW | "
            f"Effective C-rate (DC-side) = {s3_meta.get('effective_c_rate', 0.0):.3f} C"
        )
        doc.add_paragraph(
            f"SOH profile ID = {s3_meta.get('soh_profile_id')} "
            f"(C-rate ≈ {s3_meta.get('chosen_soh_c_rate')}, cycles/year = {s3_meta.get('chosen_soh_cycles_per_year')}); "
            f"RTE profile ID = {s3_meta.get('rte_profile_id')} (C-rate ≈ {s3_meta.get('chosen_rte_c_rate')})."
        )
        doc.add_paragraph(f"Guarantee Year (from COD) = {guarantee_year} | POI Energy Target = {poi_target:.2f} MWh")

        if MATPLOTLIB_AVAILABLE:
            try:
                png = _plot_poi_usable_png(
                    s3_df=s3_df,
                    poi_target=poi_target,
                    title=f"POI Usable Energy vs Year – {key}"
                )
                doc.add_picture(png, width=Inches(6.7))
            except Exception:
                doc.add_paragraph("Chart export skipped due to plotting error.")
        else:
            doc.add_paragraph("Chart export skipped (matplotlib not available).")

        _docx_add_lifetime_table(doc, s3_df)
        doc.add_paragraph("")

    p_final = doc.add_paragraph()
    p_fmt = p_final.paragraph_format
    p_fmt.space_before = Pt(0)
    p_fmt.space_after = Pt(0)
    p_fmt.line_spacing = Pt(0)
    run_final = p_final.add_run()
    run_final.font.size = Pt(1)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ==========================================
# 6. LOAD DATA
# ==========================================
try:
    defaults, df_blocks, df_soh_profile, df_soh_curve, df_rte_profile, df_rte_curve = load_data(DATA_FILE)
except Exception as exc:
    st.error(f"❌ Failed to load data file '{DATA_FILE}': {exc}")
    st.stop()

def get_default_numeric(field_name: str, fallback: float):
    raw = defaults.get(field_name, fallback)
    return to_float(raw, fallback)

def get_default_str(field_name: str, fallback: str):
    raw = defaults.get(field_name, fallback)
    if raw is None:
        return fallback
    return str(raw)

def get_default_percent_val(field_name: str, fallback: float) -> float:
    raw = defaults.get(field_name, fallback)
    val = to_float(raw, fallback)
    if val <= 1.5 and val > 0:
        return val * 100.0
    return val

NAV_OPTIONS = ["Stage 1–3 Inputs", "DC Block Results & Export"]


def render_results_view(payload: dict | None):
    if not payload:
        st.info("Run the sizing from the inputs tab to view DC Block results and export reports.")
        return

    s1 = payload["stage1"]
    results = payload["results"]
    modes_to_run = payload["modes_to_run"]
    poi_nominal_voltage_kv = float(payload.get("poi_nominal_voltage_kv", st.session_state.get("poi_nominal_voltage_kv", 22.0)))

    st.markdown("<div class='calb-card'>", unsafe_allow_html=True)
    st.subheader("3 · Stage 1 – DC Requirement @ FAT/BOL")

    m1, m2, m3, m4 = st.columns(4)
    with m1:
        st.markdown("<div class='metric-label'>Theoretical DC Required @FAT/BOL</div>", unsafe_allow_html=True)
        st.markdown(f"<div class='metric-value'>{s1['dc_energy_capacity_required_mwh']:.2f} MWh</div>", unsafe_allow_html=True)
    with m2:
        st.markdown("<div class='metric-label'>DC→POI Efficiency Chain</div>", unsafe_allow_html=True)
        st.markdown(f"<div class='metric-value'>{s1['eff_dc_to_poi_frac']*100:.2f} %</div>", unsafe_allow_html=True)
    with m3:
        st.markdown("<div class='metric-label'>DC Usable @BOL (DOD × √RTE)</div>", unsafe_allow_html=True)
        st.markdown(f"<div class='metric-value'>{s1['dc_usable_bol_frac']*100:.2f} %</div>", unsafe_allow_html=True)
    with m4:
        st.markdown("<div class='metric-label'>POI→DC Equivalent Power</div>", unsafe_allow_html=True)
        st.markdown(f"<div class='metric-value'>{s1['dc_power_required_mw']:.2f} MW</div>", unsafe_allow_html=True)

    st.write(
        f"S&C Time = **{int(round(s1['sc_time_months']))} months**, "
        f"S&C Loss = **{s1['sc_loss_pct']:.2f} %**, "
        f"POI Guarantee Year (from COD) = **{int(s1['poi_guarantee_year'])}**"
    )
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div class='calb-card'>", unsafe_allow_html=True)
    st.subheader("4 · Stage 2 & 3 – DC Block Configurations (Compare)")

    show_key = first_success_key(results, ["hybrid", "cabinet_only", "container_only"])
    if show_key:
        _, s3_df_ref, _, _, _, _ = results[show_key]
        st.markdown("**Pre-COD Storage & Commissioning Summary (FAT → COD)**")
        sc_m = int(round(s1["sc_time_months"]))
        year0_row = s3_df_ref[s3_df_ref["Year_Index"] == 0]
        if not year0_row.empty and "SOH_Absolute_Pct" in s3_df_ref.columns:
            soh_cod_abs = float(year0_row["SOH_Absolute_Pct"].iloc[0])
        else:
            soh_cod_abs = (1.0 - s1["sc_loss_frac"]) * 100.0

        pre_cod_df = pd.DataFrame(
            [
                {"Point": "FAT / BOL", "Timeline (From FAT)": "0 months", "SOH (% Of FAT)": "100.0", "Remark": "Factory-accepted capacity"},
                {"Point": "SAT / COD", "Timeline (From FAT)": f"{sc_m} months", "SOH (% Of FAT)": f"{soh_cod_abs:.1f}", "Remark": "Baseline before operation at site"},
            ]
        )
        st.table(pre_cod_df)

    def _sanitize_stage2_for_stage4(s2: dict) -> dict:
        if not isinstance(s2, dict):
            return {}
        out = {}
        for k in [
            "mode",
            "dc_nameplate_bol_mwh",
            "oversize_mwh",
            "config_adjustment_frac",
            "container_count",
            "cabinet_count",
            "busbars_needed",
        ]:
            if k in s2:
                out[k] = s2.get(k)

        bct = s2.get("block_config_table", None)
        try:
            if isinstance(bct, pd.DataFrame):
                out["block_config_table_records"] = bct.to_dict("records")
        except Exception:
            pass

        return out

    def _good_dc_split_for_acblock(dc_blocks: int, per_ac: int = 4) -> bool:
        if dc_blocks <= 0:
            return False
        r = int(dc_blocks) % int(per_ac)
        return r in (0, 2)

    available_stage4_sources = []
    for k in ["container_only", "hybrid", "cabinet_only"]:
        if isinstance(results.get(k), tuple) and results[k][0] != "ERROR":
            available_stage4_sources.append(k)

    if "stage4_source_user" not in st.session_state:
        st.session_state["stage4_source_user"] = "AUTO"

    if available_stage4_sources:
        opts = ["AUTO"] + available_stage4_sources
        if st.session_state["stage4_source_user"] not in opts:
            st.session_state["stage4_source_user"] = "AUTO"

        stage4_source_user = st.selectbox(
            "Stage 4 Source Scenario (DC Blocks)",
            options=opts,
            index=opts.index(st.session_state["stage4_source_user"]),
            help="AUTO defaults to Container-Only; if DC block qty causes an unbalanced AC block (e.g., 3 DC blocks), AUTO will fallback to Hybrid when available.",
        )
        st.session_state["stage4_source_user"] = stage4_source_user
    else:
        stage4_source_user = "AUTO"

    preferred_default = "container_only" if "container_only" in available_stage4_sources else (available_stage4_sources[0] if available_stage4_sources else None)
    stage4_source_mode = preferred_default

    if stage4_source_user != "AUTO":
        stage4_source_mode = stage4_source_user
    else:
        if preferred_default == "container_only":
            try:
                s2_c, *_ = results["container_only"]
                n_c = int(s2_c.get("container_count", 0))
            except Exception:
                n_c = 0

            if (not _good_dc_split_for_acblock(n_c, per_ac=4)) and ("hybrid" in available_stage4_sources):
                stage4_source_mode = "hybrid"

    stage13_output_all = {}
    try:
        for k in available_stage4_sources:
            s2_k, _s3_df_k, s3_meta_k, _iter_k, _poi_g_k, _conv_k = results[k]
            s2_clean = _sanitize_stage2_for_stage4(s2_k)
            dc_qty_k = int(s2_k.get("container_count", 0))
            stage13_output_all[k] = pack_stage13_output(
                stage1=s1,
                stage2=s2_clean,
                stage3=s3_meta_k,
                dc_block_total_qty=dc_qty_k,
                selected_scenario=k,
                poi_nominal_voltage_kv=poi_nominal_voltage_kv,
            )
    except Exception as e:
        st.warning(f"Stage 4 pre-pack failed: {e}")

    if stage4_source_mode and stage4_source_mode in stage13_output_all:
        st.session_state["stage13_output"] = stage13_output_all[stage4_source_mode]
        st.session_state["stage13_output_all"] = stage13_output_all
        st.session_state["stage4_source_mode"] = stage4_source_mode

        st.session_state["dc_block_total_qty"] = int(st.session_state["stage13_output"].get("dc_block_total_qty", 0))
        st.session_state["dc_block_container_count"] = int(st.session_state["stage13_output"].get("container_count", 0))
        st.session_state["dc_block_cabinet_count"] = int(st.session_state["stage13_output"].get("cabinet_count", 0))

        st.caption(
            f"Stage 4 input prepared: source = **{stage4_source_mode}**, "
            f"DC blocks (containers) = **{st.session_state['dc_block_total_qty']}**."
        )
    else:
        st.session_state.pop("stage13_output", None)
        st.session_state.pop("stage13_output_all", None)
        st.session_state.pop("dc_block_total_qty", None)
        st.session_state.pop("dc_block_container_count", None)
        st.session_state.pop("dc_block_cabinet_count", None)
        st.session_state.pop("stage4_source_mode", None)

    tab_info = []
    if "hybrid" in modes_to_run:
        tab_info.append(("hybrid", f"Hybrid (5MWh + 418kWh, K={K_MAX_FIXED})"))
    if "cabinet_only" in modes_to_run:
        tab_info.append(("cabinet_only", f"Cabinet Only (418kWh, K={K_MAX_FIXED})"))
    tab_info.append(("container_only", "Container Only (5MWh)"))

    tabs = st.tabs([t[1] for t in tab_info])

    def render_mode(tab, key, title):
        with tab:
            st.markdown(f"### {title}")
            if not isinstance(results.get(key), tuple):
                st.error("Unknown error.")
                return
            if results[key][0] == "ERROR":
                st.error(results[key][1])
                return

            s2, s3_df, s3_meta, iter_count, poi_g, converged = results[key]

            cA, cB, cC, cD = st.columns(4)
            with cA:
                st.markdown("<div class='metric-label'>Installed DC Nameplate @BOL</div>", unsafe_allow_html=True)
                st.markdown(f"<div class='metric-value'>{s2['dc_nameplate_bol_mwh']:.3f} MWh</div>", unsafe_allow_html=True)
            with cB:
                st.markdown("<div class='metric-label'>Oversize vs Stage1</div>", unsafe_allow_html=True)
                st.markdown(f"<div class='metric-value'>{s2['oversize_mwh']:.3f} MWh</div>", unsafe_allow_html=True)
            with cC:
                st.markdown("<div class='metric-label'>Container Count</div>", unsafe_allow_html=True)
                st.markdown(f"<div class='metric-value'>{int(s2.get('container_count', 0))}</div>", unsafe_allow_html=True)
            with cD:
                st.markdown("<div class='metric-label'>418kWh Cabinet Count</div>", unsafe_allow_html=True)
                st.markdown(f"<div class='metric-value'>{int(s2.get('cabinet_count', 0))}</div>", unsafe_allow_html=True)

            if key in ("hybrid", "cabinet_only"):
                st.write(f"Busbars Needed (K={K_MAX_FIXED}) = **{int(s2.get('busbars_needed', 0))}**")

            if key in ("hybrid", "cabinet_only") and int(s2.get("cabinet_count", 0)) > 0:
                cab_total = int(s2.get("cabinet_count", 0))
                groups = []
                if key == "hybrid":
                    groups = [cab_total]
                else:
                    full = cab_total // K_MAX_FIXED
                    rem = cab_total % K_MAX_FIXED
                    groups = [K_MAX_FIXED] * full + ([rem] if rem > 0 else [])
                busbar_df = pd.DataFrame(
                    [{"DC Busbar": i + 1, "418kWh Cabinets": g} for i, g in enumerate(groups)]
                )
                st.caption("DC busbar grouping preview (for design/check):")
                st.table(busbar_df)

            df_ui = s2["block_config_table"].copy()
            if "Total DC Nameplate @BOL (MWh)" in df_ui.columns:
                df_ui = df_ui.drop(columns=["Total DC Nameplate @BOL (MWh)"])
            st_dataframe_full_width(df_ui)

            if not converged:
                st.warning(f"⚠️ Guarantee not met after {iter_count} iterations. Consider augmentation or reduce guarantee year.")
            if poi_g is not None:
                margin = poi_g - s1["poi_energy_req_mwh"]
                st.write(
                    f"POI usable energy @ guarantee year (Year {int(s1['poi_guarantee_year'])}) "
                    f"= **{poi_g:.2f} MWh**, requirement = **{s1['poi_energy_req_mwh']:.2f} MWh**, "
                    f"margin = **{margin:.2f} MWh**, iterations = **{iter_count}**."
                )

            st.write(
                f"POI Power = **{s3_meta['poi_power_mw']:.2f} MW**, "
                f"DC-equivalent Power ≈ **{s3_meta['dc_power_mw']:.2f} MW**, "
                f"Effective C-rate (DC-side) ≈ **{s3_meta['effective_c_rate']:.3f} C**"
            )

            st.write(
                f"SOH profile ID **{s3_meta['soh_profile_id']}** "
                f"(C-rate ≈ {s3_meta['chosen_soh_c_rate']}, "
                f"Cycles/year = {s3_meta['chosen_soh_cycles_per_year']}), "
                f"RTE profile ID **{s3_meta['rte_profile_id']}** "
                f"(C-rate ≈ {s3_meta['chosen_rte_c_rate']})."
            )

            s3_df_sorted = s3_df.sort_values("Year_Index")
            poi_target = float(s1["poi_energy_req_mwh"])

            bars = (
                alt.Chart(s3_df_sorted)
                .mark_bar(color=CALB_SKY_BLUE)
                .encode(
                    x=alt.X("Year_Index:O", title="Year (from COD)"),
                    y=alt.Y("POI_Usable_Energy_MWh:Q", title="POI Usable Energy (MWh)"),
                    tooltip=[
                        alt.Tooltip("Year_Index:O", title="Year"),
                        alt.Tooltip("POI_Usable_Energy_MWh:Q", title="POI Usable (MWh)", format=".2f"),
                        alt.Tooltip("SOH_Display_Pct:Q", title="SOH @ COD Baseline (%)", format=".1f"),
                        alt.Tooltip("System_RTE_Pct:Q", title="System RTE (%)", format=".1f"),
                    ],
                )
            )

            rule_df = pd.DataFrame({"y": [poi_target]})
            rule = alt.Chart(rule_df).mark_rule(color="#ff0000", strokeWidth=2).encode(y="y:Q")

            text = (
                alt.Chart(s3_df_sorted)
                .mark_text(
                    align="center",
                    baseline="bottom",
                    dy=-3,
                    color=CHART_TEXT_COLOR,
                    font="Segoe UI",
                )
                .encode(
                    x=alt.X("Year_Index:O"),
                    y=alt.Y("POI_Usable_Energy_MWh:Q"),
                    text=alt.Text("POI_Usable_Energy_MWh:Q", format=".1f"),
                )
            )

            chart = (bars + rule + text).properties(height=360, title="System Capacity / POI Usable Chart (MWh)")
            st.altair_chart(chart, use_container_width=True)

            with st.expander("Show yearly data table", expanded=False):
                display_cols = [
                    "Year_Index",
                    "DC_Usable_MWh",
                    "POI_Usable_Energy_MWh",
                    "DC_RTE_Pct",
                    "System_RTE_Pct",
                    "Meets_POI_Req",
                    "Is_Guarantee_Year",
                ]
                for c in display_cols:
                    if c not in s3_df_sorted.columns:
                        s3_df_sorted[c] = np.nan

                disp = s3_df_sorted[display_cols].copy()
                disp = disp.rename(columns={
                    "Year_Index": "Year_Index",
                    "DC_Usable_MWh": "DC Usable (MWh)",
                    "POI_Usable_Energy_MWh": "POI Usable (MWh)",
                    "DC_RTE_Pct": "DC RTE (%)",
                    "System_RTE_Pct": "System RTE (%)",
                })

                disp["DC Usable (MWh)"] = disp["DC Usable (MWh)"].map(lambda x: f"{x:.4f}")
                disp["POI Usable (MWh)"] = disp["POI Usable (MWh)"].map(lambda x: f"{x:.4f}")
                disp["DC RTE (%)"] = disp["DC RTE (%)"].map(lambda x: f"{x:.2f}")
                disp["System RTE (%)"] = disp["System RTE (%)"].map(lambda x: f"{x:.2f}")

                st_dataframe_full_width(disp)

    for i, (key, label) in enumerate(tab_info):
        render_mode(tabs[i], key, label)

    st.markdown("---")

    report_buf = None
    file_name = None
    if DOCX_AVAILABLE:
        ok_results = {}
        for k in ["hybrid", "cabinet_only", "container_only"]:
            if isinstance(results.get(k), tuple) and results[k][0] != "ERROR":
                ok_results[k] = results[k]

        report_order = []
        for k, label in tab_info:
            if k in ok_results:
                report_order.append((k, label))

        report_buf = build_report_bytes(s1, ok_results, report_order)
        if report_buf is not None:
            file_name = make_report_filename(s1.get("project_name", "CALB_ESS_Project"))

    col_export, col_stage4 = st.columns([1, 1])

    with col_export:
        if DOCX_AVAILABLE and report_buf is not None and file_name is not None:
            st.download_button(
                label="📄 Export Technical Sizing Report",
                data=report_buf,
                file_name=file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        elif not DOCX_AVAILABLE:
            st.info("python-docx not installed; cannot export DOCX report in this environment.")
        else:
            st.error("⚠️ Report generation returned empty data. Please verify environment.")

    with col_stage4:
        stage4_path = "stage4_app.py" if os.path.exists("stage4_app.py") else None
        ready = ("stage13_output" in st.session_state) and (int(st.session_state.get("dc_block_total_qty", 0)) > 0)

        def _stage4_url_pathname(page_file: str) -> str | None:
            """
            Resolve a safe url_pathname for Streamlit page links.
            Prevents KeyError when page metadata is missing or incomplete.
            """
            if not hasattr(st, "experimental_get_pages"):
                return None
            try:
                pages = st.experimental_get_pages()
            except Exception:
                return None
            default_url = f"/{Path(page_file).stem}"
            for page_data in pages.values():
                if page_data.get("script_path") == page_file:
                    if "url_pathname" not in page_data or not page_data.get("url_pathname"):
                        # Initialize the link target to avoid KeyError downstream.
                        page_data["url_pathname"] = default_url
                    return page_data["url_pathname"]
            return None

        if stage4_path and hasattr(st, "page_link"):
            stage4_url = _stage4_url_pathname(stage4_path)
            if stage4_url:
                st.page_link(stage4_url, label="➡️ Open Stage 4 (AC Block Sizing)", disabled=not ready)
                if not ready:
                    st.caption("Run sizing first to enable Stage 4.")
            else:
                st.caption("Stage 4 (AC Block) runs via stage4_app.py in a separate Streamlit session.")
        else:
            st.caption("Stage 4 (AC Block) runs via stage4_app.py in a separate Streamlit session.")

    st.markdown("</div>", unsafe_allow_html=True)

# ==========================================
# 7. HEADER
# ==========================================
nav_default = st.session_state.get("dc_nav", NAV_OPTIONS[0])
if nav_default not in NAV_OPTIONS:
    nav_default = NAV_OPTIONS[0]
use_external_nav = bool(st.session_state.get("dc_nav_external", False))

if use_external_nav:
    nav_choice = nav_default
else:
    nav_choice = st.sidebar.radio(
        "Navigation",
        NAV_OPTIONS,
        index=NAV_OPTIONS.index(nav_default),
        help="Toggle between inputs and results on a single page.",
    )

st.session_state["dc_nav"] = nav_choice

st.markdown(
    """
    <div style="padding-top:0.6rem; padding-bottom:0.6rem;">
        <h1 class="calb-page-title">Utility-Scale ESS Sizing Tool V1.0</h1>
    </div>
    """,
    unsafe_allow_html=True,
)
st.markdown("<br/>", unsafe_allow_html=True)

# ==========================================
# 8. MAIN FORM + RESULTS
# ==========================================
run_btn = False
last_run = st.session_state.get("dc_last_run")

if nav_choice == NAV_OPTIONS[0]:
    with st.container():
        st.markdown("<div class='calb-card'>", unsafe_allow_html=True)
        st.subheader("1 · Project Inputs")

        with st.form("main_form"):
            project_name = st.text_input(
                "Project Name",
                value=get_default_str("project_name", "CALB ESS Project"),
            )

            c1, c2, c3 = st.columns(3)
            poi_power = c1.number_input(
                "POI Required Power (MW)",
                value=get_default_numeric("poi_power_req_mw", 100.0),
                min_value=0.0,
            )
            poi_energy = c2.number_input(
                "POI Required Capacity (MWh)",
                value=get_default_numeric("poi_energy_req_mwh", 400.0),
                min_value=0.0,
            )
            project_life = int(
                c3.number_input(
                    "Project Life (Years)",
                    value=int(get_default_numeric("project_life_years", 20)),
                    min_value=1,
                    step=1,
                    format="%d",
                )
            )

            poi_nominal_voltage_kv = st.number_input(
                "POI / MV Voltage (kV)",
                value=float(st.session_state.get("poi_nominal_voltage_kv", 22.0)),
                min_value=1.0,
                max_value=60.0,
                step=0.1,
                help="Typical MV range 10–35kV. Used for Stage 4 AC Block / MV equipment selection.",
            )

            c4, c5, c6 = st.columns(3)
            cycles_year = int(
                c4.number_input(
                    "Cycles Per Year",
                    value=int(get_default_numeric("cycles_per_year", 365)),
                    min_value=1,
                    step=1,
                    format="%d",
                )
            )
            guarantee_year = int(
                c5.number_input(
                    "POI Guarantee Year (Default @ COD )",
                    value=int(get_default_numeric("poi_guarantee_year", 0)),
                    min_value=0,
                    max_value=project_life,
                    step=1,
                    format="%d",
                )
            )

            def_sc = int(get_default_numeric("sc_time_months", 6.0))
            if def_sc < 3:
                def_sc = 3
            sc_time_months = int(
                c6.number_input(
                    "S&C Time (Months, From FAT To COD)",
                    value=def_sc,
                    min_value=3,
                    max_value=60,
                    step=1,
                    format="%d",
                )
            )

            st.markdown("---")
            st.subheader("2 · DC Parameters")

            c7, c8 = st.columns(2)
            dod_pct = c7.number_input(
                "DOD (%)",
                value=get_default_percent_val("dod_pct", 97.0),
                min_value=0.0,
                max_value=100.0,
            )
            dc_rte_pct = c8.number_input(
                "DC RTE (%) – Cell Only Base on Default Data",
                value=get_default_percent_val("dc_round_trip_efficiency_pct", 94.0),
                min_value=0.0,
                max_value=100.0,
            )

            st.info(f"Design Rule: Max 418kWh Cabinets per DC Busbar (K) = {K_MAX_FIXED} (fixed)")

            st.markdown("**3 · Configuration Options**")
            copt1, copt2, copt3 = st.columns([2, 2, 3])
            enable_hybrid = copt1.checkbox("Enable Hybrid Mode (5MWh + Cabinet)", value=True)
            enable_cabinet_only = copt2.checkbox("Enable Cabinet-Only Mode (Cabinet Modular Shipping)", value=True)
            hybrid_disable_threshold = copt3.number_input(
                "Disable Hybrid When POI Required Capacity ≥ (MWh) (0 = no limit)",
                value=9999.0,
                min_value=0.0,
                help="If POI required energy ≥ this threshold, Hybrid mode will be skipped. Set 0 for no limit.",
            )

            with st.expander(
                "Advanced: DC → POI @ MV Efficiency Chain. If the POI is located at the DC side, you may enable the checkbox below to force all efficiencies to 100%.",
                expanded=False
            ):

                poi_is_dc_side = st.checkbox(
                    "POI Is Located At DC Side (Force All Efficiencies To 100%)",
                    value=False,
                    help="When enabled, the DC→POI efficiency chain is bypassed (all set to 100%).",
                )

                c10, c11, c12 = st.columns(3)
                eff_dc_cables = c10.number_input(
                    "DC Cables Efficiency (%)",
                    value=get_default_percent_val("eff_dc_cables", 99.50),
                    min_value=0.0,
                    max_value=100.0,
                )
                eff_pcs = c11.number_input(
                    "PCS Efficiency (%)",
                    value=get_default_percent_val("eff_pcs", 98.50),
                    min_value=0.0,
                    max_value=100.0,
                )
                eff_mvt = c12.number_input(
                    "MV Transformer Efficiency (%)",
                    value=get_default_percent_val("eff_mvt", 99.50),
                    min_value=0.0,
                    max_value=100.0,
                )
                c13, c14 = st.columns(2)
                eff_ac_sw = c13.number_input(
                    "AC Cables + SW/RMU Efficiency (%)",
                    value=get_default_percent_val("eff_ac_cables_sw_rmu", 99.20),
                    min_value=0.0,
                    max_value=100.0,
                )
                eff_hvt = c14.number_input(
                    "Other--HVT Efficiency (%)",
                    value=get_default_percent_val("eff_hvt_others", 100.0),
                    min_value=0.0,
                    max_value=100.0,
                )

            st.markdown("---")
            run_btn = st.form_submit_button("🔄 Run Sizing")

        st.markdown("</div>", unsafe_allow_html=True)

    if run_btn:
        st.session_state["poi_nominal_voltage_kv"] = float(poi_nominal_voltage_kv)

        if poi_is_dc_side:
            eff_dc_cables = 100.0
            eff_pcs = 100.0
            eff_mvt = 100.0
            eff_ac_sw = 100.0
            eff_hvt = 100.0

        inputs = {
            "project_name": project_name,
            "poi_power_req_mw": poi_power,
            "poi_energy_req_mwh": poi_energy,
            "project_life_years": project_life,
            "cycles_per_year": cycles_year,
            "poi_guarantee_year": guarantee_year,
            "sc_time_months": sc_time_months,
            "dod_pct": dod_pct,
            "dc_round_trip_efficiency_pct": dc_rte_pct,
            "eff_dc_cables": eff_dc_cables,
            "eff_pcs": eff_pcs,
            "eff_mvt": eff_mvt,
            "eff_ac_cables_sw_rmu": eff_ac_sw,
            "eff_hvt_others": eff_hvt,
        }

        s1 = run_stage1(inputs, defaults)

        modes_to_run = ["container_only"]
        if enable_cabinet_only:
            modes_to_run.insert(0, "cabinet_only")
        if enable_hybrid:
            if not (hybrid_disable_threshold > 0 and poi_energy >= hybrid_disable_threshold):
                modes_to_run.insert(0, "hybrid")

        results = {}
        for mode in modes_to_run:
            try:
                results[mode] = size_with_guarantee(
                    s1, mode,
                    df_blocks,
                    df_soh_profile, df_soh_curve,
                    df_rte_profile, df_rte_curve,
                    k_max=K_MAX_FIXED
                )
            except Exception as e:
                results[mode] = ("ERROR", str(e))

        last_run = {
            "stage1": s1,
            "results": results,
            "modes_to_run": modes_to_run,
            "poi_nominal_voltage_kv": poi_nominal_voltage_kv,
        }
        st.session_state["dc_last_run"] = last_run
        st.session_state["dc_nav"] = NAV_OPTIONS[1]
        render_results_view(last_run)

elif nav_choice == NAV_OPTIONS[1]:
    render_results_view(last_run)
